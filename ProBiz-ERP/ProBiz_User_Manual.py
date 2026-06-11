from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import docx.opc.constants

doc = Document()

# ── Page margins ──
section = doc.sections[0]
section.page_width  = Inches(8.27)   # A4
section.page_height = Inches(11.69)
section.left_margin   = Cm(2.5)
section.right_margin  = Cm(2.5)
section.top_margin    = Cm(2.5)
section.bottom_margin = Cm(2.5)

# ── Styles helper ──
def heading1(text):
    p = doc.add_heading(text, level=1)
    p.runs[0].font.color.rgb = RGBColor(0x1e, 0x40, 0xaf)
    return p

def heading2(text):
    p = doc.add_heading(text, level=2)
    p.runs[0].font.color.rgb = RGBColor(0x1e, 0x40, 0xaf)
    return p

def heading3(text):
    p = doc.add_heading(text, level=3)
    if p.runs:
        p.runs[0].font.color.rgb = RGBColor(0x1e, 0x3a, 0x8a)
    return p

def para(text, bold=False, italic=False, size=11):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    return p

def bullet(text):
    p = doc.add_paragraph(text, style='List Bullet')
    p.runs[0].font.size = Pt(11)
    return p

def numbered(text):
    p = doc.add_paragraph(text, style='List Number')
    p.runs[0].font.size = Pt(11)
    return p

def add_table(headers, rows, col_widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        hdr[i].paragraphs[0].runs[0].bold = True
        hdr[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(0xff, 0xff, 0xff)
        tc = hdr[i]._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), '1E40AF')
        tcPr.append(shd)
    for ri, row in enumerate(rows):
        cells = table.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = str(val)
            cells[i].paragraphs[0].runs[0].font.size = Pt(10)
            if ri % 2 == 1:
                tc = cells[i]._tc
                tcPr = tc.get_or_add_tcPr()
                shd = OxmlElement('w:shd')
                shd.set(qn('w:val'), 'clear')
                shd.set(qn('w:color'), 'auto')
                shd.set(qn('w:fill'), 'EFF6FF')
                tcPr.append(shd)
    if col_widths:
        for i, row in enumerate(table.rows):
            for j, cell in enumerate(row.cells):
                if j < len(col_widths):
                    cell.width = Cm(col_widths[j])
    return table

# ════════════════════════════════════════════════════════
# COVER PAGE
# ════════════════════════════════════════════════════════
cover = doc.add_paragraph()
cover.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = cover.add_run('\n\n\n')

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = title.add_run('ProBiz ERP')
r.bold = True
r.font.size = Pt(36)
r.font.color.rgb = RGBColor(0x1e, 0x40, 0xaf)

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
r2 = subtitle.add_run('User Manual')
r2.font.size = Pt(22)
r2.font.color.rgb = RGBColor(0x47, 0x55, 0x69)

doc.add_paragraph()
divider = doc.add_paragraph()
divider.alignment = WD_ALIGN_PARAGRAPH.CENTER
r3 = divider.add_run('─' * 40)
r3.font.color.rgb = RGBColor(0x1e, 0x40, 0xaf)

doc.add_paragraph()
sub2 = doc.add_paragraph()
sub2.alignment = WD_ALIGN_PARAGRAPH.CENTER
r4 = sub2.add_run('Complete Guide to Using ProBiz ERP\nFor Pharmacies, Hospitals & Business Owners')
r4.font.size = Pt(13)
r4.font.color.rgb = RGBColor(0x64, 0x74, 0x8b)

doc.add_paragraph('\n\n')
contact = doc.add_paragraph()
contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
r5 = contact.add_run('Dr Muhammad Raees RPh\n0316-8818693  |  raees.malik89@gmail.com\nIslamabad, Pakistan')
r5.font.size = Pt(11)
r5.font.color.rgb = RGBColor(0x1e, 0x40, 0xaf)

doc.add_page_break()

# ════════════════════════════════════════════════════════
# TABLE OF CONTENTS (manual)
# ════════════════════════════════════════════════════════
heading1('Table of Contents')
toc_items = [
    ('1', 'Getting Started', '3'),
    ('2', 'Dashboard', '5'),
    ('3', 'Inventory Management', '6'),
    ('4', 'Sales & Point of Sale (POS)', '9'),
    ('5', 'Purchase Management', '12'),
    ('6', 'Customer Management', '14'),
    ('7', 'Supplier Management', '15'),
    ('8', 'Accounting Module', '16'),
    ('9', 'HR & Payroll', '18'),
    ('10', 'Reports & Analytics', '20'),
    ('11', 'Settings', '22'),
    ('12', 'Subscription & Plans', '24'),
    ('13', 'Security & User Roles', '25'),
    ('14', 'Frequently Asked Questions', '26'),
    ('15', 'Contact & Support', '28'),
]
for num, title_text, page in toc_items:
    p = doc.add_paragraph()
    p.add_run(f'{num}.  {title_text}').font.size = Pt(11)
    p.paragraph_format.space_after = Pt(4)

doc.add_page_break()

# ════════════════════════════════════════════════════════
# 1. GETTING STARTED
# ════════════════════════════════════════════════════════
heading1('1. Getting Started')
para('Welcome to ProBiz ERP — a complete cloud-based business management system designed for Pakistani pharmacies, hospitals, cash & carry stores and wholesale businesses. This manual will guide you through every feature of the system.')

heading2('1.1  Accessing the System')
para('Open your web browser (Google Chrome recommended) and go to:')
p = doc.add_paragraph()
r = p.add_run('https://probiz-erp-poru.vercel.app')
r.bold = True
r.font.color.rgb = RGBColor(0x1e, 0x40, 0xaf)
r.font.size = Pt(12)
para('The system works on any device — desktop, laptop, tablet or mobile phone. No installation is required.')

heading2('1.2  Logging In')
numbered('Open the URL above in your browser.')
numbered('Enter your Email address and Password.')
numbered('Click the Login button.')
numbered('You will be taken to the Dashboard.')
para('If you forget your password, contact your system administrator or WhatsApp 0316-8818693.')

heading2('1.3  Login Credentials')
add_table(
    ['Role', 'Email', 'Password', 'Access Level'],
    [
        ['Superadmin', 'raees.malik89@gmail.com', 'admin123', 'Full access — all modules'],
        ['Admin', 'admin@probiz.pk', 'admin123', 'All modules except system settings'],
        ['Manager', 'manager@probiz.pk', 'manager123', 'Inventory, Sales, Purchases, Reports'],
        ['Cashier', 'cashier@probiz.pk', 'cashier123', 'Sales / POS only'],
        ['Accountant', 'accountant@probiz.pk', 'acc123', 'Accounting and financial reports'],
    ]
)
doc.add_paragraph()
para('⚠️  Important: Change your password immediately after first login. Go to Settings → Change Password.', bold=True)

heading2('1.4  First-Time Setup (5 Steps)')
numbered('Log in as Superadmin.')
numbered('Go to Settings → update Company Name, Address, Phone — this prints on all invoices.')
numbered('Go to Inventory → add your products with cost price, sale price and opening stock.')
numbered('Go to Customers → add your regular customers with credit limits.')
numbered('Go to Suppliers → add your medicine/product suppliers.')
para('After these 5 steps, you are ready to start making sales.')

heading2('1.5  Important Tips')
bullet('Use Google Chrome or Microsoft Edge for the best experience.')
bullet('Bookmark the URL so you can access it quickly.')
bullet('The system auto-saves — do not close the browser while saving.')
bullet('All data is stored in the cloud — closing the browser does not lose any data.')
bullet('Multiple staff can log in at the same time from different devices.')

doc.add_page_break()

# ════════════════════════════════════════════════════════
# 2. DASHBOARD
# ════════════════════════════════════════════════════════
heading1('2. Dashboard')
para('The Dashboard is the first screen you see after login. It gives you a complete overview of your business in real time.')

heading2('2.1  KPI Cards')
add_table(
    ['Card', 'What It Shows'],
    [
        ["Today's Sales", "Total revenue from all completed sales today"],
        ['Month Revenue', 'Total sales for the current calendar month'],
        ['Total Products', 'Number of products in inventory + low-stock alert count'],
        ['Customers', 'Total customer accounts with total outstanding receivables (Rs.)'],
        ['Suppliers', 'Total suppliers with total outstanding payables (Rs.)'],
        ['Employees', 'Number of active staff members'],
    ]
)

doc.add_paragraph()
heading2('2.2  Charts')
bullet('Sales — Last 7 Days: Line/area chart showing daily sales trend for the past week.')
bullet('Top Products: Pie chart of best-selling products by revenue this month.')
bullet('Sales vs Purchases (6 Months): Bar chart comparing monthly sales and purchase totals.')
bullet('Recent Sales: Table of the latest invoices with customer name, amount and status.')

heading2('2.3  Low Stock Alerts')
para('When any product falls below its minimum stock level, a yellow warning panel appears at the bottom of the Dashboard. It shows the product name, current stock quantity, and the minimum threshold. Use this to reorder before you run out of stock.')

doc.add_page_break()

# ════════════════════════════════════════════════════════
# 3. INVENTORY
# ════════════════════════════════════════════════════════
heading1('3. Inventory Management')
para('The Inventory module lets you manage all your products — medicines, surgical items, cosmetics, vitamins, or any other goods you sell.')

heading2('3.1  Adding a New Product')
numbered('Go to Inventory from the left sidebar.')
numbered('Click the "Add Product" button (top right corner).')
numbered('Fill in all required fields (marked with *).')
numbered('For pharmacy products: fill in Batch No., Mfg. Date and Expiry Date.')
numbered('Click Save. The product is now available in POS and Purchase screens.')

heading2('3.2  Product Fields Explained')
add_table(
    ['Field', 'Description', 'Example'],
    [
        ['Name *', 'Full product name', 'Panadol 500mg Tablets'],
        ['SKU', 'Unique product code', 'MED-001'],
        ['Barcode', 'Product barcode (optional, must be unique)', '6281234567890'],
        ['Category', 'Product group for filtering', 'Medicines'],
        ['Unit *', 'How the product is sold', 'strips / pcs / kg / box / bottle'],
        ['Cost Price *', 'Price you paid the supplier (Rs.)', 'Rs. 80.00'],
        ['Sale Price *', 'Price you charge the customer (Rs.)', 'Rs. 100.00'],
        ['Stock *', 'Current quantity in stock', '150'],
        ['Min Stock', 'Alert threshold for low-stock warning', '20'],
        ['Batch No.', 'Manufacturer batch/lot number', 'B-2024-P01'],
        ['Mfg. Date', 'Manufacturing date', '2024-01-15'],
        ['Expiry Date', 'Expiry date — shown in RED on invoices', '2026-06-01'],
    ]
)

doc.add_paragraph()
heading2('3.3  Batch & Expiry Tracking (Pharmacy Feature)')
para('This is a special feature for pharmacies and medicine stores:')
bullet('Every product can have a Batch Number, Manufacturing Date, and Expiry Date.')
bullet('These details appear automatically on every printed invoice.')
bullet('Expiry dates are displayed in RED BOLD on invoices as a safety alert to the customer.')
bullet('Helps with medicine recalls — find which customers received a specific batch.')
bullet('Keeps you compliant with DRAP (Drug Regulatory Authority Pakistan) requirements.')

heading2('3.4  Editing a Product')
numbered('Find the product in the Inventory list (use search bar to find quickly).')
numbered('Click the Edit button (pencil icon) on the product row.')
numbered('Update any field you need to change.')
numbered('Click Save.')

heading2('3.5  How Stock Updates Automatically')
add_table(
    ['Action', 'Effect on Stock'],
    [
        ['Sale completed', 'Stock decreases by sold quantity instantly'],
        ['Purchase received', 'Stock increases by received quantity instantly'],
        ['Manual adjustment', 'Edit the product and change the Stock field directly'],
    ]
)

doc.add_paragraph()
heading2('3.6  Searching Products')
para('Use the search bar at the top of the Inventory page. You can search by product name, SKU, or category. Results appear instantly as you type.')

heading2('3.7  Categories')
para('Organize products into categories such as Medicines, Surgical, Cosmetics, Vitamins, Baby Care, etc. Categories help you filter the product list and appear in sales reports so you can see which product group is most profitable.')

doc.add_page_break()

# ════════════════════════════════════════════════════════
# 4. SALES / POS
# ════════════════════════════════════════════════════════
heading1('4. Sales & Point of Sale (POS)')
para('The Sales module is where you record every sale made to customers and generate professional invoices.')

heading2('4.1  Creating a New Sale')
numbered('Go to Sales / POS from the sidebar.')
numbered('Click "New Sale" button.')
numbered('Select the customer from the dropdown (or choose "Walk-in Customer" for cash sales).')
numbered('Search for a product and click Add, then enter the quantity.')
numbered('Repeat for all products in the sale.')
numbered('Enter a discount amount if applicable.')
numbered('Select the payment method.')
numbered('Click "Complete Sale".')
para('The invoice is generated instantly with a unique invoice number.')

heading2('4.2  Payment Methods')
add_table(
    ['Payment Method', 'When to Use', 'Customer Balance After Sale'],
    [
        ['Cash', 'Customer pays in cash immediately', 'Rs. 0 — fully paid'],
        ['Card', 'Customer pays by debit or credit card', 'Rs. 0 — fully paid'],
        ['Bank Transfer', 'Customer transfers via online banking', 'Rs. 0 — fully paid'],
        ['Credit', 'Customer takes goods now and pays later', 'Full amount added to their balance'],
    ]
)

doc.add_paragraph()
heading2('4.3  Invoice Number Format')
para('Every invoice gets a unique number automatically:')
p = doc.add_paragraph()
r = p.add_run('INV-YYYYMM-NNNNN   (Example: INV-202606-00001)')
r.bold = True
r.font.size = Pt(12)
r.font.color.rgb = RGBColor(0x1e, 0x40, 0xaf)
para('Numbers are sequential and never repeat — required for proper accounting records.')

heading2('4.4  Printing an Invoice')
numbered('After completing a sale, click the Print icon (🖨️) on the invoice.')
numbered('A print-optimized invoice opens automatically.')
numbered('Press Ctrl+P (Windows) or Cmd+P (Mac).')
numbered('Select your printer or choose "Save as PDF".')
numbered('Click Print.')
para('The printed invoice includes:')
bullet('Business name, address and phone number')
bullet('Invoice number, date and customer details')
bullet('Product table with: Name, Batch No., Mfg. Date, Expiry Date, Qty, Unit Price, Total')
bullet('Discount and grand total')
bullet('Three signature lines: Pharmacist / Licensed Pharmacist, Manager / Authorized Signatory, Customer / Received By')

heading2('4.5  Credit Sales')
para('When you select Credit as the payment method:')
bullet("The full invoice amount is added to the customer's outstanding balance.")
bullet("The system checks the customer's credit limit — if exceeded, a warning appears.")
bullet('You can still complete the sale by overriding the warning if needed.')
bullet("When the customer pays later, record the payment in the Customers section to reduce their balance.")

heading2('4.6  What Happens Automatically on Each Sale')
bullet('Stock of each sold product decreases immediately.')
bullet('Invoice saved with all line items, prices, discounts and totals.')
bullet('Credit sale: customer receivable balance increases.')
bullet('Profit margin calculated: (Sale Price − Cost Price) × Quantity.')
bullet('Accounting journal entry created automatically.')
bullet('Dashboard charts and KPI cards update.')

doc.add_page_break()

# ════════════════════════════════════════════════════════
# 5. PURCHASES
# ════════════════════════════════════════════════════════
heading1('5. Purchase Management')
para('Record all purchases from your suppliers — medicines, surgical items, or any goods you buy to resell.')

heading2('5.1  Creating a Purchase Order')
numbered('Go to Purchases from the sidebar.')
numbered('Click "New Purchase".')
numbered('Select the supplier from the dropdown.')
numbered('Add products and enter the quantity being received.')
numbered('Enter the cost price for each item (may differ from the stored cost price).')
numbered('Select payment method: Cash (paid now) or Credit (payable later).')
numbered('Add notes if needed (e.g. supplier invoice number, delivery date).')
numbered('Click Save.')

heading2('5.2  What Happens Automatically')
bullet('Stock of each purchased product increases immediately.')
bullet('Credit purchase: supplier payable balance increases.')
bullet('Purchase gets a unique PO number (PO-YYYYMM-NNNNN).')
bullet('Accounting journal entry created: Inventory Dr / Cash or Payable Cr.')

heading2('5.3  Viewing Purchase History')
para('All purchase orders are listed with date, supplier name, total amount, and payment status. Use the date filter or supplier dropdown to find any historical purchase quickly.')

doc.add_page_break()

# ════════════════════════════════════════════════════════
# 6. CUSTOMERS
# ════════════════════════════════════════════════════════
heading1('6. Customer Management')
para('Maintain customer accounts, track credit balances and view purchase history.')

heading2('6.1  Adding a Customer')
numbered('Go to Customers from the sidebar.')
numbered('Click "Add Customer".')
numbered('Fill in: Name, Phone, Email (optional), City.')
numbered('Set Credit Limit (Rs.) — maximum credit you allow this customer. Set 0 for no credit.')
numbered('Click Save.')

heading2('6.2  Customer Fields')
add_table(
    ['Field', 'Description', 'Example'],
    [
        ['Name', 'Full name or business name', 'Ahmad Medical Store'],
        ['Phone', 'Contact number', '0321-1234567'],
        ['Email', 'Email address (optional)', 'ahmad@gmail.com'],
        ['City', 'City for records and delivery', 'Islamabad'],
        ['Credit Limit', 'Maximum credit allowed (Rs.)', 'Rs. 50,000'],
        ['Balance', 'Outstanding amount owed — auto-calculated', 'Rs. 12,500'],
    ]
)

doc.add_paragraph()
heading2('6.3  Managing Credit')
bullet("When a credit sale is made, the customer's balance increases automatically.")
bullet("When the customer pays, record the payment to reduce their balance.")
bullet("If a new credit sale would exceed their limit, a warning is shown.")
bullet("The Receivables Report shows all customers with outstanding balances.")

heading2('6.4  Walk-in Customers')
para('For cash sales to customers you do not need to track, select "Walk-in Customer" in POS. No account is needed. The sale records as a cash transaction with no receivable balance.')

doc.add_page_break()

# ════════════════════════════════════════════════════════
# 7. SUPPLIERS
# ════════════════════════════════════════════════════════
heading1('7. Supplier Management')
para('Track all your suppliers, purchase history and how much you owe each one.')

heading2('7.1  Adding a Supplier')
numbered('Go to Suppliers from the sidebar.')
numbered('Click "Add Supplier".')
numbered('Fill in: Company Name, Contact Person, Phone, Email, City.')
numbered('Add any payment terms or notes.')
numbered('Click Save.')

heading2('7.2  Payables Management')
para('Every credit purchase from a supplier increases their payable balance automatically. When you pay the supplier, record the payment to reduce the balance. The Suppliers list shows the total payable for each supplier so you always know who to pay.')

doc.add_page_break()

# ════════════════════════════════════════════════════════
# 8. ACCOUNTING
# ════════════════════════════════════════════════════════
heading1('8. Accounting Module')
para('ProBiz uses professional double-entry bookkeeping. Every transaction automatically creates the correct journal entries — you do not need to do this manually.')

heading2('8.1  What is Double-Entry Accounting?')
para('Every financial transaction affects two accounts — one is debited and one is credited for the same amount. This ensures the books always balance. Example: when you make a cash sale of Rs. 1,000 — Cash account is debited Rs. 1,000 AND Sales Revenue is credited Rs. 1,000.')

heading2('8.2  Account Types')
add_table(
    ['Type', 'Examples', 'Normal Balance'],
    [
        ['Asset', 'Cash, Bank, Inventory, Accounts Receivable', 'Debit'],
        ['Liability', 'Accounts Payable, Loans Payable', 'Credit'],
        ['Equity', "Owner's Equity, Retained Earnings", 'Credit'],
        ['Income', 'Sales Revenue, Other Income', 'Credit'],
        ['Expense', 'Cost of Goods Sold, Salaries, Rent', 'Debit'],
    ]
)

doc.add_paragraph()
heading2('8.3  Automatic Journal Entries')
add_table(
    ['Transaction', 'Debit Account', 'Credit Account'],
    [
        ['Cash Sale Rs. 1,000', 'Cash Rs. 1,000', 'Sales Revenue Rs. 1,000'],
        ['Credit Sale Rs. 1,000', 'Accounts Receivable Rs. 1,000', 'Sales Revenue Rs. 1,000'],
        ['Cash Purchase Rs. 500', 'Inventory Rs. 500', 'Cash Rs. 500'],
        ['Credit Purchase Rs. 500', 'Inventory Rs. 500', 'Accounts Payable Rs. 500'],
        ['Customer Payment Rs. 1,000', 'Cash Rs. 1,000', 'Accounts Receivable Rs. 1,000'],
        ['Supplier Payment Rs. 500', 'Accounts Payable Rs. 500', 'Cash Rs. 500'],
    ]
)

doc.add_paragraph()
heading2('8.4  Manual Journal Entries')
para('For adjustments, depreciation, owner drawings or any non-standard transactions, go to Accounting → Journal Entries → New Entry. Select the debit account, credit account, amount and description.')

heading2('8.5  Financial Reports')
bullet('Trial Balance — all account balances side by side to verify books balance.')
bullet('Profit & Loss Statement — total income minus total expenses for any period.')
bullet('Balance Sheet — snapshot of all assets, liabilities and equity at any date.')
bullet('General Ledger — complete transaction history for any account.')

doc.add_page_break()

# ════════════════════════════════════════════════════════
# 9. HR & PAYROLL
# ════════════════════════════════════════════════════════
heading1('9. HR & Payroll')
para('Manage your employees, process monthly salaries and generate payslips.')

heading2('9.1  Adding an Employee')
numbered('Go to HR & Payroll from the sidebar.')
numbered('Click "Add Employee".')
numbered('Enter: Full Name, Employee ID, Department, Designation, Phone, CNIC.')
numbered('Set Basic Salary, Allowances and Deductions.')
numbered('Assign to a branch and set Join Date.')
numbered('Mark as Active and Save.')

heading2('9.2  Salary Structure')
add_table(
    ['Component', 'Description', 'Example'],
    [
        ['Basic Salary', 'Fixed monthly base pay', 'Rs. 25,000'],
        ['Allowances', 'House rent, transport, medical, other benefits', 'Rs. 8,000'],
        ['Deductions', 'EOBI, income tax, loan repayment', 'Rs. 1,200'],
        ['Net Salary', 'Basic + Allowances − Deductions', 'Rs. 31,800'],
    ]
)

doc.add_paragraph()
heading2('9.3  Processing Monthly Payroll')
numbered('Go to HR & Payroll → click "Process Payroll".')
numbered('Select the month and year.')
numbered('System calculates net salary for every active employee automatically.')
numbered('Review the summary — make any manual adjustments if needed.')
numbered('Confirm — payslips are generated for all employees.')
numbered('Print payslips for distribution to staff.')

heading2('9.4  Attendance')
para('Mark daily attendance for each employee: Present, Absent, Half Day, or Leave. Attendance records are stored per employee per day and can be used to calculate salary deductions for absences.')

doc.add_page_break()

# ════════════════════════════════════════════════════════
# 10. REPORTS
# ════════════════════════════════════════════════════════
heading1('10. Reports & Analytics')
para('Generate any business report instantly for any date range.')

heading2('10.1  Available Reports')
add_table(
    ['Report Name', 'What It Shows'],
    [
        ['Sales Report', 'All invoices — customer, items, amount, payment method, status'],
        ['Purchase Report', 'All purchase orders — supplier, items, amounts, dates'],
        ['Inventory Report', 'Current stock levels, values, low stock warnings'],
        ['Profit & Loss', 'Total revenue − costs = gross profit and net profit'],
        ['Balance Sheet', 'Assets, liabilities and equity at any date'],
        ['Summary Report', 'Quick overview: sales, purchases, profit in one view'],
        ['Receivables', 'Customers with outstanding balances'],
        ['Payables', 'Suppliers with outstanding amounts owed'],
        ['Payroll Report', 'Monthly salary summary per employee'],
        ['Top Products', 'Best selling products by quantity and revenue'],
    ]
)

doc.add_paragraph()
heading2('10.2  How to Generate a Report')
numbered('Go to Reports from the sidebar.')
numbered('Click on the report tab you need.')
numbered('Set the From Date and To Date.')
numbered('Click Generate or Apply Filter.')
numbered('The report appears on screen instantly.')
numbered('Click Print to print or save as PDF.')

heading2('10.3  Profit & Loss Statement')
para('The P&L shows:')
bullet('Total Revenue — sum of all sales for the selected period')
bullet('Cost of Goods Sold — total cost price of all items sold')
bullet('Gross Profit — Revenue minus Cost of Goods Sold')
bullet('Operating Expenses — salaries, rent, utilities etc.')
bullet('Net Profit — Gross Profit minus all expenses')

heading2('10.4  Printing Reports')
para('Every report has a Print button. Click it and press Ctrl+P to print or save as PDF. The print layout is clean — it hides the sidebar and navigation buttons for a professional output.')

doc.add_page_break()

# ════════════════════════════════════════════════════════
# 11. SETTINGS
# ════════════════════════════════════════════════════════
heading1('11. Settings')
para('Configure your company information, change your password, and manage system preferences.')

heading2('11.1  Company Information')
para('Go to Settings and update:')
bullet('Company Name — appears on every invoice')
bullet('Company Address — appears on invoices and reports')
bullet('Phone Number — appears on invoices')
bullet('Email Address — for contact purposes')
para('Click "Save All Settings" when done.')

heading2('11.2  Changing Your Password')
numbered('Go to Settings from the sidebar.')
numbered('Scroll down to the "Change Password" section.')
numbered('Enter your Current Password.')
numbered('Enter New Password (minimum 6 characters).')
numbered('Enter Confirm New Password (must match exactly).')
numbered('Click "Change Password".')
para('A success message confirms the change. Use the new password next time you log in.')

heading2('11.3  Financial Settings')
bullet('Default Tax Rate (%) — applied to invoices if GST is applicable.')
bullet('Fiscal Year Start — January, July, or April.')
bullet('Invoice Prefix — customize invoice numbering (default: INV-).')
bullet('PO Prefix — customize purchase order prefix (default: PO-).')

heading2('11.4  User Roles & Permissions')
add_table(
    ['Permission', 'Superadmin', 'Admin', 'Manager', 'Cashier', 'Accountant'],
    [
        ['View Dashboard', '✅', '✅', '✅', '✅', '✅'],
        ['Manage Inventory', '✅', '✅', '✅', '❌', '❌'],
        ['Process Sales', '✅', '✅', '✅', '✅', '❌'],
        ['Manage Purchases', '✅', '✅', '✅', '❌', '❌'],
        ['View Accounting', '✅', '✅', '❌', '❌', '✅'],
        ['Manage HR/Payroll', '✅', '✅', '❌', '❌', '❌'],
        ['System Settings', '✅', '❌', '❌', '❌', '❌'],
    ]
)

doc.add_paragraph()
heading2('11.5  Reset to Demo Data')
para('⚠️  WARNING: This deletes ALL data (sales, inventory, customers, purchases, employees, accounting) and replaces with a small demo dataset. Use ONLY for fresh setup or testing. Your admin login is preserved. This action CANNOT be undone.')

doc.add_page_break()

# ════════════════════════════════════════════════════════
# 12. SUBSCRIPTION & PLANS
# ════════════════════════════════════════════════════════
heading1('12. Subscription & Plans')

heading2('12.1  Free Trial')
para('Every new ProBiz installation starts with a 14-day free trial. You get full access to all features with no limitations. The trial banner at the top of the screen shows how many days remain. When the trial ends, a paywall appears and you need to activate a paid plan to continue.')

heading2('12.2  Available Plans')
add_table(
    ['Plan', 'Price', 'Branches', 'Users', 'Best For'],
    [
        ['Free Trial', 'Free / 14 days', '1', '3', 'Evaluation'],
        ['Basic', 'Rs. 2,500/month', '1', '3', 'Single pharmacy or small shop'],
        ['Professional', 'Rs. 5,000/month', '3', '10', 'Clinic + pharmacy, small chain'],
        ['Enterprise', 'Rs. 15,000/month', 'Unlimited', 'Unlimited', 'Hospital, large wholesale'],
    ]
)

doc.add_paragraph()
heading2('12.3  How to Activate a Plan')
numbered('Contact us: WhatsApp 0316-8818693 or raees.malik89@gmail.com.')
numbered('Tell us which plan you want and for how many months.')
numbered('Make payment via JazzCash / EasyPaisa / Bank Transfer.')
numbered('Send payment screenshot on WhatsApp.')
numbered('We activate your plan within minutes — system unlocks immediately.')

doc.add_page_break()

# ════════════════════════════════════════════════════════
# 13. SECURITY
# ════════════════════════════════════════════════════════
heading1('13. Security & User Roles')

heading2('13.1  User Roles')
add_table(
    ['Role', 'Who It Is For', 'Key Access'],
    [
        ['Superadmin', 'Business owner', 'Full access — all modules, settings, subscriptions'],
        ['Admin', 'General manager', 'All modules except system settings'],
        ['Manager', 'Branch manager', 'Inventory, sales, purchases, HR, reports'],
        ['Cashier', 'Counter staff', 'Create sales and view invoices only'],
        ['Accountant', 'Finance staff', 'Accounting module and financial reports only'],
    ]
)

doc.add_paragraph()
heading2('13.2  Security Features')
bullet('Passwords are hashed with bcrypt — never stored in plain text.')
bullet('JWT tokens expire after 1 hour — abandoned sessions auto-close.')
bullet('All API endpoints require a valid login token.')
bullet('Role checks are enforced on the server — cannot be bypassed.')

heading2('13.3  Best Practices')
bullet('Change default passwords immediately after setup.')
bullet('Create individual accounts for each staff member — do not share logins.')
bullet('Assign the lowest role needed — cashiers do not need admin access.')
bullet('Log out when leaving your computer unattended.')

doc.add_page_break()

# ════════════════════════════════════════════════════════
# 14. FAQ
# ════════════════════════════════════════════════════════
heading1('14. Frequently Asked Questions')

faqs = [
    ('Can I use ProBiz on my mobile phone?',
     'Yes. ProBiz ERP works on any browser including Chrome on Android and Safari on iPhone. Desktop gives the best experience for data entry.'),
    ('What happens if I close the browser accidentally?',
     'Nothing is lost. All data is saved to the cloud instantly when you click Save or Complete Sale. Log back in and everything is exactly as you left it.'),
    ('Can multiple staff use the system at the same time?',
     'Yes. Multiple users can be logged in simultaneously from different devices with no conflicts.'),
    ('Why am I getting a "barcode unique constraint" error?',
     'Each product barcode must be unique. If you see this error, either use a different barcode or leave the barcode field empty — the system works fine without barcodes.'),
    ('How do I print an invoice?',
     'Open the sale record, click the print icon (🖨️), then press Ctrl+P to print or save as PDF.'),
    ('How do I record a payment from a customer?',
     'Go to Customers → find the customer → open their record → click Record Payment → enter amount and date.'),
    ('Can I give a discount on a sale?',
     'Yes. When creating a sale in POS, enter the discount amount in the Discount field before completing the sale.'),
    ('What happens when my trial expires?',
     'A paywall screen appears. All your data is preserved — nothing is deleted. Contact us on 0316-8818693 to activate a paid plan. Access is restored within minutes after payment.'),
    ('How do I add my real products and remove demo data?',
     'Go to Settings → Reset to Clean Demo → confirm. This wipes demo records. Then add your real products via Inventory → Add Product.'),
    ('Can I export reports to Excel?',
     'Currently reports can be printed as PDF. Full Excel/CSV export is coming in the next version. Contact us if you need a data export urgently.'),
    ('Is my data backed up?',
     'Data is stored on PythonAnywhere cloud servers. Contact us to set up automatic daily backups for extra safety.'),
    ('How do I change my password?',
     'Go to Settings → Change Password section → enter current password → enter new password twice → click Change Password.'),
]

for q, a in faqs:
    heading3(q)
    para(a)
    doc.add_paragraph()

doc.add_page_break()

# ════════════════════════════════════════════════════════
# 15. CONTACT & SUPPORT
# ════════════════════════════════════════════════════════
heading1('15. Contact & Support')

heading2('15.1  Contact Information')
add_table(
    ['Channel', 'Details', 'Best For'],
    [
        ['WhatsApp', '0316-8818693', 'Fastest — plan activation, urgent issues'],
        ['Phone Call', '0316-8818693', 'Complex questions, training'],
        ['Email', 'raees.malik89@gmail.com', 'Non-urgent queries, data export'],
        ['Office', 'House 124, Street 39, I-14/3, Islamabad', 'In-person training and setup'],
    ]
)

doc.add_paragraph()
heading2('15.2  Support Hours')
bullet('Monday – Saturday: 9:00 AM – 9:00 PM (Pakistan Standard Time)')
bullet('Sunday: 10:00 AM – 6:00 PM')
bullet('Emergency support for paid plans: WhatsApp any time')

heading2('15.3  What We Help With')
bullet('Plan activation and upgrades')
bullet('Initial setup and data migration from Excel')
bullet('Staff training — in-person or via Zoom/WhatsApp video call')
bullet('Custom reports or features for your specific business')
bullet('Technical issues and bug fixes')
bullet('Bulk product import from your existing inventory records')
bullet('Adding your business logo to invoices')

heading2('15.4  How to Report a Problem')
numbered('Take a screenshot of the error or unexpected behavior.')
numbered('Note what you were doing when the problem occurred.')
numbered('WhatsApp the screenshot to 0316-8818693.')
numbered('Include your business name and the time of the problem.')
para('Most issues are resolved within 2–4 hours during support hours.')

# ════════════════════════════════════════════════════════
# SAVE
# ════════════════════════════════════════════════════════
output_path = r'C:\Users\Acer\ProBiz-ERP\ProBiz_ERP_User_Manual.docx'
doc.save(output_path)
print('Word document saved: ' + output_path)
