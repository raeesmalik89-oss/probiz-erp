from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

section = doc.sections[0]
section.page_width  = Inches(8.27)
section.page_height = Inches(11.69)
section.left_margin = section.right_margin = Cm(2.5)
section.top_margin  = section.bottom_margin = Cm(2.5)

# ── Helpers ──────────────────────────────────────────────────────────────────
def h1(text):
    p = doc.add_heading(text, level=1)
    if p.runs: p.runs[0].font.color.rgb = RGBColor(0x1e,0x40,0xaf)
    return p

def h2(text):
    p = doc.add_heading(text, level=2)
    if p.runs: p.runs[0].font.color.rgb = RGBColor(0x1e,0x40,0xaf)
    return p

def h3(text):
    p = doc.add_heading(text, level=3)
    if p.runs: p.runs[0].font.color.rgb = RGBColor(0x0f,0x28,0x6e)
    return p

def para(text, bold=False, size=11):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = bold
    r.font.size = Pt(size)
    return p

def bullet(text):
    p = doc.add_paragraph(text, style='List Bullet')
    p.runs[0].font.size = Pt(11)

def numbered(text):
    p = doc.add_paragraph(text, style='List Number')
    p.runs[0].font.size = Pt(11)

def code(text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.name = 'Courier New'
    r.font.size = Pt(9)
    r.font.color.rgb = RGBColor(0x1e,0x29,0x3b)
    p.paragraph_format.left_indent = Cm(1)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(2)
    tc = p._p
    pPr = tc.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'),'clear'); shd.set(qn('w:color'),'auto'); shd.set(qn('w:fill'),'F1F5F9')
    pPr.append(shd)
    return p

def tbl(headers, rows):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Table Grid'
    hdr = table.rows[0].cells
    for i,h in enumerate(headers):
        hdr[i].text = h
        hdr[i].paragraphs[0].runs[0].bold = True
        hdr[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(0xff,0xff,0xff)
        hdr[i].paragraphs[0].runs[0].font.size = Pt(10)
        tcPr = hdr[i]._tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'),'clear'); shd.set(qn('w:color'),'auto'); shd.set(qn('w:fill'),'1E40AF')
        tcPr.append(shd)
    for ri,row in enumerate(rows):
        cells = table.add_row().cells
        for i,val in enumerate(row):
            cells[i].text = str(val)
            cells[i].paragraphs[0].runs[0].font.size = Pt(10)
            if ri % 2 == 1:
                tcPr = cells[i]._tc.get_or_add_tcPr()
                shd = OxmlElement('w:shd')
                shd.set(qn('w:val'),'clear'); shd.set(qn('w:color'),'auto'); shd.set(qn('w:fill'),'EFF6FF')
                tcPr.append(shd)
    doc.add_paragraph()

def note(text):
    p = doc.add_paragraph()
    r = p.add_run('NOTE:  ' + text)
    r.font.size = Pt(10)
    r.font.color.rgb = RGBColor(0xd9,0x77,0x06)
    r.bold = True
    p.paragraph_format.left_indent = Cm(0.5)

# ════════════════════════════════════════════════════════════════════════════
# COVER
# ════════════════════════════════════════════════════════════════════════════
for _ in range(4): doc.add_paragraph()

t = doc.add_paragraph()
t.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = t.add_run('ProBiz ERP')
r.bold = True; r.font.size = Pt(38)
r.font.color.rgb = RGBColor(0x1e,0x40,0xaf)

t2 = doc.add_paragraph()
t2.alignment = WD_ALIGN_PARAGRAPH.CENTER
r2 = t2.add_run('Technical Development & Sales Guide')
r2.font.size = Pt(20); r2.font.color.rgb = RGBColor(0x47,0x55,0x69)

doc.add_paragraph()
div = doc.add_paragraph()
div.alignment = WD_ALIGN_PARAGRAPH.CENTER
div.add_run('-' * 50).font.color.rgb = RGBColor(0x1e,0x40,0xaf)

doc.add_paragraph()
sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
r3 = sub.add_run(
    'How to Build a SaaS ERP System from Scratch\n'
    'and How to Sell It as a Profitable Software Business'
)
r3.font.size = Pt(13); r3.font.color.rgb = RGBColor(0x64,0x74,0x8b)

for _ in range(3): doc.add_paragraph()
c = doc.add_paragraph()
c.alignment = WD_ALIGN_PARAGRAPH.CENTER
r4 = c.add_run('Dr Muhammad Raees RPh\n0316-8818693  |  raees.malik89@gmail.com')
r4.font.size = Pt(11); r4.font.color.rgb = RGBColor(0x1e,0x40,0xaf)

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════════
# TABLE OF CONTENTS
# ════════════════════════════════════════════════════════════════════════════
h1('Table of Contents')
toc = [
    ('PART A — BUILDING THE SOFTWARE',''),
    ('1','Complete Tech Stack & Why Each Tool Was Chosen','3'),
    ('2','Project Folder Structure','5'),
    ('3','Database Design (All Tables)','6'),
    ('4','Backend — FastAPI (Python)','9'),
    ('5','Frontend — React','13'),
    ('6','Authentication System (JWT)','16'),
    ('7','SaaS Subscription System','18'),
    ('8','Deployment — PythonAnywhere + Vercel','20'),
    ('9','How to Build a Similar ERP from Scratch','22'),
    ('PART B — SELLING THE SOFTWARE',''),
    ('10','SaaS Business Model Explained','26'),
    ('11','Pricing Strategy','27'),
    ('12','Target Customers & Industries','28'),
    ('13','Sales & Marketing Strategy','29'),
    ('14','How to Onboard a New Client','31'),
    ('15','Common Objections & How to Handle Them','32'),
    ('16','Scaling the Business','33'),
]
for item in toc:
    if len(item) == 2:
        p = doc.add_paragraph()
        r = p.add_run(item[0])
        r.bold = True; r.font.size = Pt(12)
        r.font.color.rgb = RGBColor(0x1e,0x40,0xaf)
        p.paragraph_format.space_before = Pt(8)
    else:
        p = doc.add_paragraph()
        p.add_run(f'{item[0]}.  {item[1]}').font.size = Pt(11)
        p.paragraph_format.space_after = Pt(3)

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════════
# PART A HEADER
# ════════════════════════════════════════════════════════════════════════════
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('PART A — BUILDING THE SOFTWARE')
r.bold = True; r.font.size = Pt(16)
r.font.color.rgb = RGBColor(0xff,0xff,0xff)
pPr = p._p.get_or_add_pPr()
shd = OxmlElement('w:shd')
shd.set(qn('w:val'),'clear'); shd.set(qn('w:color'),'auto'); shd.set(qn('w:fill'),'1E3A8A')
pPr.append(shd)
doc.add_paragraph()

# ════════════════════════════════════════════════════════════════════════════
# 1. TECH STACK
# ════════════════════════════════════════════════════════════════════════════
h1('1.  Complete Tech Stack & Why Each Tool Was Chosen')
para(
    'ProBiz ERP is a full-stack web application built with modern, lightweight, free/open-source tools. '
    'Every tool was chosen for a specific reason — especially to keep hosting costs near zero while '
    'delivering a professional product.'
)

h2('1.1  Backend — FastAPI (Python)')
para('FastAPI is a modern Python web framework for building REST APIs.')
tbl(
    ['Why FastAPI','Detail'],
    [
        ['Speed','One of the fastest Python frameworks — comparable to Node.js'],
        ['Auto documentation','/docs endpoint auto-generates interactive API docs (Swagger UI)'],
        ['Type safety','Python type hints catch bugs before runtime'],
        ['Easy to learn','Simple syntax — less code than Django or Flask'],
        ['Pydantic','Built-in data validation — request bodies are auto-validated'],
        ['Async support','Handles many concurrent requests efficiently'],
    ]
)

h2('1.2  Database — SQLite / PostgreSQL via SQLAlchemy')
para('SQLite is used for development and small deployments. PostgreSQL for production scale.')
tbl(
    ['Tool','Purpose'],
    [
        ['SQLite','Zero-config file database — perfect for single-tenant small business ERP'],
        ['PostgreSQL','Production database for multi-tenant or high-traffic deployments'],
        ['SQLAlchemy','Python ORM — write Python classes instead of raw SQL'],
        ['Alembic','Database migration tool — safe schema changes without losing data'],
    ]
)

h2('1.3  Frontend — React 18')
tbl(
    ['Library / Tool','Purpose'],
    [
        ['React 18','UI framework — component-based, fast re-rendering'],
        ['React Router v6','Client-side routing — /app/dashboard, /app/sales etc.'],
        ['Axios','HTTP client — makes API calls to the backend'],
        ['Recharts','Charts library — sales graphs, pie charts, bar charts'],
        ['React Hot Toast','Toast notifications (success/error messages)'],
        ['Lucide React','Icon library — all icons in the sidebar and UI'],
        ['date-fns','Date formatting and manipulation'],
    ]
)

h2('1.4  Authentication — JWT')
tbl(
    ['Tool','Purpose'],
    [
        ['python-jose','JWT token creation and verification (backend)'],
        ['passlib + bcrypt','Password hashing — never store plain text passwords'],
        ['localStorage','Frontend stores access_token and refresh_token'],
        ['Axios interceptor','Auto-attaches token to every request; auto-refreshes on 401'],
    ]
)

h2('1.5  Hosting — Zero Cost Stack')
tbl(
    ['Service','What Runs There','Monthly Cost'],
    [
        ['PythonAnywhere (Free)','FastAPI backend + SQLite database','Rs. 0 (free tier)'],
        ['Vercel (Free)','React frontend — global CDN','Rs. 0 (free tier)'],
        ['PythonAnywhere (Paid)','Upgraded backend for production clients','~$5/month'],
        ['Vercel Pro','Custom domains, more bandwidth','~$20/month'],
    ]
)
note('Free tier is sufficient for a single client. Upgrade only when revenue justifies it.')

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════════
# 2. FOLDER STRUCTURE
# ════════════════════════════════════════════════════════════════════════════
h1('2.  Project Folder Structure')
para('Understanding where every file lives is essential before you build or modify anything.')

h2('2.1  Root Structure')
code('ProBiz-ERP/')
code('  frontend/          <- React app (what users see in the browser)')
code('  backend/           <- FastAPI Python app (the brain / API server)')
code('  ProBiz_ERP_User_Manual.docx')
code('  ProBiz_Technical_Doc.docx')

h2('2.2  Backend Structure')
code('backend/')
code('  main.py            <- App entry point. Registers all routers.')
code('  models.py          <- ALL database table definitions (SQLAlchemy models)')
code('  database.py        <- DB connection setup (engine, session, get_db)')
code('  auth.py            <- JWT logic: create_token, verify_token, get_current_user')
code('  config.py          <- App settings: SECRET_KEY, DATABASE_URL, APP_NAME')
code('  requirements.txt   <- Python packages to install')
code('  routers/')
code('    auth.py          <- /api/auth/* : login, register, change-password, refresh')
code('    inventory.py     <- /api/inventory/* : products, categories')
code('    sales.py         <- /api/sales/* : invoices, POS')
code('    purchases.py     <- /api/purchases/* : purchase orders')
code('    accounting.py    <- /api/accounting/* : journal entries, accounts')
code('    payroll.py       <- /api/payroll/* : employees, payslips')
code('    dashboard.py     <- /api/dashboard/* : stats, charts data')
code('    admin.py         <- /api/admin/* : reset demo data')
code('    subscription.py  <- /api/subscription/* : trial/plan management')
code('    demo_requests.py <- /api/demo/* : landing page leads')

h2('2.3  Frontend Structure')
code('frontend/src/')
code('  App.js             <- Root component. Defines all routes.')
code('  index.js           <- React entry point')
code('  api/')
code('    client.js        <- Axios instance with base URL + auth interceptors')
code('  contexts/')
code('    AuthContext.js   <- React context for login state (user, logout)')
code('  components/')
code('    Layout.js        <- Sidebar + header wrapper for all /app/* pages')
code('    SubscriptionGuard.js  <- Trial banner + paywall screen')
code('  pages/')
code('    Landing.js       <- Public marketing page + demo request form')
code('    Login.js         <- Login form')
code('    Dashboard.js     <- KPI cards + charts')
code('    Inventory.js     <- Product table + add/edit modal')
code('    Sales.js         <- Invoice list + new sale + print invoice')
code('    Purchases.js     <- Purchase orders')
code('    Customers.js     <- Customer accounts')
code('    Suppliers.js     <- Supplier accounts')
code('    Accounting.js    <- Journal entries + accounts')
code('    Payroll.js       <- Employees + payslips')
code('    Reports.js       <- All reports with tabs')
code('    Settings.js      <- Company settings + admin tools')
code('    Docs.js          <- In-app documentation')

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════════
# 3. DATABASE DESIGN
# ════════════════════════════════════════════════════════════════════════════
h1('3.  Database Design (All Tables)')
para(
    'The database uses SQLAlchemy ORM. Every table is defined as a Python class in models.py. '
    'SQLAlchemy creates the tables automatically when the app starts. '
    'All tables are in a single SQLite file: probiz.db'
)

h2('3.1  Complete Table List')
tbl(
    ['Table Name','Purpose','Key Fields'],
    [
        ['users','Login accounts with roles','id, name, email, hashed_password, role, branch_id, is_active'],
        ['branches','Business locations / branches','id, name, city, address, phone'],
        ['categories','Product categories','id, name, description'],
        ['products','Inventory items','id, name, sku, barcode, category_id, unit, cost_price, sale_price, stock, min_stock, batch_no, mfg_date, expiry_date'],
        ['customers','Customer accounts','id, name, phone, email, city, credit_limit, balance'],
        ['suppliers','Supplier accounts','id, company_name, contact_person, phone, email, balance'],
        ['sales','Invoice headers','id, invoice_number, customer_id, branch_id, payment_method, discount, total, status, created_at'],
        ['sale_items','Invoice line items','id, sale_id, product_id, quantity, unit_price, total'],
        ['purchases','Purchase order headers','id, po_number, supplier_id, payment_method, total, status'],
        ['purchase_items','PO line items','id, purchase_id, product_id, quantity, cost_price, total'],
        ['accounts','Chart of accounts','id, name, account_type, code, balance'],
        ['journal_entries','Double-entry transactions','id, description, debit_account_id, credit_account_id, amount, date'],
        ['departments','Employee departments','id, name, branch_id'],
        ['employees','Staff records','id, name, employee_id, department_id, designation, basic_salary, allowances, deductions'],
        ['attendances','Daily attendance','id, employee_id, date, status (present/absent/half_day)'],
        ['payslips','Monthly payroll records','id, employee_id, month, year, net_salary, status'],
        ['subscriptions','SaaS plan status (1 row)','id, plan, status, trial_end, paid_until, max_users'],
        ['demo_requests','Landing page leads','id, name, business, phone, email, status'],
        ['notifications','System alerts','id, user_id, title, message, is_read'],
        ['audit_logs','Action history','id, user_id, action, table_name, record_id, created_at'],
    ]
)

h2('3.2  How to Define a Table (models.py Pattern)')
code('from sqlalchemy import Column, Integer, String, Float, Boolean, DateTime, ForeignKey')
code('from sqlalchemy.orm import relationship')
code('from sqlalchemy.sql import func')
code('from database import Base')
code('')
code('class Product(Base):')
code('    __tablename__ = "products"')
code('    id         = Column(Integer, primary_key=True, index=True)')
code('    name       = Column(String(200), nullable=False)')
code('    sku        = Column(String(100), nullable=True, unique=True)')
code('    cost_price = Column(Float, default=0.0)')
code('    sale_price = Column(Float, default=0.0)')
code('    stock      = Column(Float, default=0.0)')
code('    batch_no   = Column(String(100), nullable=True)')
code('    expiry_date= Column(String(20),  nullable=True)')
code('    created_at = Column(DateTime(timezone=True), server_default=func.now())')
code('    # Relationship:')
code('    category   = relationship("Category", back_populates="products")')

h2('3.3  Adding a New Column to an Existing Table')
para('If the app is already running with data, use ALTER TABLE (not create_all):')
code('import sqlite3')
code("conn = sqlite3.connect('/home/Raees1989/probiz/backend/probiz.db')")
code("conn.execute('ALTER TABLE products ADD COLUMN new_field TEXT')")
code('conn.commit()')
code('conn.close()')
para('Then add the same column to the Python model class in models.py, clear .pyc cache, and reload.')

h2('3.4  Two Databases Gotcha')
para('PythonAnywhere has TWO database files. Only ONE is used by the web app:')
tbl(
    ['File','Used By','Notes'],
    [
        ['probiz.db','Web app (live)','Set in WSGI file via DATABASE_URL env var'],
        ['probiz_erp.db','Not used','Default in config.py — IGNORE this file'],
    ]
)
note('Always edit probiz.db. The WSGI file sets: os.environ["DATABASE_URL"] = "sqlite:////home/Raees1989/probiz/backend/probiz.db"')

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════════
# 4. BACKEND — FASTAPI
# ════════════════════════════════════════════════════════════════════════════
h1('4.  Backend — FastAPI (Python)')

h2('4.1  main.py — The Entry Point')
para('main.py is the first file FastAPI reads. It:')
bullet('Creates the FastAPI app instance')
bullet('Adds CORS middleware (allows the React frontend to call the API)')
bullet('Calls create_all() to create any missing tables on startup')
bullet('Registers all routers (each router = one module like sales, inventory etc.)')
code('from fastapi import FastAPI')
code('from fastapi.middleware.cors import CORSMiddleware')
code('from database import engine')
code('import models')
code('from routers import auth, inventory, sales, subscription')
code('')
code('models.Base.metadata.create_all(bind=engine)  # Create tables if missing')
code('')
code('app = FastAPI(title="ProBiz ERP")')
code('app.add_middleware(CORSMiddleware, allow_origins=["*"],')
code('                   allow_methods=["*"], allow_headers=["*"],')
code('                   allow_credentials=True)')
code('')
code('app.include_router(auth.router)')
code('app.include_router(inventory.router)')
code('app.include_router(sales.router)')
code('app.include_router(subscription.router)')

h2('4.2  database.py — DB Connection')
code('from sqlalchemy import create_engine')
code('from sqlalchemy.ext.declarative import declarative_base')
code('from sqlalchemy.orm import sessionmaker')
code('import os')
code('')
code('DATABASE_URL = os.environ.get("DATABASE_URL", "sqlite:///./probiz_erp.db")')
code('engine = create_engine(DATABASE_URL, connect_args={"check_same_thread": False})')
code('SessionLocal = sessionmaker(autocommit=False, autoflush=False, bind=engine)')
code('Base = declarative_base()')
code('')
code('def get_db():')
code('    db = SessionLocal()')
code('    try:    yield db')
code('    finally: db.close()')

h2('4.3  A Typical Router File Pattern')
para('Every router follows the same structure:')
code('from fastapi import APIRouter, Depends, HTTPException')
code('from sqlalchemy.orm import Session')
code('from pydantic import BaseModel')
code('from database import get_db')
code('from auth import get_current_user')
code('import models')
code('')
code('router = APIRouter(prefix="/api/inventory", tags=["inventory"])')
code('')
code('# --- Pydantic schema (validates request body) ---')
code('class ProductIn(BaseModel):')
code('    name: str')
code('    cost_price: float')
code('    sale_price: float')
code('    stock: float = 0')
code('')
code('# --- GET all products ---')
code('@router.get("/products")')
code('def get_products(db: Session = Depends(get_db),')
code('                 current_user = Depends(get_current_user)):')
code('    return db.query(models.Product).filter(models.Product.is_active==True).all()')
code('')
code('# --- POST create product ---')
code('@router.post("/products")')
code('def create_product(data: ProductIn,')
code('                   db: Session = Depends(get_db),')
code('                   current_user = Depends(get_current_user)):')
code('    product = models.Product(**data.dict())')
code('    db.add(product)')
code('    db.commit()')
code('    db.refresh(product)')
code('    return product')

h2('4.4  All API Endpoints Overview')
tbl(
    ['Method','Endpoint','What It Does','Auth Required'],
    [
        ['POST','/api/auth/login','Log in, get JWT tokens','No'],
        ['POST','/api/auth/refresh','Get new access token','No (refresh token)'],
        ['POST','/api/auth/change-password','Change own password','Yes'],
        ['GET','/api/auth/me','Get logged-in user info','Yes'],
        ['GET','/api/inventory/products','List all products','Yes'],
        ['POST','/api/inventory/products','Create new product','Yes'],
        ['PUT','/api/inventory/products/{id}','Update product','Yes'],
        ['DELETE','/api/inventory/products/{id}','Delete product','Yes'],
        ['GET','/api/sales/','List all invoices','Yes'],
        ['POST','/api/sales/','Create new sale/invoice','Yes'],
        ['GET','/api/purchases/','List all purchases','Yes'],
        ['POST','/api/purchases/','Create new purchase','Yes'],
        ['GET','/api/dashboard/stats','Dashboard KPI numbers','Yes'],
        ['GET','/api/subscription/status','Get trial/plan status','Yes'],
        ['POST','/api/subscription/activate','Activate paid plan','Admin only'],
        ['POST','/api/demo/request','Submit demo request','No (public)'],
        ['GET','/api/demo/requests','List all demo requests','Admin only'],
        ['POST','/api/admin/reset-demo','Reset to demo data','Superadmin only'],
    ]
)

h2('4.5  How a Sale is Processed (Backend Logic)')
para('When POST /api/sales/ is called:')
numbered('Validate request body (Pydantic model checks all fields).')
numbered('Verify JWT token — get current user from DB.')
numbered('Create Sale record (auto-generate invoice number INV-YYYYMM-NNNNN).')
numbered('Loop through each item: create SaleItem, decrease product.stock.')
numbered('If payment_method == "credit": increase customer.balance.')
numbered('Create accounting journal entry automatically.')
numbered('Commit all changes to DB in one transaction.')
numbered('Return the saved sale object with generated invoice number.')
code('# Invoice number generation example:')
code('from datetime import datetime')
code('def generate_invoice_number(db):')
code('    now = datetime.now()')
code('    prefix = f"INV-{now.year}{now.month:02d}-"')
code('    count = db.query(models.Sale).filter(')
code('        models.Sale.invoice_number.like(prefix + "%")).count()')
code('    return f"{prefix}{(count+1):05d}"')

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════════
# 5. FRONTEND — REACT
# ════════════════════════════════════════════════════════════════════════════
h1('5.  Frontend — React')

h2('5.1  How the App is Structured')
para(
    'The React app uses a single-page application (SPA) architecture. There is only ONE HTML file '
    '(index.html). React Router handles navigation — when you click "Inventory" in the sidebar, '
    'the URL changes to /app/inventory and React renders the Inventory component without a full page reload.'
)

h2('5.2  App.js — Route Definitions')
code('// Every page route is defined here')
code('<Routes>')
code('  <Route path="/" element={<Landing />} />')
code('  <Route path="/login" element={<Login />} />')
code('  <Route path="/app" element={<PrivateRoute><Layout /></PrivateRoute>}>')
code('    <Route path="dashboard"  element={<Dashboard />} />')
code('    <Route path="inventory"  element={<Inventory />} />')
code('    <Route path="sales"      element={<Sales />} />')
code('    <Route path="reports"    element={<Reports />} />')
code('    ...etc')
code('  </Route>')
code('</Routes>')
para('PrivateRoute checks if user is logged in. If not, redirects to /login.')

h2('5.3  API Client (api/client.js)')
para('All HTTP calls go through one Axios instance:')
code("const API = axios.create({ baseURL: 'https://raees1989.pythonanywhere.com' })")
code('')
code('// Auto-attach token to every request:')
code("API.interceptors.request.use(config => {")
code("  const token = localStorage.getItem('token')")
code("  if (token) config.headers.Authorization = `Bearer ${token}`")
code("  return config")
code("})")
code('')
code('// Auto-refresh on 401 (token expired):')
code('API.interceptors.response.use(res => res, async err => {')
code('  if (err.response?.status === 401) {')
code("    const refresh = localStorage.getItem('refresh_token')")
code("    const res = await axios.post('/api/auth/refresh', { refresh_token: refresh })")
code("    localStorage.setItem('token', res.data.access_token)")
code('    return API(err.config)  // Retry original request')
code('  }')
code('})')

h2('5.4  A Typical Page Component Pattern')
para('Every page follows the same structure:')
code('export default function Inventory() {')
code('  const [products, setProducts] = useState([])')
code('  const [loading, setLoading]   = useState(true)')
code('')
code('  // 1. Load data on mount')
code('  useEffect(() => {')
code('    API.get("/api/inventory/products")')
code('       .then(r => setProducts(r.data))')
code('       .finally(() => setLoading(false))')
code('  }, [])')
code('')
code('  // 2. Create/update via API')
code('  const handleSave = async () => {')
code('    await API.post("/api/inventory/products", form)')
code('    toast.success("Product saved!")')
code('    loadProducts()  // Refresh list')
code('  }')
code('')
code('  // 3. Render')
code('  return (')
code('    <div>')
code('      {loading ? <p>Loading...</p> : products.map(p => <div>{p.name}</div>)}')
code('    </div>')
code('  )')
code('}')

h2('5.5  AuthContext — Login State')
para(
    'AuthContext.js stores the logged-in user globally so every component can access it. '
    'When the user logs in, the user object is saved to context AND to localStorage. '
    'When they refresh the page, context is rebuilt from localStorage.'
)
code('const { user, logout } = useAuth()  // Use anywhere in the app')

h2('5.6  Layout.js — The Shell')
para(
    'Layout.js renders the sidebar + top header around every /app/* page. '
    'It uses <Outlet /> (React Router) to render the active page inside the content area. '
    'SubscriptionGuard is wrapped around <Outlet /> so the trial banner appears on every page.'
)

h2('5.7  Invoice Printing')
para('Invoices are printed using a JavaScript window.open() with inline HTML:')
code("const printInvoice = (sale) => {")
code("  const win = window.open('', '_blank')")
code("  win.document.write(`")
code("    <html><head>")
code("      <style>@media print { body { margin: 10mm; } }</style>")
code("    </head><body>")
code("      <h2>INVOICE ${sale.invoice_number}</h2>")
code("      <table>")
code("        <tr><th>Product</th><th>Batch</th><th>Expiry</th><th>Qty</th><th>Total</th></tr>")
code("        ${sale.items.map(item => `<tr>...")
code("      </table>")
code("      <div>Pharmacist: ________________</div>")
code("    </body></html>`)")
code("  win.print()")
code("}")

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════════
# 6. AUTHENTICATION
# ════════════════════════════════════════════════════════════════════════════
h1('6.  Authentication System (JWT)')

h2('6.1  How JWT Authentication Works')
para('JWT = JSON Web Token. It is the industry standard for API authentication.')
tbl(
    ['Step','What Happens'],
    [
        ['1. User logs in','Frontend sends email + password to POST /api/auth/login'],
        ['2. Backend verifies','Checks password with bcrypt.verify(password, hashed_password)'],
        ['3. Tokens issued','Backend creates access_token (1hr) + refresh_token (7 days)'],
        ['4. Frontend stores','Saves both tokens to localStorage'],
        ['5. Every API call','Axios attaches: Authorization: Bearer <access_token>'],
        ['6. Backend validates','Decodes JWT, extracts email, loads user from DB'],
        ['7. Token expires','Axios interceptor catches 401, calls /api/auth/refresh'],
        ['8. New token','Fresh access_token issued without requiring password again'],
        ['9. Refresh fails','localStorage.clear() + redirect to /login'],
    ]
)

h2('6.2  Backend Auth Code (auth.py)')
code('from jose import jwt, JWTError')
code('from passlib.context import CryptContext')
code('from datetime import datetime, timedelta')
code('')
code('SECRET_KEY = "your-secret-key-here"  # Keep this private!')
code('ALGORITHM  = "HS256"')
code('pwd_context = CryptContext(schemes=["bcrypt"])')
code('')
code('def get_password_hash(password):')
code('    return pwd_context.hash(password)')
code('')
code('def verify_password(plain, hashed):')
code('    return pwd_context.verify(plain, hashed)')
code('')
code('def create_access_token(data: dict):')
code('    expire = datetime.utcnow() + timedelta(hours=1)')
code('    data.update({"exp": expire, "type": "access"})')
code('    return jwt.encode(data, SECRET_KEY, algorithm=ALGORITHM)')
code('')
code('def get_current_user(token: str = Depends(oauth2_scheme), db = Depends(get_db)):')
code('    payload = jwt.decode(token, SECRET_KEY, algorithms=[ALGORITHM])')
code('    email = payload.get("sub")')
code('    user = db.query(models.User).filter(models.User.email == email).first()')
code('    if not user: raise HTTPException(401, "User not found")')
code('    return user')

h2('6.3  Password Security')
bullet('Passwords are NEVER stored in plain text — always hashed with bcrypt.')
bullet('bcrypt automatically adds a salt — same password produces different hash each time.')
bullet('Even if someone steals the database file, passwords cannot be recovered.')

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════════
# 7. SUBSCRIPTION SYSTEM
# ════════════════════════════════════════════════════════════════════════════
h1('7.  SaaS Subscription System')

h2('7.1  How It Works')
para(
    'The subscription system is the business model engine. It controls who has access based on '
    'whether they have paid. There is one row in the subscriptions table per installation.'
)
tbl(
    ['Component','File','Purpose'],
    [
        ['Subscription model','models.py','Stores plan, status, trial_end, paid_until'],
        ['Subscription API','routers/subscription.py','GET status, POST activate, POST extend-trial'],
        ['SubscriptionGuard','components/SubscriptionGuard.js','React component — shows banner or paywall'],
        ['Layout.js','components/Layout.js','Wraps every page with SubscriptionGuard'],
    ]
)

h2('7.2  The Flow')
code('New install')
code('   -> get_or_create_subscription()  # Creates 1 row: plan=trial, trial_end=today+14')
code('')
code('Every page load after login')
code('   -> SubscriptionGuard calls GET /api/subscription/status')
code('   -> Returns: { plan, is_active, days_remaining, trial_end }')
code('')
code('is_active=True + plan=trial  -> Show blue banner "14 days remaining"')
code('days_remaining <= 3          -> Banner turns red, urgent warning')
code('is_active=False (expired)    -> Show full-screen PAYWALL')
code('')
code('Client pays -> Admin goes to Settings -> POST /api/subscription/activate')
code('   -> Updates: plan=basic, status=active, paid_until=today+30days')
code('   -> Paywall disappears immediately')

h2('7.3  subscription.py Key Functions')
code('def get_or_create_subscription(db):')
code('    sub = db.query(models.Subscription).first()')
code('    if not sub:')
code('        sub = models.Subscription(')
code('            plan="trial",')
code('            status="active",')
code('            trial_end=str(date.today() + timedelta(days=14))')
code('        )')
code('        db.add(sub); db.commit()')
code('    return sub')
code('')
code('def check_subscription_active(db):')
code('    sub = get_or_create_subscription(db)')
code('    today = str(date.today())')
code('    if sub.plan == "trial" and today > sub.trial_end:')
code('        sub.status = "expired"; db.commit()')
code('        return sub, False')
code('    return sub, True')

h2('7.4  Adding the Subscription Check to Any Endpoint')
para('To protect any API endpoint based on subscription status:')
code('@router.get("/sales/")')
code('def get_sales(db = Depends(get_db), user = Depends(get_current_user)):')
code('    sub, is_active = check_subscription_active(db)')
code('    if not is_active:')
code('        raise HTTPException(403, "Subscription expired")')
code('    return db.query(models.Sale).all()')

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════════
# 8. DEPLOYMENT
# ════════════════════════════════════════════════════════════════════════════
h1('8.  Deployment — PythonAnywhere + Vercel')

h2('8.1  Backend Deployment (PythonAnywhere)')
tbl(
    ['What','Detail'],
    [
        ['Server URL','https://raees1989.pythonanywhere.com'],
        ['Files location','/home/Raees1989/probiz/backend/'],
        ['Virtualenv','~/.virtualenvs/probizenv/'],
        ['Database file','/home/Raees1989/probiz/backend/probiz.db'],
        ['Web reload URL','https://www.pythonanywhere.com/user/Raees1989/webapps/'],
    ]
)
para('Steps to deploy a backend change:')
numbered('Edit files locally (on your PC).')
numbered('Connect to PythonAnywhere Bash console.')
numbered('Edit the file directly on server (nano or vi), or upload via Files tab.')
numbered('Clear .pyc cache: find ~/probiz/backend -name "*.pyc" -delete')
numbered('Go to Web tab -> click Reload button.')
note('After ANY Python file change you MUST clear .pyc cache AND reload the web app.')

h2('8.2  WSGI File — The Critical Config')
para('The WSGI file at /var/www/raees1989_pythonanywhere_com_wsgi.py tells PythonAnywhere how to run the app:')
code("import sys, os")
code("sys.path.insert(0, '/home/Raees1989/probiz/backend')")
code("os.environ['DATABASE_URL'] = 'sqlite:////home/Raees1989/probiz/backend/probiz.db'")
code("")
code("_asgi_app = None")
code("def application(environ, start_response):")
code("    global _asgi_app")
code("    if _asgi_app is None:")
code("        from main import app")
code("        import uvicorn")
code("        _asgi_app = app")
code("    ...")

h2('8.3  Frontend Deployment (Vercel)')
para('One command deploys the React app:')
code('cd C:/Users/Acer/ProBiz-ERP/frontend')
code('npx vercel --prod --yes --token "YOUR_VERCEL_TOKEN"')
para('Vercel builds the React app and deploys to a global CDN. The live URL is:')
code('https://probiz-erp-poru.vercel.app')
para('Every deployment creates a new URL but also updates the main alias automatically.')

h2('8.4  Environment Variables')
tbl(
    ['Variable','Where Set','Value'],
    [
        ['DATABASE_URL','WSGI file (server)','sqlite:////home/Raees1989/probiz/backend/probiz.db'],
        ['REACT_APP_API_URL','frontend/.env or Vercel','https://raees1989.pythonanywhere.com'],
        ['SECRET_KEY','backend/config.py','Your long random secret string'],
    ]
)

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════════
# 9. BUILD SIMILAR SOFTWARE FROM SCRATCH
# ════════════════════════════════════════════════════════════════════════════
h1('9.  How to Build a Similar ERP from Scratch')
para(
    'This section gives you a step-by-step roadmap to build a similar SaaS ERP for a '
    'different industry (clinic, restaurant, school, retail store, etc.).'
)

h2('9.1  Phase 1 — Setup (Day 1)')
numbered('Install Python 3.10+ and Node.js 18+ on your PC.')
numbered('Create project folder: mkdir MyERP && cd MyERP')
numbered('Create virtualenv: python -m venv venv && venv\\Scripts\\activate')
numbered('Install backend: pip install fastapi uvicorn sqlalchemy pydantic python-jose passlib[bcrypt] python-multipart')
numbered('Create React app: npx create-react-app frontend')
numbered('Install frontend libs: npm install axios react-router-dom react-hot-toast recharts lucide-react')

h2('9.2  Phase 2 — Backend Core (Days 2-4)')
numbered('Create database.py (copy the pattern from section 4.2).')
numbered('Create models.py — define your tables as Python classes.')
numbered('Create auth.py — copy JWT logic from section 6.2.')
numbered('Create main.py — set up FastAPI app with CORS.')
numbered('Create routers/ folder — add one file per module.')
numbered('Test with: uvicorn main:app --reload')
numbered('Open http://localhost:8000/docs to see auto-generated API documentation.')

h2('9.3  Phase 3 — Frontend Core (Days 5-7)')
numbered('Set up React Router in App.js with all your page routes.')
numbered('Create api/client.js with Axios + auth interceptors.')
numbered('Create AuthContext.js for login state management.')
numbered('Build Layout.js with sidebar navigation.')
numbered('Build Login.js page.')
numbered('Start building pages one by one — Dashboard first, then the core module.')

h2('9.4  Phase 4 — Business Logic (Days 8-14)')
numbered('Implement the core transaction (e.g. sales/POS for retail, patient billing for clinic).')
numbered('Add inventory or stock management.')
numbered('Add customer/patient accounts.')
numbered('Add basic reporting.')
numbered('Test end-to-end: create a sale, check stock decreased, check customer balance.')

h2('9.5  Phase 5 — SaaS Features (Days 15-16)')
numbered('Add the Subscription model and router (copy from subscription.py).')
numbered('Add SubscriptionGuard React component.')
numbered('Add demo request form on landing page.')
numbered('Test trial expiry by setting trial_end to yesterday in the DB.')

h2('9.6  Phase 6 — Deployment (Day 17)')
numbered('Create PythonAnywhere free account.')
numbered('Upload backend files to ~/myerp/backend/.')
numbered('Set up virtualenv and install requirements.')
numbered('Configure WSGI file.')
numbered('Create Vercel account and deploy frontend with npx vercel --prod.')
numbered('Test the live URL.')

h2('9.7  Adapting for Different Industries')
tbl(
    ['Industry','Core Module to Change','Extra Fields Needed'],
    [
        ['Pharmacy','Inventory: add batch/expiry. Invoice: add signatures','batch_no, mfg_date, expiry_date, drug_license'],
        ['Restaurant','Replace inventory with menu items. Add table management','table_no, order_type (dine-in/takeaway), kitchen_status'],
        ['Clinic / Hospital','Replace sales with patient billing. Add patient records','patient_id, diagnosis, doctor_id, visit_date'],
        ['School','Replace sales with fee collection. Add student records','student_id, class, fee_type, due_date'],
        ['Garment Store','Add size/color variants to products','size, color, variant_sku, season'],
        ['Spare Parts','Add vehicle compatibility to products','vehicle_make, vehicle_model, part_number'],
        ['Hotel','Replace inventory with rooms. Add booking system','room_no, room_type, check_in, check_out, guest_id'],
    ]
)

h2('9.8  Common Mistakes to Avoid')
bullet('Unique constraint on barcodes — empty string "" violates uniqueness. Convert to null before saving.')
bullet('Two database files on PythonAnywhere — always check WSGI file to find which DB is actually used.')
bullet('.pyc cache — after any model change, always delete .pyc files before reloading.')
bullet('CORS errors — make sure allow_origins=["*"] in main.py during development.')
bullet('Token key names — be consistent. Pick one name ("token" not "access_token") and use it everywhere.')
bullet('Foreign key without data — always add None check before creating relationships.')

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════════
# PART B HEADER
# ════════════════════════════════════════════════════════════════════════════
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('PART B — SELLING THE SOFTWARE')
r.bold = True; r.font.size = Pt(16)
r.font.color.rgb = RGBColor(0xff,0xff,0xff)
pPr = p._p.get_or_add_pPr()
shd = OxmlElement('w:shd')
shd.set(qn('w:val'),'clear'); shd.set(qn('w:color'),'auto'); shd.set(qn('w:fill'),'065F46')
pPr.append(shd)
doc.add_paragraph()

# ════════════════════════════════════════════════════════════════════════════
# 10. SAAS BUSINESS MODEL
# ════════════════════════════════════════════════════════════════════════════
h1('10.  SaaS Business Model Explained')
para(
    'SaaS = Software as a Service. Instead of selling software once (like a CD), '
    'you charge clients a monthly or yearly fee to use the software. '
    'This creates recurring revenue — clients pay every month as long as they use the system.'
)

h2('10.1  Why SaaS is Better Than One-Time Sale')
tbl(
    ['One-Time Sale','SaaS Monthly Subscription'],
    [
        ['Sell for Rs. 50,000 once','Sell for Rs. 2,500/month = Rs. 30,000/year'],
        ['No income after sale','Income every month without doing new work'],
        ['Client uses outdated version','Client always gets latest version automatically'],
        ['No relationship after sale','Ongoing relationship — you can upsell upgrades'],
        ['Need many new sales every month','10 clients = Rs. 25,000/month stable income'],
    ]
)

h2('10.2  Revenue Example')
tbl(
    ['Clients','Plan','Monthly Revenue','Annual Revenue'],
    [
        ['5 clients','Basic Rs. 2,500','Rs. 12,500/month','Rs. 150,000/year'],
        ['10 clients','Basic Rs. 2,500','Rs. 25,000/month','Rs. 300,000/year'],
        ['5 clients','Professional Rs. 5,000','Rs. 25,000/month','Rs. 300,000/year'],
        ['20 clients','Mix of Basic + Pro','Rs. 60,000-80,000/month','Rs. 720,000-960,000/year'],
        ['50 clients','Mix of all plans','Rs. 150,000+/month','Rs. 1.8M+/year'],
    ]
)
note('With PythonAnywhere free tier, hosting cost is near zero for the first 5-10 clients.')

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════════
# 11. PRICING STRATEGY
# ════════════════════════════════════════════════════════════════════════════
h1('11.  Pricing Strategy')

h2('11.1  Current ProBiz Plans')
tbl(
    ['Plan','Monthly Price','Annual Price','Target Client','Key Limits'],
    [
        ['Free Trial','Free','Free','Anyone evaluating','14 days, 1 branch, 3 users'],
        ['Basic','Rs. 2,500','Rs. 25,000 (save Rs. 5,000)','Single pharmacy/shop','1 branch, 3 users'],
        ['Professional','Rs. 5,000','Rs. 50,000','Clinic + pharmacy, small chain','3 branches, 10 users'],
        ['Enterprise','Rs. 15,000','Rs. 150,000','Hospital, wholesale, chain','Unlimited everything'],
    ]
)

h2('11.2  Pricing Psychology Tips')
bullet('Always show 3 plans — clients almost never choose the cheapest OR most expensive.')
bullet('Mark the middle plan as "Popular" — it anchors the comparison.')
bullet('Offer annual discount (2 months free) — improves cash flow and reduces churn.')
bullet('Never compete on price with large ERP companies — compete on local support and simplicity.')
bullet('Rs. 2,500/month = Rs. 83/day — frame it as less than a cup of chai per day.')

h2('11.3  Setup / Installation Fee')
para('Consider charging a one-time setup fee in addition to monthly subscription:')
tbl(
    ['Service','Suggested Price'],
    [
        ['Basic installation + training (2 hrs)','Rs. 5,000 one-time'],
        ['Data migration from Excel / old software','Rs. 3,000 - 10,000 depending on volume'],
        ['Custom invoice/report design','Rs. 2,000 - 5,000'],
        ['On-site training (per day)','Rs. 5,000'],
        ['Remote training (Zoom/WhatsApp)','Rs. 2,000'],
    ]
)

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════════
# 12. TARGET CUSTOMERS
# ════════════════════════════════════════════════════════════════════════════
h1('12.  Target Customers & Industries')

h2('12.1  Primary Targets for ProBiz ERP (Pharmacy Focus)')
tbl(
    ['Customer Type','Why They Need This','Pain Points You Solve'],
    [
        ['Independent pharmacies','Track medicines, batch numbers, expiry dates','Manual stock register, expiry wastage, no profit tracking'],
        ['Pharmacy chains (2-5 branches)','Centralized control of all branches','Cannot see which branch is profitable'],
        ['Hospital pharmacies','Patient billing + medicine dispensing','Manual billing errors, no stock control'],
        ['Medical stores + distributors','Purchase and sales tracking with suppliers','Outstanding balance confusion, no payables tracking'],
        ['Cash & carry with medicine section','Mixed inventory (medicine + FMCG)','Cannot separate medicine sales from general'],
    ]
)

h2('12.2  Expanding to Other Industries')
para('The same codebase can be adapted with minor changes:')
tbl(
    ['Industry','Modifications Needed','Market Size (Pakistan)'],
    [
        ['General retail stores','Remove batch/expiry, add variants (size/color)','Massive — every kiryana shop'],
        ['Restaurants / cafes','Add menu, table management, kitchen display','Growing fast in cities'],
        ['Clinics / specialist doctors','Add patient records, appointment booking','Every clinic needs this'],
        ['Schools / academies','Add student accounts, fee collection','Thousands of private schools'],
        ['Garment / fashion','Add size, color, season to products','Lahore and Faisalabad wholesale'],
        ['Auto spare parts','Add vehicle compatibility fields','Every spare parts market'],
    ]
)

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════════
# 13. SALES & MARKETING
# ════════════════════════════════════════════════════════════════════════════
h1('13.  Sales & Marketing Strategy')

h2('13.1  Your Biggest Advantage — Trust + Local Knowledge')
para(
    'You are Dr Muhammad Raees RPh — a licensed pharmacist. Your target customers are pharmacies. '
    'You understand their problems better than any tech company. '
    'This is your biggest competitive advantage. Use it in every sales conversation.'
)
bullet('"I built this system myself because I was using the same manual registers you are using."')
bullet('"I understand DRAP requirements — that is why the system has batch numbers and expiry dates on every invoice."')
bullet('"I am available on WhatsApp 24/7 — no call center, no ticket system."')

h2('13.2  Free Demo Strategy')
para('The free trial is your most powerful sales tool. The workflow:')
numbered('Prospect hears about your ERP (WhatsApp, referral, LinkedIn, walk-in).')
numbered('Give them the live demo URL: https://probiz-erp-poru.vercel.app')
numbered('They explore it for free for 14 days — no pressure.')
numbered('After 3-4 days, follow up on WhatsApp: "How is the system? Any questions?"')
numbered('Before trial ends, call or WhatsApp: "Your trial ends in 3 days — shall I activate your plan?"')
numbered('If they say yes — activate from Settings and send payment details.')
numbered('If they say not yet — extend trial by 7-14 days and follow up again.')

h2('13.3  Where to Find Clients')
tbl(
    ['Channel','Approach','Cost'],
    [
        ['Your own pharmacy network','You know pharmacists personally — start there','Free'],
        ['WhatsApp groups (pharmacists)','Share demo link in pharmacy WhatsApp groups','Free'],
        ['LinkedIn','Post about ProBiz ERP — pharmacist + tech audience','Free'],
        ['PIMA (Pakistan Islami Medical Association)','Present at local pharmacy events','Low'],
        ['Referrals from happy clients','Offer 1 month free for every referral that converts','Free'],
        ['Walk-ins to pharmacies','Take laptop, demo on the spot, leave your card','Time only'],
        ['Facebook pharmacy groups','Post in local business and pharmacy groups','Free'],
        ['Google Ads (later)','Target "pharmacy software Pakistan"','Paid - invest when revenue allows'],
    ]
)

h2('13.4  Demo Script (What to Say in First Meeting)')
numbered('Start with a question: "How do you currently track your stock and sales?"')
numbered('Listen to their pain: manual registers, expiry wastage, outstanding balances.')
numbered('Show the relevant module first: if they complain about stock — show Inventory.')
numbered('Run a live demo: add a product, make a sale, show the invoice with batch number.')
numbered('Show the trial banner: "It is free for 14 days — nothing to install, works on your phone."')
numbered('Close: "Shall I set up your pharmacy name and products today? Takes 30 minutes."')

h2('13.5  WhatsApp Marketing Message Template')
para('Send this to pharmacy contacts:')
code('Assalam o Alaikum,')
code('')
code('I am Dr Muhammad Raees RPh. I have built a complete')
code('Pharmacy Management System specifically for Pakistani')
code('pharmacies — POS, inventory with batch/expiry tracking,')
code('invoicing with pharmacist signature, accounts, and payroll.')
code('')
code('FREE 14-day trial — no installation, works on any browser.')
code('')
code('Live demo: https://probiz-erp-poru.vercel.app')
code('Login: raees.malik89@gmail.com / admin123')
code('')
code('Plans start from Rs. 2,500/month.')
code('WhatsApp me for a free live walkthrough: 0316-8818693')

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════════
# 14. ONBOARDING A NEW CLIENT
# ════════════════════════════════════════════════════════════════════════════
h1('14.  How to Onboard a New Client')
para('When a client agrees to subscribe, follow this process:')

h2('14.1  Option A — Shared Instance (Same Server)')
para('The client uses the same PythonAnywhere server as your demo. Easiest to start with.')
numbered('Take payment.')
numbered('Go to Settings on live app -> Subscription Management -> Activate their plan.')
numbered('Change the company name in Settings to their pharmacy name.')
numbered('Help them add their products (via Inventory or bulk import).')
numbered('Create user accounts for their staff.')
numbered('Train them for 1-2 hours (in person or WhatsApp video).')
note('Limitation: all clients share the same database on the same server. Works for 5-10 clients.')

h2('14.2  Option B — Dedicated Instance (Recommended for Paying Clients)')
para('Each client gets their own PythonAnywhere account and database. More professional.')
numbered('Create a new PythonAnywhere account for the client (use their business email).')
numbered('Upload backend files to the new account.')
numbered('Set up a new probiz.db with clean data.')
numbered('Deploy frontend with the new backend URL as REACT_APP_API_URL.')
numbered('Create the superadmin login with the client\'s credentials.')
numbered('Activate their subscription plan.')
numbered('Train them and hand over.')
para('Cost: PythonAnywhere free tier handles 1 web app per account — no extra hosting cost for basic plan clients.')

h2('14.3  Data Migration')
para('Most clients have data in Excel or handwritten registers:')
bullet('For products: ask client for Excel file with columns: Name, Unit, Cost Price, Sale Price, Stock')
bullet('Write a simple Python script to bulk import products via the API')
bullet('For customers: import customer list with name, phone, outstanding balance')
bullet('Charge Rs. 3,000-10,000 for data migration service')

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════════
# 15. OBJECTIONS
# ════════════════════════════════════════════════════════════════════════════
h1('15.  Common Objections & How to Handle Them')

objections = [
    (
        '"It is too expensive — Rs. 2,500 per month is a lot."',
        'Rs. 2,500/month = Rs. 83/day. Less than the cost of one medicine that expires because '
        'you had no expiry tracking. How much do you lose per month from expired medicines? '
        'This system pays for itself in the first week.'
    ),
    (
        '"We already use a software / Excel."',
        '"Can your current software show batch numbers and expiry dates on invoices? '
        'Can it track which customer owes you how much? Can you access it from your mobile '
        'when you are away from the shop?" Show them something their current system cannot do.'
    ),
    (
        '"What if the internet is down?"',
        '"The system needs internet to work, but internet outages in Pakistan are short and rare. '
        'For critical moments, your phone\'s mobile data works. Most pharmacies have WiFi anyway. '
        'All data is safely backed up in the cloud — nothing is lost."'
    ),
    (
        '"What if you stop supporting it?"',
        '"I am a pharmacist — this is my own business tool too, not just a product I sell. '
        'I have been running it for my own use. And all your data is stored in your own '
        'database — I can give you a full backup of your data at any time."'
    ),
    (
        '"My staff is not tech-savvy."',
        '"The POS screen is as simple as a calculator. Most cashiers learn it in 30 minutes. '
        'I provide free training for your staff and I am on WhatsApp if anyone has questions."'
    ),
    (
        '"Let me think about it / I will call you later."',
        '"Of course. Let me set up your 14-day free trial right now — takes 5 minutes. '
        'You can explore it at your own pace and decide after seeing it with your own products."'
        ' Always leave with them using the trial, not just thinking about it.'
    ),
]
for objection, response in objections:
    h3(objection)
    para(response)
    doc.add_paragraph()

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════════
# 16. SCALING
# ════════════════════════════════════════════════════════════════════════════
h1('16.  Scaling the Business')

h2('16.1  Growth Roadmap')
tbl(
    ['Stage','Clients','Monthly Revenue','What to Do'],
    [
        ['Start','1-5 clients','Rs. 5,000-15,000','Focus on pharmacies in Islamabad. Deliver excellent support.'],
        ['Early Growth','5-20 clients','Rs. 15,000-60,000','Add referral program. Hire 1 part-time support person.'],
        ['Scale','20-50 clients','Rs. 60,000-150,000','Add more industry verticals. Upgrade to paid hosting per client.'],
        ['Established','50-100+ clients','Rs. 150,000-400,000','Hire developer to add features. Expand to Lahore/Karachi.'],
    ]
)

h2('16.2  Features to Add for Higher Pricing')
bullet('Urdu language interface — huge demand, very few ERP systems have this.')
bullet('SMS notifications — auto SMS on sale, low stock, payment reminder.')
bullet('WhatsApp integration — send invoice directly to customer on WhatsApp.')
bullet('Barcode scanner support — scan product barcodes to add to sale.')
bullet('Custom logo and letterhead on invoices.')
bullet('Mobile app (React Native) — same codebase, runs as Android app.')
bullet('Multi-tenant SaaS — one codebase, many clients on same server with data isolation.')
bullet('Analytics dashboard with AI insights — which products to reorder, peak sales times.')

h2('16.3  Technical Scaling')
tbl(
    ['When','Problem','Solution'],
    [
        ['5+ clients on one server','Server gets slow','Upgrade PythonAnywhere to paid ($5/mo)'],
        ['10+ clients','Data isolation concerns','Switch to separate DB per client (already designed for this)'],
        ['50+ clients','Manual client setup is slow','Build admin panel to auto-provision new clients'],
        ['100+ clients','SQLite limits','Migrate to PostgreSQL (change DATABASE_URL only)'],
        ['500+ clients','PythonAnywhere not enough','Move to AWS/DigitalOcean with Docker containers'],
    ]
)

h2('16.4  Protecting Your Product')
bullet('Clients do not see the source code — they only use the deployed web app.')
bullet('Your competitive advantage is not the code — it is your support, local knowledge, and relationships.')
bullet('Add a watermark or your branding to all printouts.')
bullet('Backend is hosted on your PythonAnywhere account — client cannot copy the database logic.')
bullet('Consider trademarking "ProBiz ERP" once revenue justifies it.')

h2('16.5  Building a Team')
tbl(
    ['Role','When to Hire','What They Do'],
    [
        ['Support staff (part-time)','When you have 15+ clients','Answer WhatsApp, do basic training'],
        ['Frontend developer','When you have 25+ clients','Add features, fix UI issues'],
        ['Sales person','When you have steady revenue','Cold outreach to new pharmacies/businesses'],
        ['Full-stack developer','When you have 50+ clients','Build new industry verticals'],
    ]
)

doc.add_paragraph()
para('Final Note:', bold=True, size=13)
para(
    'You have already built a production-ready SaaS ERP system. The hardest part is done. '
    'The next step is simply getting it in front of the right people. '
    'Every pharmacy owner in Pakistan who sees this system working will want it. '
    'Start with your own network, deliver excellent support to the first 5 clients, '
    'and let word of mouth do the rest.',
    size=12
)

doc.add_paragraph()
c = doc.add_paragraph()
c.alignment = WD_ALIGN_PARAGRAPH.CENTER
r5 = c.add_run('Dr Muhammad Raees RPh  |  0316-8818693  |  raees.malik89@gmail.com')
r5.bold = True; r5.font.color.rgb = RGBColor(0x1e,0x40,0xaf); r5.font.size = Pt(12)

# ── SAVE ──────────────────────────────────────────────────────────────────
out = r'C:\Users\Acer\ProBiz-ERP\ProBiz_ERP_Technical_and_Sales_Guide.docx'
doc.save(out)
print('Saved: ' + out)
