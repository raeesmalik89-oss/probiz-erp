from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime

doc = Document()

# ── Page setup ────────────────────────────────────────────────────────────────
for sec in doc.sections:
    sec.top_margin    = Inches(1)
    sec.bottom_margin = Inches(1)
    sec.left_margin   = Inches(1.1)
    sec.right_margin  = Inches(1.1)

# ── Helpers ───────────────────────────────────────────────────────────────────
def shade(para, fill):
    pPr = para._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill)
    pPr.append(shd)

def set_cell_color(cell, fill):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill)
    tcPr.append(shd)

def file_header(title, subtitle, color_fill='1F497D'):
    """Big coloured banner for each file section."""
    p = doc.add_paragraph()
    shade(p, color_fill)
    r = p.add_run(f'  {title}')
    r.font.name = 'Calibri'; r.font.size = Pt(14); r.font.bold = True
    r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(0)

    p2 = doc.add_paragraph()
    shade(p2, '2E75B6')
    r2 = p2.add_run(f'  {subtitle}')
    r2.font.name = 'Calibri'; r2.font.size = Pt(10)
    r2.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    p2.paragraph_format.space_before = Pt(0)
    p2.paragraph_format.space_after  = Pt(6)

def code_block(lines):
    """Render a list of strings as a monospace code block."""
    for line in lines:
        p = doc.add_paragraph()
        shade(p, 'F5F5F5')
        display = line if line.strip() != '' else ' '
        r = p.add_run(display)
        r.font.name = 'Courier New'
        r.font.size = Pt(9)
        r.font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after  = Pt(0)
        p.paragraph_format.left_indent  = Inches(0.15)
    # small gap after block
    gap = doc.add_paragraph()
    gap.paragraph_format.space_before = Pt(0)
    gap.paragraph_format.space_after  = Pt(6)

def label(text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.name = 'Calibri'; r.font.size = Pt(10); r.font.bold = True
    r.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after  = Pt(2)

# ══════════════════════════════════════════════════════════════════════════════
# TITLE PAGE
# ══════════════════════════════════════════════════════════════════════════════
doc.add_paragraph()
doc.add_paragraph()

t = doc.add_paragraph(); t.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = t.add_run('AIOpsCare')
r.font.name='Calibri'; r.font.size=Pt(40); r.font.bold=True
r.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)

s = doc.add_paragraph(); s.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = s.add_run('Complete Source Code — All Files')
r.font.name='Calibri'; r.font.size=Pt(18)
r.font.color.rgb = RGBColor(0x2E, 0x75, 0xB6)

d = doc.add_paragraph(); d.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = d.add_run('Serverless AIOps Platform  |  Final Year Project — Topic 7')
r.font.name='Calibri'; r.font.size=Pt(12)
r.font.color.rgb = RGBColor(0x60, 0x60, 0x60)

doc.add_paragraph()
dt = doc.add_paragraph(); dt.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = dt.add_run(datetime.datetime.now().strftime('%B %d, %Y'))
r.font.name='Calibri'; r.font.size=Pt(11)
r.font.color.rgb = RGBColor(0x80, 0x80, 0x80)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# FILE INDEX TABLE
# ══════════════════════════════════════════════════════════════════════════════
idx_title = doc.add_paragraph()
r = idx_title.add_run('File Index')
r.font.name='Calibri'; r.font.size=Pt(16); r.font.bold=True
r.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)
idx_title.paragraph_format.space_after = Pt(8)

files = [
    ('01', 'app/main.py',                         'FastAPI app — routes, tracing, Prometheus',     'UPDATED'),
    ('02', 'app/schemas.py',                       'Pydantic request model',                        'ORIGINAL'),
    ('03', 'app/model_loader.py',                  'ML model loader with path fix',                 'UPDATED'),
    ('04', 'app/alerts.py',                        'ICU threshold alert engine',                    'ORIGINAL'),
    ('05', 'app/auth.py',                          'JWT + OPA authorization middleware',            'NEW'),
    ('06', 'ml/train.py',                          'Model training pipeline with scaler',           'UPDATED'),
    ('07', 'events/producer.py',                   'Kafka patient event producer',                  'ORIGINAL'),
    ('08', 'events/consumer.py',                   'Kafka consumer → OpenFaaS trigger',            'UPDATED'),
    ('09', 'functions/predict/handler.py',         'OpenFaaS serverless function logic',            'NEW'),
    ('10', 'functions/predict/entrypoint.py',      'Flask HTTP wrapper for OpenFaaS',              'NEW'),
    ('11', 'functions/predict/Dockerfile',         'Function container image definition',           'NEW'),
    ('12', 'functions/predict/requirements.txt',   'Function dependencies',                         'NEW'),
    ('13', 'functions/stack.yml',                  'OpenFaaS deployment manifest',                  'NEW'),
    ('14', 'security/opa/policy.rego',             'Open Policy Agent authorization rules',         'NEW'),
    ('15', 'security/keycloak/realm-export.json',  'Keycloak realm, roles and users',              'NEW'),
    ('16', 'infrastructure/docker-compose.yml',    'All 11 services orchestration',                'UPDATED'),
    ('17', 'monitoring/prometheus/prometheus.yml',  'Prometheus scrape config',                     'UPDATED'),
    ('18', 'scripts/zap_scan.sh',                  'OWASP ZAP security scan script',               'NEW'),
    ('19', 'Dockerfile',                           'FastAPI container image (root)',                'FILLED'),
    ('20', 'requirements.txt',                     'All Python dependencies',                       'UPDATED'),
]

table = doc.add_table(rows=1 + len(files), cols=4)
table.style = 'Table Grid'
col_w = [Inches(0.4), Inches(2.6), Inches(3.2), Inches(0.9)]

# Header
hdr = table.rows[0]
for i, (cell, w, txt) in enumerate(zip(hdr.cells, col_w,
        ['#', 'File Path', 'Description', 'Status'])):
    cell.width = w
    cell.paragraphs[0].clear()
    r = cell.paragraphs[0].add_run(txt)
    r.bold=True; r.font.name='Calibri'; r.font.size=Pt(10)
    r.font.color.rgb = RGBColor(0xFF,0xFF,0xFF)
    set_cell_color(cell, '1F497D')

status_colors = {'NEW':'1E7E34', 'UPDATED':'0056B3', 'ORIGINAL':'555555', 'FILLED':'7B3F00'}
status_bg     = {'NEW':'D4EDDA', 'UPDATED':'CCE5FF', 'ORIGINAL':'EEEEEE', 'FILLED':'FFF3CD'}

for ri, (num, path, desc, status) in enumerate(files):
    row = table.rows[ri+1]
    bg = 'F7FBFF' if ri%2==0 else 'FFFFFF'
    data = [num, path, desc, status]
    for ci, (cell, w, val) in enumerate(zip(row.cells, col_w, data)):
        cell.width = w
        cell.paragraphs[0].clear()
        r = cell.paragraphs[0].add_run(val)
        r.font.name = 'Calibri'; r.font.size = Pt(9)
        if ci == 3:
            r.bold = True
            r.font.color.rgb = RGBColor(
                int(status_colors[status][0:2],16),
                int(status_colors[status][2:4],16),
                int(status_colors[status][4:6],16))
            set_cell_color(cell, status_bg[status])
        elif ci == 1:
            r.font.name = 'Courier New'
            set_cell_color(cell, bg)
        else:
            set_cell_color(cell, bg)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# HELPER: emit a complete file
# ══════════════════════════════════════════════════════════════════════════════
def emit_file(number, path, description, status, lines, lang='python'):
    color_map = {'NEW':'1E7E34','UPDATED':'0056B3','ORIGINAL':'555555','FILLED':'7B3F00'}
    hdr_color = color_map.get(status, '1F497D')
    file_header(f'{number}.  {path}', f'{description}  |  Status: {status}', hdr_color)
    code_block(lines)

# ══════════════════════════════════════════════════════════════════════════════
# 01  app/main.py
# ══════════════════════════════════════════════════════════════════════════════
emit_file('01', 'app/main.py', 'FastAPI Application Entry Point', 'UPDATED', [
    'import os',
    '',
    'from fastapi import FastAPI, Depends',
    '',
    'from app.schemas import PatientData',
    'from app.model_loader import model',
    'from app.alerts import generate_alerts',
    'from app.auth import authorize',
    '',
    'from prometheus_fastapi_instrumentator import Instrumentator',
    '',
    '# -----------------------------------',
    '# OpenTelemetry Imports',
    '# -----------------------------------',
    '',
    'from opentelemetry import trace',
    '',
    'from opentelemetry.sdk.resources import Resource',
    'from opentelemetry.sdk.trace import TracerProvider',
    'from opentelemetry.sdk.trace.export import BatchSpanProcessor',
    '',
    'from opentelemetry.exporter.otlp.proto.grpc.trace_exporter import (',
    '    OTLPSpanExporter',
    ')',
    '',
    'from opentelemetry.instrumentation.fastapi import (',
    '    FastAPIInstrumentor',
    ')',
    '',
    '',
    '# -----------------------------------',
    '# FastAPI App',
    '# -----------------------------------',
    '',
    'app = FastAPI(',
    '    title="AIOpsCare",',
    '    description="Real-Time ICU Monitoring & Sepsis Prediction Platform",',
    '    version="1.0.0",',
    ')',
    '',
    '',
    '# -----------------------------------',
    '# Prometheus Metrics',
    '# -----------------------------------',
    '',
    'Instrumentator().instrument(app).expose(app)',
    '',
    '',
    '# -----------------------------------',
    '# OpenTelemetry Configuration',
    '# -----------------------------------',
    '',
    'resource = Resource.create({',
    '    "service.name": "aiopscare-fastapi"',
    '})',
    '',
    'trace.set_tracer_provider(',
    '    TracerProvider(resource=resource)',
    ')',
    '',
    'tracer = trace.get_tracer(__name__)',
    '',
    '',
    '# -----------------------------------',
    '# OTLP Exporter',
    '# -----------------------------------',
    '',
    'otlp_exporter = OTLPSpanExporter(',
    '    endpoint="http://jaeger:4317",',
    '    insecure=True',
    ')',
    '',
    'span_processor = BatchSpanProcessor(',
    '    otlp_exporter',
    ')',
    '',
    'trace.get_tracer_provider().add_span_processor(',
    '    span_processor',
    ')',
    '',
    '',
    '# -----------------------------------',
    '# FastAPI Instrumentation',
    '# -----------------------------------',
    '',
    'FastAPIInstrumentor.instrument_app(app)',
    '',
    '',
    '# -----------------------------------',
    '# API Routes',
    '# -----------------------------------',
    '',
    '@app.get("/")',
    'async def root():',
    '    return {',
    '        "message": "AIOpsCare API running successfully"',
    '    }',
    '',
    '',
    '@app.post("/predict")',
    'async def predict(',
    '    data: PatientData,',
    '    token: str = Depends(authorize)',
    '):',
    '',
    '    with tracer.start_as_current_span(',
    '        "predict-sepsis"',
    '    ):',
    '',
    '        prediction = model.predict([[',
    '            data.heart_rate,',
    '            data.temperature,',
    '            data.respiratory_rate,',
    '        ]])',
    '',
    '        result = int(prediction[0])',
    '',
    '        alerts = generate_alerts(data)',
    '',
    '        return {',
    '            "sepsis_prediction": result,',
    '            "alerts": alerts',
    '        }',
])
doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 02  app/schemas.py
# ══════════════════════════════════════════════════════════════════════════════
emit_file('02', 'app/schemas.py', 'Pydantic Request / Response Models', 'ORIGINAL', [
    'from pydantic import BaseModel',
    '',
    '',
    'class PatientData(BaseModel):',
    '    heart_rate: float',
    '    temperature: float',
    '    respiratory_rate: float',
])
doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 03  app/model_loader.py
# ══════════════════════════════════════════════════════════════════════════════
emit_file('03', 'app/model_loader.py', 'ML Model Loader — Path-Safe with Error Handling', 'UPDATED', [
    'import os',
    'import joblib',
    '',
    '_model_path = os.path.abspath(',
    '    os.path.join(os.path.dirname(__file__), "..", "ml", "model.joblib")',
    ')',
    '',
    'try:',
    '    model = joblib.load(_model_path)',
    'except FileNotFoundError:',
    "    raise RuntimeError(",
    "        f\"Model not found at {_model_path}. Run 'python ml/train.py' first.\"",
    '    )',
])
doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 04  app/alerts.py
# ══════════════════════════════════════════════════════════════════════════════
emit_file('04', 'app/alerts.py', 'ICU Clinical Alert Engine', 'ORIGINAL', [
    '# -----------------------------------',
    '# AIOpsCare ICU Alert Engine',
    '# -----------------------------------',
    '# Detects abnormal patient conditions',
    '# and generates ICU monitoring alerts',
    '',
    '',
    'def generate_alerts(data):',
    '',
    '    alerts = []',
    '',
    '    # -----------------------------------',
    '    # High Heart Rate Alert',
    '    # -----------------------------------',
    '    if data.heart_rate > 100:',
    '',
    '        alerts.append(',
    '            "High Heart Rate Detected"',
    '        )',
    '',
    '    # -----------------------------------',
    '    # High Fever Alert',
    '    # -----------------------------------',
    '    if data.temperature > 38:',
    '',
    '        alerts.append(',
    '            "High Fever Detected"',
    '        )',
    '',
    '    # -----------------------------------',
    '    # Respiratory Risk Alert',
    '    # -----------------------------------',
    '    if data.respiratory_rate > 24:',
    '',
    '        alerts.append(',
    '            "Abnormal Respiratory Rate"',
    '        )',
    '',
    '    # -----------------------------------',
    '    # Return ICU Alerts',
    '    # -----------------------------------',
    '    return alerts',
])
doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 05  app/auth.py
# ══════════════════════════════════════════════════════════════════════════════
emit_file('05', 'app/auth.py', 'JWT + OPA Authorization Middleware  [NEW]', 'NEW', [
    'import os',
    'import httpx',
    'from fastapi import HTTPException, Security, Request',
    'from fastapi.security import HTTPBearer, HTTPAuthorizationCredentials',
    '',
    'OPA_URL = os.getenv("OPA_URL", "http://opa:8181")',
    '',
    'security = HTTPBearer()',
    '',
    '',
    'async def authorize(',
    '    request: Request,',
    '    credentials: HTTPAuthorizationCredentials = Security(security),',
    '):',
    '    token = credentials.credentials',
    '    payload = {',
    '        "input": {',
    '            "token": f"Bearer {token}",',
    '            "method": request.method,',
    '            "path": request.url.path,',
    '        }',
    '    }',
    '    try:',
    '        async with httpx.AsyncClient(timeout=5.0) as client:',
    '            resp = await client.post(',
    '                f"{OPA_URL}/v1/data/aiopscare/authz/allow",',
    '                json=payload,',
    '            )',
    '        if not resp.json().get("result", False):',
    '            raise HTTPException(status_code=403, detail="Forbidden: insufficient role")',
    '    except httpx.RequestError:',
    '        raise HTTPException(status_code=503, detail="Authorization service unavailable")',
    '    return token',
])
doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 06  ml/train.py
# ══════════════════════════════════════════════════════════════════════════════
emit_file('06', 'ml/train.py', 'ML Model Training Pipeline with StandardScaler + Cross-Validation', 'UPDATED', [
    'import os',
    'from sklearn.linear_model import LogisticRegression',
    'from sklearn.preprocessing import StandardScaler',
    'from sklearn.pipeline import Pipeline',
    'from sklearn.model_selection import cross_val_score',
    'import pandas as pd',
    'import joblib',
    '',
    '# Synthetic training data',
    '# Replace with real ICU dataset (e.g. MIMIC-III) before production',
    'data = {',
    '    "heart_rate":       [80, 95, 110, 70, 120, 85, 105, 72, 115, 90],',
    '    "temperature":      [36.5, 38.2, 39.1, 36.8, 40.0, 37.0, 38.8, 36.6, 39.5, 37.5],',
    '    "respiratory_rate": [18, 22, 30, 16, 35, 19, 28, 15, 32, 21],',
    '    "sepsis":           [0, 1, 1, 0, 1, 0, 1, 0, 1, 0],',
    '}',
    '',
    'df = pd.DataFrame(data)',
    'X = df[["heart_rate", "temperature", "respiratory_rate"]]',
    'y = df["sepsis"]',
    '',
    '# Pipeline: scale features then classify',
    '# (StandardScaler is important for logistic regression with different-range features)',
    'pipeline = Pipeline([',
    '    ("scaler", StandardScaler()),',
    '    ("classifier", LogisticRegression(random_state=42)),',
    '])',
    '',
    'scores = cross_val_score(pipeline, X, y, cv=5, scoring="accuracy")',
    'print(f"Cross-val accuracy: {scores.mean():.2f} (+/- {scores.std():.2f})")',
    '',
    'pipeline.fit(X, y)',
    '',
    'output_path = os.path.join(os.path.dirname(__file__), "model.joblib")',
    'joblib.dump(pipeline, output_path)',
    'print(f"Model saved to {output_path}")',
])
doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 07  events/producer.py
# ══════════════════════════════════════════════════════════════════════════════
emit_file('07', 'events/producer.py', 'Kafka Patient Event Producer', 'ORIGINAL', [
    '# -----------------------------------',
    '# Kafka ICU Event Producer',
    '# -----------------------------------',
    '',
    'from kafka import KafkaProducer',
    'import json',
    '',
    '# Kafka Producer Configuration',
    'producer = KafkaProducer(',
    "    bootstrap_servers='localhost:9092',",
    '    value_serializer=lambda v: json.dumps(v).encode("utf-8")',
    ')',
    '',
    '# Send ICU Patient Event',
    'patient_event = {',
    '    "heart_rate": 120,',
    '    "temperature": 39.2,',
    '    "respiratory_rate": 30',
    '}',
    '',
    'producer.send("patient-events", patient_event)',
    '',
    'producer.flush()',
    '',
    'print("Patient event sent successfully.")',
])
doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 08  events/consumer.py
# ══════════════════════════════════════════════════════════════════════════════
emit_file('08', 'events/consumer.py', 'Kafka Consumer — Triggers OpenFaaS Serverless Function', 'UPDATED', [
    'import os',
    'import json',
    'import httpx',
    'from kafka import KafkaConsumer',
    '',
    'KAFKA_BROKER      = os.getenv("KAFKA_BROKER", "kafka:29092")',
    'OPENFAAS_URL      = os.getenv("OPENFAAS_URL", "http://openfaas-gateway:8080")',
    'FUNCTION_ENDPOINT = f"{OPENFAAS_URL}/function/sepsis-predict"',
    '',
    'consumer = KafkaConsumer(',
    '    "patient-events",',
    '    bootstrap_servers=KAFKA_BROKER,',
    '    auto_offset_reset="earliest",',
    '    value_deserializer=lambda x: json.loads(x.decode("utf-8")),',
    ')',
    '',
    'print("Consumer listening — forwarding events to OpenFaaS predict function...")',
    '',
    'for message in consumer:',
    '    patient_data = message.value',
    '    print(f"\\nReceived event: {patient_data}")',
    '',
    '    try:',
    '        response = httpx.post(FUNCTION_ENDPOINT, json=patient_data, timeout=10.0)',
    '        result = response.json()',
    "        print(f\"Prediction: {result['sepsis_prediction']}\")",
    '        if result.get("alerts"):',
    '            print(f"Alerts: {result[\'alerts\']}")',
    '        if result["sepsis_prediction"] == 1:',
    '            print("ACTION: High sepsis risk — escalating to ICU team.")',
    '        else:',
    '            print("STATUS: Patient stable.")',
    '    except httpx.RequestError as e:',
    '        print(f"ERROR: Could not reach predict function — {e}")',
])
doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 09  functions/predict/handler.py
# ══════════════════════════════════════════════════════════════════════════════
emit_file('09', 'functions/predict/handler.py', 'OpenFaaS Serverless Function — Prediction Logic  [NEW]', 'NEW', [
    'import json',
    'import joblib',
    'import os',
    '',
    'MODEL_PATH           = os.getenv("MODEL_PATH", "/app/ml/model.joblib")',
    'HEART_RATE_THRESHOLD = 100',
    'TEMP_THRESHOLD       = 38.0',
    'RR_THRESHOLD         = 24',
    '',
    'model = joblib.load(MODEL_PATH)',
    '',
    '',
    'def handle(event, context):',
    '    try:',
    '        body             = json.loads(event.body)',
    '        heart_rate       = float(body["heart_rate"])',
    '        temperature      = float(body["temperature"])',
    '        respiratory_rate = float(body["respiratory_rate"])',
    '',
    '        prediction = int(model.predict(',
    '            [[heart_rate, temperature, respiratory_rate]]',
    '        )[0])',
    '',
    '        alerts = []',
    '        if heart_rate > HEART_RATE_THRESHOLD:',
    '            alerts.append("High Heart Rate Detected")',
    '        if temperature > TEMP_THRESHOLD:',
    '            alerts.append("High Fever Detected")',
    '        if respiratory_rate > RR_THRESHOLD:',
    '            alerts.append("Abnormal Respiratory Rate")',
    '',
    '        return {',
    '            "statusCode": 200,',
    '            "body": json.dumps({',
    '                "sepsis_prediction": prediction,',
    '                "alerts": alerts',
    '            }),',
    '        }',
    '',
    '    except (KeyError, ValueError) as e:',
    '        return {',
    '            "statusCode": 400,',
    '            "body": json.dumps({"error": f"Invalid input: {e}"})',
    '        }',
    '    except Exception as e:',
    '        return {',
    '            "statusCode": 500,',
    '            "body": json.dumps({"error": str(e)})',
    '        }',
])
doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 10  functions/predict/entrypoint.py
# ══════════════════════════════════════════════════════════════════════════════
emit_file('10', 'functions/predict/entrypoint.py', 'OpenFaaS Flask HTTP Wrapper  [NEW]', 'NEW', [
    '"""',
    'Minimal HTTP wrapper that makes the predict function behave',
    'like an OpenFaaS function.',
    'OpenFaaS gateway calls POST / on this container and forwards',
    'the response.',
    '"""',
    'from flask import Flask, request, Response',
    'from handler import handle',
    '',
    'app = Flask(__name__)',
    '',
    '',
    'class Event:',
    '    def __init__(self, body: bytes):',
    '        self.body = body',
    '',
    '',
    '@app.post("/")',
    'def invoke():',
    '    event = Event(body=request.data)',
    '    result = handle(event, context=None)',
    '    return Response(',
    '        result["body"],',
    '        status=result["statusCode"],',
    '        mimetype="application/json",',
    '    )',
    '',
    '',
    'if __name__ == "__main__":',
    '    app.run(host="0.0.0.0", port=5000)',
])
doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 11  functions/predict/Dockerfile
# ══════════════════════════════════════════════════════════════════════════════
emit_file('11', 'functions/predict/Dockerfile', 'Serverless Function Container Image  [NEW]', 'NEW', [
    'FROM python:3.11-slim',
    '',
    'WORKDIR /app',
    '',
    'COPY requirements.txt .',
    'RUN pip install --no-cache-dir -r requirements.txt',
    '',
    'COPY handler.py entrypoint.py ./',
    '',
    'EXPOSE 5000',
    '',
    'CMD ["python", "entrypoint.py"]',
])
doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 12  functions/predict/requirements.txt
# ══════════════════════════════════════════════════════════════════════════════
emit_file('12', 'functions/predict/requirements.txt', 'Serverless Function Python Dependencies  [NEW]', 'NEW', [
    'flask==3.1.1',
    'scikit-learn==1.8.0',
    'joblib==1.5.3',
    'numpy==2.4.4',
])
doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 13  functions/stack.yml
# ══════════════════════════════════════════════════════════════════════════════
emit_file('13', 'functions/stack.yml', 'OpenFaaS CLI Deployment Manifest  [NEW]', 'NEW', [
    'version: 1.0',
    'provider:',
    '  name: openfaas',
    '  gateway: http://localhost:8080',
    '',
    'functions:',
    '  sepsis-predict:',
    '    lang: dockerfile',
    '    handler: ./predict',
    '    image: aiopscare/sepsis-predict:latest',
    '    environment:',
    '      MODEL_PATH: /app/ml/model.joblib',
    '    labels:',
    '      com.openfaas.scale.min: "1"',
    '      com.openfaas.scale.max: "5"',
])
doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 14  security/opa/policy.rego
# ══════════════════════════════════════════════════════════════════════════════
emit_file('14', 'security/opa/policy.rego', 'Open Policy Agent Authorization Rules  [NEW]', 'NEW', [
    'package aiopscare.authz',
    '',
    'default allow = false',
    '',
    '# Public endpoints — no token required',
    'allow {',
    '    input.method == "GET"',
    '    input.path == "/"',
    '}',
    '',
    'allow {',
    '    input.method == "GET"',
    '    input.path == "/metrics"',
    '}',
    '',
    'allow {',
    '    input.method == "GET"',
    '    input.path == "/docs"',
    '}',
    '',
    '# Protected endpoint — requires valid token with icu-staff role',
    'allow {',
    '    input.method == "POST"',
    '    input.path == "/predict"',
    '    token.payload.realm_access.roles[_] == "icu-staff"',
    '}',
    '',
    '# Decode the Bearer token',
    'token := {"payload": payload} {',
    '    [_, encoded] := split(input.token, " ")',
    '    [_, payload, _] := io.jwt.decode(encoded)',
    '}',
])
doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 15  security/keycloak/realm-export.json
# ══════════════════════════════════════════════════════════════════════════════
emit_file('15', 'security/keycloak/realm-export.json', 'Keycloak Identity Configuration — Realm, Roles & Users  [NEW]', 'NEW', [
    '{',
    '  "realm": "aiopscare",',
    '  "enabled": true,',
    '  "sslRequired": "none",',
    '  "registrationAllowed": false,',
    '  "clients": [',
    '    {',
    '      "clientId": "aiopscare-api",',
    '      "enabled": true,',
    '      "publicClient": false,',
    '      "secret": "aiopscare-secret",',
    '      "directAccessGrantsEnabled": true,',
    '      "serviceAccountsEnabled": true,',
    '      "standardFlowEnabled": true,',
    '      "redirectUris": ["*"],',
    '      "webOrigins": ["*"]',
    '    }',
    '  ],',
    '  "roles": {',
    '    "realm": [',
    '      { "name": "icu-staff", "description": "ICU medical staff — can call /predict" },',
    '      { "name": "admin",     "description": "System administrator" }',
    '    ]',
    '  },',
    '  "users": [',
    '    {',
    '      "username": "icu-user",',
    '      "enabled": true,',
    '      "credentials": [',
    '        { "type": "password", "value": "icu-password", "temporary": false }',
    '      ],',
    '      "realmRoles": ["icu-staff"]',
    '    }',
    '  ]',
    '}',
])
doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 16  infrastructure/docker-compose.yml
# ══════════════════════════════════════════════════════════════════════════════
emit_file('16', 'infrastructure/docker-compose.yml', 'All 11 Services Orchestration  [UPDATED]', 'UPDATED', [
    'version: "3.8"',
    '',
    'services:',
    '',
    '  # ──────────────────────────────────────────',
    '  # ZOOKEEPER  (Kafka dependency)',
    '  # ──────────────────────────────────────────',
    '  zookeeper:',
    '    image: confluentinc/cp-zookeeper:7.4.0',
    '    container_name: zookeeper',
    '    environment:',
    '      ZOOKEEPER_CLIENT_PORT: 2181',
    '    ports:',
    '      - "2181:2181"',
    '    restart: unless-stopped',
    '',
    '  # ──────────────────────────────────────────',
    '  # KAFKA  (Event streaming)',
    '  # ──────────────────────────────────────────',
    '  kafka:',
    '    image: confluentinc/cp-kafka:7.4.0',
    '    container_name: kafka',
    '    depends_on:',
    '      - zookeeper',
    '    ports:',
    '      - "9092:9092"',
    '    environment:',
    '      KAFKA_BROKER_ID: 1',
    '      KAFKA_ZOOKEEPER_CONNECT: zookeeper:2181',
    '      KAFKA_LISTENERS: PLAINTEXT://0.0.0.0:9092,PLAINTEXT_INTERNAL://0.0.0.0:29092',
    '      KAFKA_ADVERTISED_LISTENERS: PLAINTEXT://localhost:9092,PLAINTEXT_INTERNAL://kafka:29092',
    '      KAFKA_LISTENER_SECURITY_PROTOCOL_MAP: PLAINTEXT:PLAINTEXT,PLAINTEXT_INTERNAL:PLAINTEXT',
    '      KAFKA_INTER_BROKER_LISTENER_NAME: PLAINTEXT_INTERNAL',
    '      KAFKA_OFFSETS_TOPIC_REPLICATION_FACTOR: 1',
    '    volumes:',
    '      - kafka-data:/var/lib/kafka/data',
    '    restart: unless-stopped',
    '',
    '  # ──────────────────────────────────────────',
    '  # POSTGRES  (Keycloak database)',
    '  # ──────────────────────────────────────────',
    '  postgres:',
    '    image: postgres:15-alpine',
    '    container_name: postgres',
    '    environment:',
    '      POSTGRES_DB: keycloak',
    '      POSTGRES_USER: keycloak',
    '      POSTGRES_PASSWORD: keycloak_pass',
    '    volumes:',
    '      - postgres-data:/var/lib/postgresql/data',
    '    restart: unless-stopped',
    '',
    '  # ──────────────────────────────────────────',
    '  # KEYCLOAK  (Identity & access management)',
    '  # ──────────────────────────────────────────',
    '  keycloak:',
    '    image: quay.io/keycloak/keycloak:24.0.3',
    '    container_name: keycloak',
    '    command: start-dev --import-realm',
    '    depends_on:',
    '      - postgres',
    '    ports:',
    '      - "8082:8080"',
    '    environment:',
    '      KC_DB: postgres',
    '      KC_DB_URL: jdbc:postgresql://postgres:5432/keycloak',
    '      KC_DB_USERNAME: keycloak',
    '      KC_DB_PASSWORD: keycloak_pass',
    '      KEYCLOAK_ADMIN: admin',
    '      KEYCLOAK_ADMIN_PASSWORD: admin',
    '    volumes:',
    '      - ../security/keycloak:/opt/keycloak/data/import',
    '    restart: unless-stopped',
    '',
    '  # ──────────────────────────────────────────',
    '  # OPA  (Open Policy Agent — authorization)',
    '  # ──────────────────────────────────────────',
    '  opa:',
    '    image: openpolicyagent/opa:latest',
    '    container_name: opa',
    '    command:',
    '      - run',
    '      - --server',
    '      - --addr=0.0.0.0:8181',
    '      - --log-level=info',
    '      - /policies',
    '    ports:',
    '      - "8181:8181"',
    '    volumes:',
    '      - ../security/opa:/policies',
    '    restart: unless-stopped',
    '',
    '  # ──────────────────────────────────────────',
    '  # OPENFAAS GATEWAY  (Serverless gateway)',
    '  # ──────────────────────────────────────────',
    '  openfaas-gateway:',
    '    image: ghcr.io/openfaas/gateway:latest',
    '    container_name: openfaas-gateway',
    '    ports:',
    '      - "8080:8080"',
    '    environment:',
    '      basic_auth: "false"',
    '      functions_provider_url: "http://predict-fn:5000/"',
    '      direct_functions: "true"',
    '    depends_on:',
    '      - predict-fn',
    '    restart: unless-stopped',
    '',
    '  # ──────────────────────────────────────────',
    '  # PREDICT FUNCTION  (OpenFaaS serverless)',
    '  # ──────────────────────────────────────────',
    '  predict-fn:',
    '    build:',
    '      context: ../functions/predict',
    '      dockerfile: Dockerfile',
    '    container_name: predict-fn',
    '    environment:',
    '      MODEL_PATH: /app/ml/model.joblib',
    '    volumes:',
    '      - ../ml/model.joblib:/app/ml/model.joblib',
    '    restart: unless-stopped',
    '',
    '  # ──────────────────────────────────────────',
    '  # FASTAPI APP  (API Gateway + Auth)',
    '  # ──────────────────────────────────────────',
    '  fastapi:',
    '    build:',
    '      context: ..',
    '      dockerfile: Dockerfile',
    '    container_name: fastapi',
    '    ports:',
    '      - "8000:8000"',
    '    environment:',
    '      OPA_URL: http://opa:8181',
    '      JAEGER_ENDPOINT: http://jaeger:4318/v1/traces',
    '    depends_on:',
    '      - opa',
    '      - keycloak',
    '      - predict-fn',
    '    restart: unless-stopped',
    '',
    '  # ──────────────────────────────────────────',
    '  # PROMETHEUS  (Metrics collection)',
    '  # ──────────────────────────────────────────',
    '  prometheus:',
    '    image: prom/prometheus:latest',
    '    container_name: prometheus',
    '    ports:',
    '      - "9090:9090"',
    '    volumes:',
    '      - ../monitoring/prometheus/prometheus.yml:/etc/prometheus/prometheus.yml',
    '      - prometheus-data:/prometheus',
    '    restart: unless-stopped',
    '',
    '  # ──────────────────────────────────────────',
    '  # GRAFANA  (Metrics dashboards)',
    '  # ──────────────────────────────────────────',
    '  grafana:',
    '    image: grafana/grafana:latest',
    '    container_name: grafana',
    '    ports:',
    '      - "3000:3000"',
    '    volumes:',
    '      - grafana-data:/var/lib/grafana',
    '    depends_on:',
    '      - prometheus',
    '    restart: unless-stopped',
    '',
    '  # ──────────────────────────────────────────',
    '  # JAEGER  (Distributed tracing)',
    '  # ──────────────────────────────────────────',
    '  jaeger:',
    '    image: jaegertracing/all-in-one:latest',
    '    container_name: jaeger',
    '    ports:',
    '      - "16686:16686"',
    '      - "4318:4318"',
    '    restart: unless-stopped',
    '',
    'volumes:',
    '  kafka-data:',
    '  postgres-data:',
    '  prometheus-data:',
    '  grafana-data:',
])
doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 17  monitoring/prometheus/prometheus.yml
# ══════════════════════════════════════════════════════════════════════════════
emit_file('17', 'monitoring/prometheus/prometheus.yml', 'Prometheus Scrape Configuration  [UPDATED]', 'UPDATED', [
    'global:',
    '  scrape_interval: 15s',
    '',
    'scrape_configs:',
    '  - job_name: "fastapi"',
    '    static_configs:',
    '      - targets: ["fastapi:8000"]',
    '',
    '  - job_name: "openfaas-gateway"',
    '    static_configs:',
    '      - targets: ["openfaas-gateway:8080"]',
])
doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 18  scripts/zap_scan.sh
# ══════════════════════════════════════════════════════════════════════════════
emit_file('18', 'scripts/zap_scan.sh', 'OWASP ZAP Baseline Security Scan Script  [NEW]', 'NEW', [
    '#!/usr/bin/env bash',
    '# OWASP ZAP baseline security scan against the AIOpsCare API.',
    '# Usage: bash scripts/zap_scan.sh [target_url]',
    '',
    'TARGET=${1:-"http://localhost:8000"}',
    'REPORT_DIR="docs/security"',
    'REPORT_FILE="$REPORT_DIR/zap-report.html"',
    '',
    'mkdir -p "$REPORT_DIR"',
    '',
    'echo "Running OWASP ZAP baseline scan against $TARGET ..."',
    '',
    'docker run --rm \\',
    '  --network host \\',
    '  -v "$(pwd)/$REPORT_DIR:/zap/wrk:rw" \\',
    '  ghcr.io/zaproxy/zaproxy:stable \\',
    '  zap-baseline.py \\',
    '    -t "$TARGET" \\',
    '    -r "zap-report.html" \\',
    '    -I',
    '',
    'echo "Report saved to $REPORT_FILE"',
])
doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 19  Dockerfile (root)
# ══════════════════════════════════════════════════════════════════════════════
emit_file('19', 'Dockerfile', 'FastAPI Application Container Image  [FILLED]', 'FILLED', [
    'FROM python:3.11-slim',
    '',
    'WORKDIR /app',
    '',
    'COPY requirements.txt .',
    'RUN pip install --no-cache-dir -r requirements.txt',
    '',
    'COPY . .',
    '',
    'RUN python ml/train.py',
    '',
    'EXPOSE 8000',
    '',
    'CMD ["uvicorn", "app.main:app", "--host", "0.0.0.0", "--port", "8000"]',
])
doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 20  requirements.txt
# ══════════════════════════════════════════════════════════════════════════════
emit_file('20', 'requirements.txt', 'All Python Dependencies  [UPDATED]', 'UPDATED', [
    '# -----------------------------------',
    '# FastAPI & API Layer',
    '# -----------------------------------',
    '',
    'fastapi==0.136.1',
    'uvicorn[standard]==0.38.0',
    'httpx==0.28.1',
    'requests==2.32.5',
    '',
    '',
    '# -----------------------------------',
    '# Machine Learning',
    '# -----------------------------------',
    '',
    'scikit-learn==1.8.0',
    'pandas==2.3.3',
    'numpy==2.4.4',
    'scipy==1.16.3',
    'joblib==1.5.3',
    '',
    '',
    '# -----------------------------------',
    '# Kafka Event Streaming',
    '# -----------------------------------',
    '',
    'kafka-python==2.3.1',
    '',
    '',
    '# -----------------------------------',
    '# Monitoring & Observability',
    '# -----------------------------------',
    '',
    'prometheus-client==0.23.1',
    'prometheus-fastapi-instrumentator==7.1.0',
    'opentelemetry-api==1.41.1',
    'opentelemetry-sdk==1.41.1',
    'opentelemetry-distro==0.62b1',
    'opentelemetry-instrumentation==0.62b1',
    'opentelemetry-instrumentation-fastapi==0.62b1',
    'opentelemetry-instrumentation-asgi==0.62b1',
    'protobuf==5.29.4',
    'googleapis-common-protos==1.70.0',
    '',
    '',
    '# -----------------------------------',
    '# Authentication & Security',
    '# -----------------------------------',
    '',
    'python-jose==3.5.0',
    'passlib==1.7.4',
    'bcrypt==4.3.0',
    '',
    '',
    '# -----------------------------------',
    '# Testing',
    '# -----------------------------------',
    '',
    'pytest==9.0.3',
    'pytest-asyncio==1.3.0',
    '',
    '',
    '# -----------------------------------',
    '# Utility Dependencies',
    '# -----------------------------------',
    '',
    'pydantic==2.12.3',
    'typing_extensions==4.15.0',
    'python-dotenv==1.1.1',
    'opentelemetry-exporter-otlp==1.41.1',
    'opentelemetry-exporter-otlp-proto-grpc==1.41.1',
])

# ── Footer ────────────────────────────────────────────────────────────────────
footer = doc.sections[0].footer
fp = footer.paragraphs[0]
fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = fp.add_run('AIOpsCare  |  Complete Source Code  |  Final Year Project — Topic 7  |  '
               + datetime.datetime.now().strftime('%B %Y'))
r.font.name='Calibri'; r.font.size=Pt(8)
r.font.color.rgb = RGBColor(0x80, 0x80, 0x80)

# ── Save ──────────────────────────────────────────────────────────────────────
output = r'C:\Users\Acer\AIOpsCare_All_Files_Code.docx'
doc.save(output)
print(f'Saved: {output}')
