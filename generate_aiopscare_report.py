from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime

doc = Document()

# ── Page margins (1 inch all around) ──────────────────────────────────────────
section = doc.sections[0]
section.top_margin    = Inches(1)
section.bottom_margin = Inches(1)
section.left_margin   = Inches(1.2)
section.right_margin  = Inches(1.2)

# ── Helper: set paragraph shading ─────────────────────────────────────────────
def shade_paragraph(para, hex_color):
    pPr = para._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    pPr.append(shd)

# ── Helper: add horizontal rule ───────────────────────────────────────────────
def add_rule(doc, color="2E75B6"):
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), color)
    pBdr.append(bottom)
    pPr.append(pBdr)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(4)

# ── Helper: styled heading ────────────────────────────────────────────────────
def add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    run = p.add_run(text)
    if level == 1:
        run.font.size  = Pt(16)
        run.font.bold  = True
        run.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)
        p.paragraph_format.space_before = Pt(18)
        p.paragraph_format.space_after  = Pt(6)
        add_rule(doc)
    elif level == 2:
        run.font.size  = Pt(13)
        run.font.bold  = True
        run.font.color.rgb = RGBColor(0x2E, 0x75, 0xB6)
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after  = Pt(4)
    elif level == 3:
        run.font.size  = Pt(11)
        run.font.bold  = True
        run.font.color.rgb = RGBColor(0x40, 0x40, 0x40)
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after  = Pt(2)
    run.font.name = 'Calibri'

# ── Helper: body paragraph ────────────────────────────────────────────────────
def add_body(doc, text, bold_parts=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    if bold_parts:
        parts = text.split('**')
        for i, part in enumerate(parts):
            run = p.add_run(part)
            run.font.name = 'Calibri'
            run.font.size = Pt(11)
            run.bold = (i % 2 == 1)
    else:
        run = p.add_run(text)
        run.font.name = 'Calibri'
        run.font.size = Pt(11)
    return p

# ── Helper: bullet point ──────────────────────────────────────────────────────
def add_bullet(doc, text, indent=0):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.left_indent  = Inches(0.3 + indent * 0.3)
    p.paragraph_format.space_after  = Pt(3)
    run = p.add_run(text)
    run.font.name = 'Calibri'
    run.font.size = Pt(11)

# ── Helper: code block ────────────────────────────────────────────────────────
def add_code(doc, lines):
    for line in lines:
        p = doc.add_paragraph()
        shade_paragraph(p, 'F2F2F2')
        run = p.add_run(line if line else ' ')
        run.font.name = 'Courier New'
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after  = Pt(0)
        p.paragraph_format.left_indent  = Inches(0.2)

# ── Helper: two-column table ──────────────────────────────────────────────────
def add_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    if col_widths is None:
        col_widths = [Inches(2)] * len(headers)

    # Header row
    hdr_row = table.rows[0]
    for i, (cell, w) in enumerate(zip(hdr_row.cells, col_widths)):
        cell.width = w
        cell.paragraphs[0].clear()
        run = cell.paragraphs[0].add_run(headers[i])
        run.bold = True
        run.font.name = 'Calibri'
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
        # Blue background
        tc_pr = cell._tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), '2E75B6')
        tc_pr.append(shd)

    # Data rows
    for r_idx, row_data in enumerate(rows):
        row = table.rows[r_idx + 1]
        fill = 'DEEAF1' if r_idx % 2 == 0 else 'FFFFFF'
        for c_idx, (cell, w) in enumerate(zip(row.cells, col_widths)):
            cell.width = w
            cell.paragraphs[0].clear()
            run = cell.paragraphs[0].add_run(row_data[c_idx])
            run.font.name = 'Calibri'
            run.font.size = Pt(10)
            tc_pr = cell._tc.get_or_add_tcPr()
            shd = OxmlElement('w:shd')
            shd.set(qn('w:val'), 'clear')
            shd.set(qn('w:color'), 'auto')
            shd.set(qn('w:fill'), fill)
            tc_pr.append(shd)
    doc.add_paragraph()

# ══════════════════════════════════════════════════════════════════════════════
# TITLE PAGE
# ══════════════════════════════════════════════════════════════════════════════
doc.add_paragraph()
doc.add_paragraph()

title_p = doc.add_paragraph()
title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
title_run = title_p.add_run('AIOpsCare')
title_run.font.name  = 'Calibri'
title_run.font.size  = Pt(32)
title_run.font.bold  = True
title_run.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)

sub_p = doc.add_paragraph()
sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub_run = sub_p.add_run('Project Modification Report')
sub_run.font.name  = 'Calibri'
sub_run.font.size  = Pt(18)
sub_run.font.color.rgb = RGBColor(0x2E, 0x75, 0xB6)

desc_p = doc.add_paragraph()
desc_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
desc_run = desc_p.add_run('Serverless Architecture with Event-Driven AIOps,\nObservability & Security Integration')
desc_run.font.name  = 'Calibri'
desc_run.font.size  = Pt(13)
desc_run.font.color.rgb = RGBColor(0x40, 0x40, 0x40)

doc.add_paragraph()
doc.add_paragraph()

date_p = doc.add_paragraph()
date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
date_run = date_p.add_run(f'Prepared: {datetime.datetime.now().strftime("%B %d, %Y")}')
date_run.font.name  = 'Calibri'
date_run.font.size  = Pt(11)
date_run.font.color.rgb = RGBColor(0x70, 0x70, 0x70)

author_p = doc.add_paragraph()
author_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
author_run = author_p.add_run('AIOpsCare Final Year Project — Topic 7')
author_run.font.name  = 'Calibri'
author_run.font.size  = Pt(11)
author_run.font.color.rgb = RGBColor(0x70, 0x70, 0x70)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 1. EXECUTIVE SUMMARY
# ══════════════════════════════════════════════════════════════════════════════
add_heading(doc, '1. Executive Summary')
add_body(doc, (
    'The AIOpsCare project is a Real-Time ICU Monitoring and Sepsis Prediction Platform. '
    'The original codebase implemented a FastAPI API, a Kafka event pipeline, a scikit-learn ML '
    'model, and basic observability using Prometheus, Grafana, and Jaeger.'
))
add_body(doc, (
    'This report documents all modifications made to bring the project into full alignment with '
    'Topic 7 requirements: Designing Serverless Architectures with Event-Driven AIOps, '
    'Observability, and Security Integration.'
))
add_body(doc, 'The key additions were:')
add_bullet(doc, 'OpenFaaS serverless function platform — wraps the ML prediction logic as a serverless function')
add_bullet(doc, 'Keycloak — identity and access management providing JWT-based authentication')
add_bullet(doc, 'Open Policy Agent (OPA) — policy-based authorization for the /predict endpoint')
add_bullet(doc, 'OWASP ZAP — automated security scanning script')
add_bullet(doc, 'Kafka listener fix — corrected advertised listeners for inter-container communication')
add_bullet(doc, 'Prometheus fix — replaced hardcoded IP with Docker service name')
add_bullet(doc, 'Dockerfile — filled in to allow containerization of the FastAPI app')
add_bullet(doc, 'ML pipeline improvements — StandardScaler, cross-validation, fixed model save path')
doc.add_paragraph()

# ══════════════════════════════════════════════════════════════════════════════
# 2. REQUIREMENTS COVERAGE
# ══════════════════════════════════════════════════════════════════════════════
add_heading(doc, '2. Requirements Coverage')
add_body(doc, 'The table below maps each mandatory requirement from Topic 7 to the implementation:')
doc.add_paragraph()

add_table(doc,
    ['Requirement', 'Tool / Standard', 'Status', 'Location'],
    [
        ['5.1 Serverless Platform',    'OpenFaaS',               'Implemented', 'functions/, docker-compose.yml'],
        ['5.2 Event Streaming',        'Apache Kafka',           'Was present, fixed', 'events/, docker-compose.yml'],
        ['5.3 API Framework',          'FastAPI',                'Was present', 'app/'],
        ['5.4 Observability',          'Prometheus + Grafana + OpenTelemetry + Jaeger', 'Was present, fixed', 'monitoring/, app/main.py'],
        ['5.5 AIOps / ML',             'scikit-learn Pipeline',  'Improved', 'ml/train.py'],
        ['5.6 Security — Auth',        'Keycloak (JWT)',          'Added', 'security/keycloak/'],
        ['5.6 Security — Policy',      'Open Policy Agent',      'Added', 'security/opa/'],
        ['5.6 Security — Scanning',    'OWASP ZAP',              'Added', 'scripts/zap_scan.sh'],
        ['4.4 GitHub Repository',      'Git + documented code',  'Present', 'Root directory'],
    ],
    col_widths=[Inches(1.8), Inches(2.1), Inches(1.2), Inches(2.1)]
)

# ══════════════════════════════════════════════════════════════════════════════
# 3. NEW FILES CREATED
# ══════════════════════════════════════════════════════════════════════════════
add_heading(doc, '3. New Files Created')

# 3.1 OpenFaaS Function
add_heading(doc, '3.1  OpenFaaS Serverless Function', level=2)
add_body(doc, (
    'The prediction logic was extracted from the FastAPI app and packaged as a standalone '
    'OpenFaaS-compatible serverless function. This separates concerns: FastAPI handles auth '
    'and routing, while the function handles ML inference.'
))

add_heading(doc, 'functions/predict/handler.py', level=3)
add_body(doc, (
    'The core function logic. Loads the trained scikit-learn model and exposes a handle(event, context) '
    'interface — the standard OpenFaaS Python function signature. Accepts patient vitals (heart_rate, '
    'temperature, respiratory_rate), returns sepsis_prediction (0 or 1) and a list of clinical alerts.'
))
add_code(doc, [
    'def handle(event, context):',
    '    body = json.loads(event.body)',
    '    prediction = int(model.predict([[hr, temp, rr]])[0])',
    '    alerts = [...]  # threshold-based ICU alerts',
    '    return {"statusCode": 200, "body": json.dumps({...})}',
])
doc.add_paragraph()

add_heading(doc, 'functions/predict/entrypoint.py', level=3)
add_body(doc, (
    'A minimal Flask HTTP server that wraps handler.py so the OpenFaaS gateway can route '
    'HTTP POST requests to it. This follows the OpenFaaS Classic Watchdog pattern — '
    'each function runs as an HTTP microservice behind the gateway.'
))

add_heading(doc, 'functions/predict/Dockerfile', level=3)
add_body(doc, 'Builds the function container image from python:3.11-slim, installs scikit-learn and Flask, and starts the Flask entrypoint on port 5000.')

add_heading(doc, 'functions/stack.yml', level=3)
add_body(doc, (
    'OpenFaaS CLI deployment manifest. Defines the function name (sepsis-predict), '
    'the container image, and scaling labels (min 1 replica, max 5). Used with '
    'faas-cli deploy to register the function with the OpenFaaS gateway.'
))
doc.add_paragraph()

# 3.2 Security Layer
add_heading(doc, '3.2  Security Layer', level=2)

add_heading(doc, 'security/opa/policy.rego', level=3)
add_body(doc, (
    'Open Policy Agent authorization policy written in Rego. Defines which endpoints '
    'are public (GET / and /metrics) and which require a valid JWT token with the '
    'icu-staff Keycloak realm role (POST /predict). OPA decodes the Bearer token, '
    'inspects the realm_access.roles claim, and returns allow = true or false.'
))
add_code(doc, [
    'allow {',
    '    input.method == "POST"',
    '    input.path  == "/predict"',
    '    token.payload.realm_access.roles[_] == "icu-staff"',
    '}',
])
doc.add_paragraph()

add_heading(doc, 'security/keycloak/realm-export.json', level=3)
add_body(doc, (
    'Keycloak realm configuration file auto-imported at startup. Defines:'
))
add_bullet(doc, 'Realm name: aiopscare')
add_bullet(doc, 'Client: aiopscare-api (with client secret for token issuance)')
add_bullet(doc, 'Roles: icu-staff and admin')
add_bullet(doc, 'Test user: icu-user / icu-password with the icu-staff role assigned')
doc.add_paragraph()

add_heading(doc, 'app/auth.py', level=3)
add_body(doc, (
    'FastAPI dependency function that enforces authorization on protected endpoints. '
    'Extracts the Bearer token from the Authorization header, then calls OPA\'s REST API '
    '(POST /v1/data/aiopscare/authz/allow) with the token, method, and path. '
    'Returns HTTP 403 if OPA denies, or HTTP 503 if OPA is unreachable.'
))
add_code(doc, [
    'async def authorize(request: Request,',
    '                    credentials = Security(HTTPBearer())):',
    '    resp = await client.post(OPA_URL + "/v1/data/...", json=payload)',
    '    if not resp.json().get("result", False):',
    '        raise HTTPException(status_code=403)',
])
doc.add_paragraph()

# 3.3 OWASP ZAP
add_heading(doc, '3.3  OWASP ZAP Security Scan', level=2)
add_heading(doc, 'scripts/zap_scan.sh', level=3)
add_body(doc, (
    'Shell script that runs an OWASP ZAP baseline security scan against the FastAPI API '
    'using the official ZAP Docker image. Saves an HTML report to docs/security/zap-report.html. '
    'Run with: bash scripts/zap_scan.sh (defaults to http://localhost:8000).'
))
add_code(doc, [
    'docker run --rm --network host \\',
    '  ghcr.io/zaproxy/zaproxy:stable \\',
    '  zap-baseline.py -t http://localhost:8000 -r zap-report.html',
])
doc.add_paragraph()

# 3.4 Dockerfile
add_heading(doc, '3.4  Dockerfile (Root)', level=2)
add_body(doc, (
    'The root Dockerfile was empty. It is now filled in to build the FastAPI application '
    'container. It installs dependencies from requirements.txt, copies the project, '
    'runs ml/train.py to generate model.joblib inside the image, and starts the '
    'Uvicorn server on port 8000.'
))
add_code(doc, [
    'FROM python:3.11-slim',
    'WORKDIR /app',
    'COPY requirements.txt .',
    'RUN pip install --no-cache-dir -r requirements.txt',
    'COPY . .',
    'RUN python ml/train.py',
    'EXPOSE 8000',
    'CMD ["uvicorn", "app.main:app", "--host", "0.0.0.0", "--port", "8000"]',
])
doc.add_paragraph()

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 4. MODIFIED FILES
# ══════════════════════════════════════════════════════════════════════════════
add_heading(doc, '4. Modified Files')

# 4.1 docker-compose.yml
add_heading(doc, '4.1  infrastructure/docker-compose.yml', level=2)
add_body(doc, 'This was the most extensively modified file. The following changes were made:')

add_heading(doc, 'Kafka — Fixed advertised listeners', level=3)
add_body(doc, (
    'The original configuration advertised localhost:9092 for all listeners. Inside Docker, '
    'localhost refers to the container itself, so other containers (consumer, FastAPI) could '
    'not reach the Kafka broker. The fix introduces two separate listeners:'
))
add_bullet(doc, 'PLAINTEXT://localhost:9092 — for external access from the host machine')
add_bullet(doc, 'PLAINTEXT_INTERNAL://kafka:29092 — for inter-container communication')
add_code(doc, [
    '# BEFORE (broken for containers)',
    'KAFKA_ADVERTISED_LISTENERS: PLAINTEXT://localhost:9092',
    '',
    '# AFTER (fixed)',
    'KAFKA_LISTENERS: PLAINTEXT://0.0.0.0:9092,PLAINTEXT_INTERNAL://0.0.0.0:29092',
    'KAFKA_ADVERTISED_LISTENERS: PLAINTEXT://localhost:9092,PLAINTEXT_INTERNAL://kafka:29092',
    'KAFKA_INTER_BROKER_LISTENER_NAME: PLAINTEXT_INTERNAL',
])
doc.add_paragraph()

add_heading(doc, 'Added Services', level=3)
add_table(doc,
    ['Service', 'Image', 'Port', 'Purpose'],
    [
        ['postgres',         'postgres:15-alpine',               '(internal)', 'Database backend for Keycloak'],
        ['keycloak',         'keycloak/keycloak:24.0.3',          '8082',      'Identity provider — issues JWT tokens'],
        ['opa',              'openpolicyagent/opa:latest',        '8181',      'Policy engine — enforces access rules'],
        ['openfaas-gateway', 'ghcr.io/openfaas/gateway:0.27.3',  '8080',      'Serverless function gateway'],
        ['predict-fn',       'Built from functions/predict/',     '(internal)', 'Serverless ML function container'],
        ['fastapi',          'Built from root Dockerfile',        '8000',      'API gateway with auth + observability'],
    ],
    col_widths=[Inches(1.5), Inches(2.0), Inches(0.8), Inches(2.9)]
)

add_heading(doc, 'Added Named Volumes', level=3)
add_body(doc, 'Named Docker volumes were added so data persists across docker compose down/up cycles:')
add_bullet(doc, 'kafka-data — Kafka message log')
add_bullet(doc, 'postgres-data — Keycloak database')
add_bullet(doc, 'prometheus-data — Prometheus metrics storage')
add_bullet(doc, 'grafana-data — Grafana dashboards and settings')
doc.add_paragraph()

add_heading(doc, 'Added Restart Policies', level=3)
add_body(doc, 'restart: unless-stopped was added to all services to survive container crashes.')
doc.add_paragraph()

# 4.2 app/main.py
add_heading(doc, '4.2  app/main.py', level=2)
add_table(doc,
    ['Change', 'Before', 'After'],
    [
        ['Jaeger endpoint', 'Hardcoded http://localhost:4318/...', 'Read from JAEGER_ENDPOINT env var (default: jaeger:4318)'],
        ['Auth on /predict', 'No authentication', 'Depends(authorize) — OPA check required'],
        ['Imports', 'Scattered, with excessive comments', 'Grouped cleanly, no redundant comments'],
        ['predict function', 'def predict(...)', 'async def predict(...) — async for OPA HTTP call'],
    ],
    col_widths=[Inches(1.6), Inches(2.5), Inches(3.1)]
)

# 4.3 app/model_loader.py
add_heading(doc, '4.3  app/model_loader.py', level=2)
add_body(doc, 'Two problems were fixed:')
add_bullet(doc, (
    'Path bug: the original hardcoded "ml/model.joblib" (relative to cwd) would break '
    'when the app was started from a different directory. The fix resolves the path '
    'relative to the file\'s own location using os.path.dirname(__file__).'
))
add_bullet(doc, (
    'Silent crash: if model.joblib was missing, the app would crash at import time with '
    'a cryptic FileNotFoundError. The fix wraps the load in a try/except and raises a '
    'RuntimeError with a clear message: "Run python ml/train.py first."'
))
doc.add_paragraph()

# 4.4 ml/train.py
add_heading(doc, '4.4  ml/train.py', level=2)
add_table(doc,
    ['Problem', 'Fix'],
    [
        ['Model saved to wrong path (model.joblib in cwd)', 'Now saved to ml/model.joblib using os.path.dirname(__file__)'],
        ['No feature scaling — logistic regression on unscaled features', 'Added StandardScaler in a Pipeline before LogisticRegression'],
        ['train_test_split on 5 samples = 1 test point', 'Replaced with 5-fold cross-validation (cross_val_score)'],
        ['Only 5 training samples', 'Extended to 10 synthetic samples (note: real data needed for production)'],
    ],
    col_widths=[Inches(3.2), Inches(4.0)]
)

# 4.5 events/consumer.py
add_heading(doc, '4.5  events/consumer.py', level=2)
add_body(doc, (
    'The consumer previously loaded the ML model directly and ran predictions in-process. '
    'This was replaced with an HTTP call to the OpenFaaS gateway — which is the correct '
    'event-driven serverless pattern: Kafka events trigger the gateway, which invokes the '
    'stateless predict function.'
))
add_code(doc, [
    '# BEFORE — direct model loading (not serverless)',
    'model = joblib.load("ml/model.joblib")',
    'prediction = model.predict(input_data)',
    '',
    '# AFTER — calls OpenFaaS gateway (serverless trigger)',
    'response = httpx.post(FUNCTION_ENDPOINT, json=patient_data)',
    'result = response.json()',
])
add_body(doc, 'Additional improvements in the consumer:')
add_bullet(doc, 'KAFKA_BROKER read from environment variable (default: kafka:29092 for inter-container use)')
add_bullet(doc, 'OPENFAAS_URL read from environment variable')
add_bullet(doc, 'httpx.RequestError caught — no longer crashes on network failure')
doc.add_paragraph()

# 4.6 prometheus.yml
add_heading(doc, '4.6  monitoring/prometheus/prometheus.yml', level=2)
add_body(doc, (
    'The scrape target was hardcoded to a WSL internal IP address (172.21.100.237:8000) '
    'that is specific to one machine. Inside Docker Compose, services communicate by '
    'service name, not IP. The target was replaced with the Docker service name.'
))
add_code(doc, [
    '# BEFORE (machine-specific, broken in Docker)',
    'targets: ["172.21.100.237:8000"]',
    '',
    '# AFTER (portable, works inside Docker)',
    'targets: ["fastapi:8000"]',
])
add_body(doc, 'The OpenFaaS gateway metrics endpoint was also added as a second scrape job.')
doc.add_paragraph()

# 4.7 requirements.txt
add_heading(doc, '4.7  requirements.txt', level=2)
add_body(doc, 'Two packages were added to support the new modules:')
add_bullet(doc, 'httpx==0.28.1 — async HTTP client used by app/auth.py (OPA calls) and events/consumer.py (OpenFaaS calls)')
add_bullet(doc, 'flask==3.1.1 — HTTP server used by the OpenFaaS function entrypoint (functions/predict/entrypoint.py)')
doc.add_paragraph()

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 5. FINAL PROJECT STRUCTURE
# ══════════════════════════════════════════════════════════════════════════════
add_heading(doc, '5. Final Project Structure')
add_code(doc, [
    'AIOpsCare/',
    '├── app/',
    '│   ├── main.py          — FastAPI app with auth + tracing',
    '│   ├── schemas.py       — Pydantic request model',
    '│   ├── model_loader.py  — Robust model loader (fixed path)',
    '│   ├── alerts.py        — ICU threshold alert engine',
    '│   └── auth.py          — NEW: JWT + OPA authorization dependency',
    '├── functions/',
    '│   ├── predict/',
    '│   │   ├── handler.py      — NEW: OpenFaaS function logic',
    '│   │   ├── entrypoint.py   — NEW: Flask HTTP wrapper (watchdog pattern)',
    '│   │   ├── Dockerfile      — NEW: Function container image',
    '│   │   └── requirements.txt',
    '│   └── stack.yml           — NEW: OpenFaaS deployment manifest',
    '├── ml/',
    '│   ├── train.py         — UPDATED: Pipeline, scaler, cross-val, fixed path',
    '│   ├── preprocess.py',
    '│   └── evaluate.py',
    '├── events/',
    '│   ├── producer.py',
    '│   └── consumer.py      — UPDATED: Calls OpenFaaS gateway via HTTP',
    '├── security/',
    '│   ├── opa/',
    '│   │   └── policy.rego  — NEW: Authorization policy (Rego)',
    '│   └── keycloak/',
    '│       └── realm-export.json  — NEW: Keycloak realm + users + roles',
    '├── monitoring/',
    '│   └── prometheus/',
    '│       └── prometheus.yml  — UPDATED: Fixed target to service name',
    '├── infrastructure/',
    '│   └── docker-compose.yml  — UPDATED: +Keycloak, OPA, OpenFaaS, volumes',
    '├── scripts/',
    '│   └── zap_scan.sh      — NEW: OWASP ZAP baseline security scan',
    '├── docs/',
    '├── tests/',
    '├── Dockerfile           — UPDATED: Was empty, now complete',
    '└── requirements.txt     — UPDATED: Added httpx, flask',
])
doc.add_paragraph()

# ══════════════════════════════════════════════════════════════════════════════
# 6. ARCHITECTURE FLOW
# ══════════════════════════════════════════════════════════════════════════════
add_heading(doc, '6. Architecture Flow')

add_heading(doc, '6.1  HTTP Request Flow (Authenticated)', level=2)
add_code(doc, [
    'Client',
    '  │',
    '  │  1. GET token from Keycloak',
    '  │     POST http://localhost:8082/realms/aiopscare/protocol/openid-connect/token',
    '  │',
    '  │  2. Call protected API with JWT',
    '  │     POST http://localhost:8000/predict  +  Authorization: Bearer <token>',
    '  │',
    '  ▼',
    'FastAPI  (/predict)',
    '  │',
    '  │  3. app/auth.py calls OPA to check token role',
    '  │     POST http://opa:8181/v1/data/aiopscare/authz/allow',
    '  │',
    '  ▼',
    'OPA  ── policy.rego: is role == "icu-staff"? ──► 403 Forbidden (if not)',
    '  │',
    '  │  4. Authorized — response returned',
    '  ▼',
    'FastAPI returns {sepsis_prediction, alerts}',
])
doc.add_paragraph()

add_heading(doc, '6.2  Event-Driven Flow (Kafka → OpenFaaS)', level=2)
add_code(doc, [
    'events/producer.py',
    '  │',
    '  │  Sends patient vitals JSON to Kafka topic: patient-events',
    '  ▼',
    'Kafka Broker  (kafka:29092)',
    '  │',
    '  │  Consumer subscribes to patient-events',
    '  ▼',
    'events/consumer.py',
    '  │',
    '  │  Forwards event via HTTP POST to OpenFaaS gateway',
    '  ▼',
    'OpenFaaS Gateway  (openfaas-gateway:8080)',
    '  │',
    '  │  Routes request to registered function: sepsis-predict',
    '  ▼',
    'predict-fn container  (functions/predict/)',
    '  │',
    '  │  handler.py runs inference + alert logic',
    '  ▼',
    'Response: {sepsis_prediction: 1, alerts: ["High Fever Detected"]}',
    '  │',
    '  ▼',
    'consumer.py prints result — escalates alert if sepsis_prediction == 1',
])
doc.add_paragraph()

# ══════════════════════════════════════════════════════════════════════════════
# 7. HOW TO RUN
# ══════════════════════════════════════════════════════════════════════════════
add_heading(doc, '7. How to Run the Project')

add_heading(doc, 'Step 1 — Train the model', level=2)
add_code(doc, [
    'cd /home/raees/AIOpsCare',
    'python ml/train.py',
    '# Creates ml/model.joblib',
])
doc.add_paragraph()

add_heading(doc, 'Step 2 — Start all services', level=2)
add_code(doc, [
    'cd infrastructure',
    'docker compose up --build',
    '',
    '# Services started:',
    '#   Zookeeper     :2181',
    '#   Kafka         :9092  (external) / :29092 (internal)',
    '#   Postgres      (internal — for Keycloak)',
    '#   Keycloak      :8082',
    '#   OPA           :8181',
    '#   OpenFaaS GW   :8080',
    '#   predict-fn    (internal)',
    '#   FastAPI       :8000',
    '#   Prometheus    :9090',
    '#   Grafana       :3000',
    '#   Jaeger        :16686',
])
doc.add_paragraph()

add_heading(doc, 'Step 3 — Get a Keycloak JWT token', level=2)
add_code(doc, [
    'curl -X POST \\',
    '  http://localhost:8082/realms/aiopscare/protocol/openid-connect/token \\',
    '  -d "client_id=aiopscare-api" \\',
    '  -d "client_secret=aiopscare-secret" \\',
    '  -d "username=icu-user" \\',
    '  -d "password=icu-password" \\',
    '  -d "grant_type=password"',
    '',
    '# Copy the access_token value from the response',
])
doc.add_paragraph()

add_heading(doc, 'Step 4 — Call the prediction API', level=2)
add_code(doc, [
    'curl -X POST http://localhost:8000/predict \\',
    '  -H "Authorization: Bearer <access_token>" \\',
    '  -H "Content-Type: application/json" \\',
    '  -d \'{"heart_rate": 120, "temperature": 39.5, "respiratory_rate": 30}\'',
    '',
    '# Expected response:',
    '# {"sepsis_prediction": 1, "alerts": ["High Heart Rate Detected",',
    '#   "High Fever Detected", "Abnormal Respiratory Rate"]}',
])
doc.add_paragraph()

add_heading(doc, 'Step 5 — Run OWASP ZAP security scan', level=2)
add_code(doc, [
    'bash scripts/zap_scan.sh',
    '# Saves report to docs/security/zap-report.html',
])
doc.add_paragraph()

add_heading(doc, 'Service UIs', level=2)
add_table(doc,
    ['Service', 'URL', 'Credentials'],
    [
        ['Grafana',      'http://localhost:3000',  'admin / admin'],
        ['Prometheus',   'http://localhost:9090',  'None'],
        ['Jaeger UI',    'http://localhost:16686', 'None'],
        ['Keycloak',     'http://localhost:8082',  'admin / admin'],
        ['OPA API',      'http://localhost:8181',  'None'],
        ['OpenFaaS GW',  'http://localhost:8080',  'None (basic_auth disabled)'],
        ['FastAPI Docs', 'http://localhost:8000/docs', 'Requires JWT token'],
    ],
    col_widths=[Inches(1.5), Inches(2.3), Inches(3.4)]
)

# ══════════════════════════════════════════════════════════════════════════════
# FOOTER
# ══════════════════════════════════════════════════════════════════════════════
section = doc.sections[0]
footer = section.footer
footer_para = footer.paragraphs[0]
footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = footer_para.add_run('AIOpsCare | Final Year Project — Topic 7 | Serverless AIOps Architecture')
run.font.name = 'Calibri'
run.font.size = Pt(8)
run.font.color.rgb = RGBColor(0x70, 0x70, 0x70)

# ── Save ──────────────────────────────────────────────────────────────────────
output = r'C:\Users\Acer\AIOpsCare_Modification_Report.docx'
doc.save(output)
print(f'Document saved: {output}')
