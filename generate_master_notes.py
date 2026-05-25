from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime

doc = Document()

# ── Page setup ────────────────────────────────────────────────────────────────
sec = doc.sections[0]
sec.top_margin    = Inches(1)
sec.bottom_margin = Inches(1)
sec.left_margin   = Inches(1.2)
sec.right_margin  = Inches(1.2)

# ── Helpers ───────────────────────────────────────────────────────────────────
def shade_para(para, hex_color):
    pPr = para._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    pPr.append(shd)

def rule(color="2E75B6"):
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bot = OxmlElement('w:bottom')
    bot.set(qn('w:val'), 'single'); bot.set(qn('w:sz'), '6')
    bot.set(qn('w:space'), '1');    bot.set(qn('w:color'), color)
    pBdr.append(bot); pPr.append(pBdr)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(2)

def h1(text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.name = 'Calibri'; r.font.size = Pt(17); r.font.bold = True
    r.font.color.rgb = RGBColor(0x1F,0x49,0x7D)
    p.paragraph_format.space_before = Pt(16); p.paragraph_format.space_after = Pt(4)
    rule()

def h2(text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.name = 'Calibri'; r.font.size = Pt(13); r.font.bold = True
    r.font.color.rgb = RGBColor(0x2E,0x75,0xB6)
    p.paragraph_format.space_before = Pt(12); p.paragraph_format.space_after = Pt(3)

def h3(text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.name = 'Calibri'; r.font.size = Pt(11); r.font.bold = True
    r.font.color.rgb = RGBColor(0x70,0x30,0x00)
    p.paragraph_format.space_before = Pt(8); p.paragraph_format.space_after = Pt(2)

def body(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(5)
    r = p.add_run(text)
    r.font.name = 'Calibri'; r.font.size = Pt(11)
    return p

def mixed(parts):
    """parts = list of (text, bold, color_hex_or_None)"""
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(5)
    for text, bold, color in parts:
        r = p.add_run(text)
        r.font.name = 'Calibri'; r.font.size = Pt(11); r.bold = bold
        if color:
            r.font.color.rgb = RGBColor(int(color[0:2],16), int(color[2:4],16), int(color[4:6],16))
    return p

def note(text):
    p = doc.add_paragraph()
    shade_para(p, 'FFF3CD')
    r = p.add_run('  NOTE:  ' + text)
    r.font.name = 'Calibri'; r.font.size = Pt(10)
    r.font.color.rgb = RGBColor(0x7B,0x5C,0x00)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.left_indent = Inches(0.1)

def tip(text):
    p = doc.add_paragraph()
    shade_para(p, 'D4EDDA')
    r = p.add_run('  HOW TO MODIFY:  ' + text)
    r.font.name = 'Calibri'; r.font.size = Pt(10)
    r.font.color.rgb = RGBColor(0x15,0x52,0x24)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.left_indent = Inches(0.1)

def code(lines):
    for line in lines:
        p = doc.add_paragraph()
        shade_para(p, 'F0F0F0')
        r = p.add_run(line if line != '' else ' ')
        r.font.name = 'Courier New'; r.font.size = Pt(9)
        r.font.color.rgb = RGBColor(0x1A,0x1A,0x1A)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after  = Pt(0)
        p.paragraph_format.left_indent  = Inches(0.2)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)

def var_table(rows):
    table = doc.add_table(rows=1+len(rows), cols=3)
    table.style = 'Table Grid'
    widths = [Inches(1.7), Inches(1.4), Inches(4.1)]
    headers = ['Name', 'Type / Value', 'Purpose & How to Modify']
    hr = table.rows[0]
    for i,(cell,w) in enumerate(zip(hr.cells, widths)):
        cell.width = w
        cell.paragraphs[0].clear()
        run = cell.paragraphs[0].add_run(headers[i])
        run.bold=True; run.font.name='Calibri'; run.font.size=Pt(10)
        run.font.color.rgb = RGBColor(0xFF,0xFF,0xFF)
        tcPr = cell._tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'),'clear'); shd.set(qn('w:color'),'auto'); shd.set(qn('w:fill'),'1F497D')
        tcPr.append(shd)
    for ri, row_data in enumerate(rows):
        row = table.rows[ri+1]
        fill = 'EBF3FB' if ri%2==0 else 'FFFFFF'
        for ci,(cell,w) in enumerate(zip(row.cells, widths)):
            cell.width = w
            cell.paragraphs[0].clear()
            run = cell.paragraphs[0].add_run(row_data[ci])
            run.font.name='Calibri'; run.font.size=Pt(9)
            tcPr = cell._tc.get_or_add_tcPr()
            shd = OxmlElement('w:shd')
            shd.set(qn('w:val'),'clear'); shd.set(qn('w:color'),'auto'); shd.set(qn('w:fill'),fill)
            tcPr.append(shd)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)

def file_banner(filename, category):
    p = doc.add_paragraph()
    shade_para(p, '1F497D')
    r = p.add_run(f'  FILE:  {filename}   |   {category}')
    r.font.name='Calibri'; r.font.size=Pt(12); r.font.bold=True
    r.font.color.rgb = RGBColor(0xFF,0xFF,0xFF)
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after  = Pt(6)

def fn_banner(signature):
    p = doc.add_paragraph()
    shade_para(p, '2E75B6')
    r = p.add_run('  def  ' + signature)
    r.font.name='Courier New'; r.font.size=Pt(10); r.font.bold=True
    r.font.color.rgb = RGBColor(0xFF,0xFF,0xFF)
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after  = Pt(4)

def bullet(text, indent=0):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.left_indent = Inches(0.3 + indent*0.25)
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(text)
    r.font.name='Calibri'; r.font.size=Pt(11)

# ══════════════════════════════════════════════════════════════════════════════
# TITLE PAGE
# ══════════════════════════════════════════════════════════════════════════════
doc.add_paragraph()
doc.add_paragraph()
tp = doc.add_paragraph(); tp.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = tp.add_run('AIOpsCare'); r.font.name='Calibri'; r.font.size=Pt(36); r.font.bold=True
r.font.color.rgb = RGBColor(0x1F,0x49,0x7D)

sp = doc.add_paragraph(); sp.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = sp.add_run('Master Code Notes'); r.font.name='Calibri'; r.font.size=Pt(20)
r.font.color.rgb = RGBColor(0x2E,0x75,0xB6)

dp = doc.add_paragraph(); dp.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = dp.add_run('Complete Code Explanation — Functions, Variables & Modification Guide')
r.font.name='Calibri'; r.font.size=Pt(12)
r.font.color.rgb = RGBColor(0x50,0x50,0x50)

doc.add_paragraph()
dt = doc.add_paragraph(); dt.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = dt.add_run(datetime.datetime.now().strftime('%B %d, %Y'))
r.font.name='Calibri'; r.font.size=Pt(11); r.font.color.rgb=RGBColor(0x80,0x80,0x80)
doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# QUICK REFERENCE INDEX
# ══════════════════════════════════════════════════════════════════════════════
h1('Files Covered in This Document')
files_index = [
    ('app/main.py',                    'FastAPI application entry point'),
    ('app/schemas.py',                 'Pydantic request/response models'),
    ('app/model_loader.py',            'ML model loading at startup'),
    ('app/alerts.py',                  'ICU clinical alert engine'),
    ('app/auth.py',                    'JWT + OPA authorization middleware  [NEW]'),
    ('ml/train.py',                    'Model training pipeline'),
    ('events/producer.py',             'Kafka event producer'),
    ('events/consumer.py',             'Kafka consumer + OpenFaaS trigger  [UPDATED]'),
    ('functions/predict/handler.py',   'OpenFaaS serverless function  [NEW]'),
    ('functions/predict/entrypoint.py','OpenFaaS HTTP wrapper  [NEW]'),
    ('security/opa/policy.rego',       'Open Policy Agent authorization rules  [NEW]'),
    ('security/keycloak/realm-export.json', 'Keycloak identity configuration  [NEW]'),
    ('infrastructure/docker-compose.yml', 'All services orchestration  [UPDATED]'),
    ('monitoring/prometheus/prometheus.yml', 'Metrics scrape config  [UPDATED]'),
    ('scripts/zap_scan.sh',            'OWASP ZAP security scan  [NEW]'),
    ('Dockerfile',                     'FastAPI container image  [FILLED]'),
]
for f, desc in files_index:
    mixed([(f'  {f}', True, '1F497D'), (f'  —  {desc}', False, None)])
doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 1. app/main.py
# ══════════════════════════════════════════════════════════════════════════════
file_banner('app/main.py', 'FastAPI Application Entry Point')
body('This is the heart of the API server. It starts the web application, sets up monitoring and tracing, and defines all HTTP endpoints. Every HTTP request to the system enters through this file.')

h2('Imports Explained')
code([
    'import os',
    'from fastapi import FastAPI, Depends',
    'from app.schemas import PatientData',
    'from app.model_loader import model',
    'from app.alerts import generate_alerts',
    'from app.auth import authorize',
    'from prometheus_fastapi_instrumentator import Instrumentator',
    'from opentelemetry import trace',
    'from opentelemetry.sdk.resources import Resource',
    'from opentelemetry.sdk.trace import TracerProvider',
    'from opentelemetry.sdk.trace.export import BatchSpanProcessor',
    'from opentelemetry.exporter.otlp.proto.http.trace_exporter import OTLPSpanExporter',
    'from opentelemetry.instrumentation.fastapi import FastAPIInstrumentor',
])
var_table([
    ['os', 'stdlib module', 'Used to read environment variables with os.getenv(). Change config without touching code.'],
    ['FastAPI', 'class', 'The web framework. Creates the app object that handles all HTTP routing.'],
    ['Depends', 'function', 'FastAPI dependency injection. Runs authorize() before the endpoint function executes.'],
    ['PatientData', 'Pydantic model', 'Defines the shape of the JSON body expected by /predict. Imported from schemas.py.'],
    ['model', 'sklearn Pipeline', 'The trained ML model loaded at startup from ml/model.joblib.'],
    ['generate_alerts', 'function', 'Checks vitals thresholds and returns a list of alert strings.'],
    ['authorize', 'async function', 'Validates JWT token via OPA before allowing access to /predict.'],
    ['Instrumentator', 'class', 'Automatically exposes /metrics endpoint for Prometheus to scrape.'],
    ['trace', 'OTel module', 'OpenTelemetry tracing API — used to create and manage trace spans.'],
    ['TracerProvider', 'class', 'Configures where traces are sent (to Jaeger in this case).'],
    ['BatchSpanProcessor', 'class', 'Batches trace spans before exporting — more efficient than sending one-by-one.'],
    ['OTLPSpanExporter', 'class', 'Sends trace data to Jaeger over HTTP using the OTLP protocol.'],
    ['FastAPIInstrumentor', 'class', 'Auto-instruments FastAPI so every request gets a trace span automatically.'],
])

h2('Environment Variable')
code(['JAEGER_ENDPOINT = os.getenv("JAEGER_ENDPOINT", "http://jaeger:4318/v1/traces")'])
body('os.getenv(variable_name, default) reads a value from the system environment. If the variable is not set, it uses the default. This allows the same code to run locally and inside Docker without changes.')
tip('To point to a different Jaeger: set JAEGER_ENDPOINT=http://your-host:4318/v1/traces in your .env file or docker-compose environment section.')

h2('App Creation')
code([
    'app = FastAPI(',
    '    title="AIOpsCare",',
    '    description="Real-Time ICU Monitoring & Sepsis Prediction Platform",',
    '    version="1.0.0",',
    ')',
])
body('Creates the FastAPI application instance. title and description appear in the auto-generated API docs at http://localhost:8000/docs. The app object is what Uvicorn runs.')
tip('Change title/description/version here to match your project name for the docs page.')

h2('Prometheus Instrumentation')
code(['Instrumentator().instrument(app).expose(app)'])
body('One-liner that does two things: instrument(app) wraps every route to collect request count, latency, and status code metrics. expose(app) adds a GET /metrics endpoint that Prometheus scrapes.')
note('This line must run before any route definitions so all routes are instrumented.')

h2('OpenTelemetry Tracing Setup')
code([
    'if not isinstance(trace.get_tracer_provider(), TracerProvider):',
    '    trace.set_tracer_provider(',
    '        TracerProvider(resource=Resource.create({"service.name": "aiopscare-fastapi"}))',
    '    )',
    '    span_processor = BatchSpanProcessor(OTLPSpanExporter(endpoint=JAEGER_ENDPOINT))',
    '    trace.get_tracer_provider().add_span_processor(span_processor)',
    '    FastAPIInstrumentor.instrument_app(app)',
])
body('The if-check prevents double-initialization if the module is imported multiple times. service.name is the label that appears in Jaeger UI to identify which service sent the trace.')
tip('Change "aiopscare-fastapi" to any name — it becomes the service name shown in Jaeger.')

fn_banner('home() -> dict')
code(['@app.get("/")', 'def home():', '    return {"message": "Welcome to AIOpsCare"}'])
body('Simple health-check endpoint. GET / returns a welcome message. No authentication required. Used to verify the server is running.')
tip('Add more fields here like version, status, uptime for a richer health check.')

fn_banner('predict(data: PatientData, token: str = Depends(authorize)) -> dict')
code([
    '@app.post("/predict")',
    'async def predict(data: PatientData, token: str = Depends(authorize)):',
    '    with tracer.start_as_current_span("predict-sepsis"):',
    '        prediction = model.predict([[data.heart_rate, data.temperature, data.respiratory_rate]])',
    '        result = int(prediction[0])',
    '        alerts = generate_alerts(data)',
    '        return {"sepsis_prediction": result, "alerts": alerts}',
])
var_table([
    ['data', 'PatientData', 'Automatically parsed and validated from the JSON request body.'],
    ['token', 'str (JWT)', 'Injected by Depends(authorize) — the raw JWT token string after OPA check passes.'],
    ['tracer.start_as_current_span', 'context manager', 'Creates a named span in Jaeger. Everything inside the with block is traced.'],
    ['"predict-sepsis"', 'str', 'The span name shown in Jaeger. Change to describe the operation.'],
    ['model.predict([[...]])', 'ndarray', 'Runs inference. Input must be a 2D list [[f1, f2, f3]]. Returns array of predictions.'],
    ['int(prediction[0])', 'int', 'Converts numpy int64 to Python int so it can be serialized to JSON. 0=no sepsis, 1=sepsis.'],
    ['generate_alerts(data)', 'list[str]', 'Returns a list of alert messages based on threshold checks in alerts.py.'],
])
tip('To add a new vital sign: add it to PatientData schema, pass it to model.predict, and add a threshold check in generate_alerts.')
doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 2. app/schemas.py
# ══════════════════════════════════════════════════════════════════════════════
file_banner('app/schemas.py', 'Pydantic Data Models')
body('Defines the structure of data coming into and going out of the API. Pydantic automatically validates the request body — if a field is missing or the wrong type, FastAPI returns a 422 error before your code runs.')

h2('PatientData Class')
code([
    'from pydantic import BaseModel',
    '',
    'class PatientData(BaseModel):',
    '    heart_rate: float',
    '    temperature: float',
    '    respiratory_rate: float',
])
var_table([
    ['BaseModel', 'Pydantic class', 'Parent class that provides automatic JSON parsing, validation, and serialization.'],
    ['heart_rate', 'float', 'Beats per minute. Must be a number. No range validation currently — can send -999.'],
    ['temperature', 'float', 'Body temperature in Celsius. Same note — no range constraint applied yet.'],
    ['respiratory_rate', 'float', 'Breaths per minute. Same note.'],
])
tip('Add validation constraints: heart_rate: float = Field(gt=0, lt=300, description="BPM"). Import Field from pydantic first.')
tip('To add a new vital (e.g., blood pressure): add bp_systolic: float to this class, then update main.py, alerts.py, and train.py to use it.')
doc.add_paragraph()

# ══════════════════════════════════════════════════════════════════════════════
# 3. app/model_loader.py
# ══════════════════════════════════════════════════════════════════════════════
file_banner('app/model_loader.py', 'ML Model Loader')
body('Loads the trained scikit-learn model once when the app starts. All prediction requests share this single loaded model — this is efficient because loading a model file is slow and should not happen per-request.')

h2('Full Code Explained')
code([
    'import os',
    'import joblib',
    '',
    '_model_path = os.path.abspath(',
    '    os.path.join(os.path.dirname(__file__), "..", "ml", "model.joblib")',
    ')',
    '',
    'try:',
    '    model = joblib.load(_model_path)',
    'except FileNotFoundError:',
    '    raise RuntimeError(',
    '        f"Model not found at {_model_path}. Run python ml/train.py first."',
    '    )',
])
var_table([
    ['__file__', 'str (built-in)', 'Python built-in — the absolute path of the current file (model_loader.py).'],
    ['os.path.dirname(__file__)', 'str', 'The directory containing model_loader.py — which is the app/ folder.'],
    ['os.path.join(..., "..", "ml", "model.joblib")', 'str', 'Navigates: app/ -> up one level -> ml/ -> model.joblib. Works regardless of where you run the app from.'],
    ['os.path.abspath(...)', 'str', 'Resolves any ".." and returns a clean absolute path like /home/raees/AIOpsCare/ml/model.joblib.'],
    ['joblib.load(_model_path)', 'sklearn Pipeline', 'Deserializes the model file into a Python object. The returned object has a .predict() method.'],
    ['model', 'module-level variable', 'Stored at module level so it is loaded once and reused by every request to /predict.'],
])
tip('To use a different model file: change the path string in os.path.join. To support multiple models: make this a function that accepts a model name parameter.')
doc.add_paragraph()

# ══════════════════════════════════════════════════════════════════════════════
# 4. app/alerts.py
# ══════════════════════════════════════════════════════════════════════════════
file_banner('app/alerts.py', 'ICU Clinical Alert Engine')
body('Pure Python rule engine. Checks each vital sign against clinical thresholds and returns a list of human-readable alert strings. No ML involved — these are hard-coded medical rules.')

fn_banner('generate_alerts(data: PatientData) -> list[str]')
code([
    'HEART_RATE_HIGH  = 100',
    'TEMPERATURE_HIGH = 38.0',
    'RESP_RATE_HIGH   = 24',
    '',
    'def generate_alerts(data):',
    '    alerts = []',
    '    if data.heart_rate > 100:',
    '        alerts.append("High Heart Rate Detected")',
    '    if data.temperature > 38:',
    '        alerts.append("High Fever Detected")',
    '    if data.respiratory_rate > 24:',
    '        alerts.append("Abnormal Respiratory Rate")',
    '    return alerts',
])
var_table([
    ['data', 'PatientData', 'The validated patient vitals object passed from the /predict endpoint.'],
    ['alerts', 'list[str]', 'Starts empty. An alert string is appended each time a condition is true.'],
    ['> 100 (heart rate)', 'threshold', 'Tachycardia threshold. Change the number to adjust sensitivity.'],
    ['> 38 (temperature)', 'threshold', 'Fever threshold in Celsius. For Fahrenheit, change to > 100.4.'],
    ['> 24 (resp. rate)', 'threshold', 'Tachypnea threshold in breaths/min.'],
])
tip('Add low-value alerts by adding elif or additional if blocks: if data.heart_rate < 40: alerts.append("Bradycardia Detected").')
tip('Add severity: instead of a plain string, append a dict like {"alert": "High Heart Rate", "severity": "critical"}.')
doc.add_paragraph()

# ══════════════════════════════════════════════════════════════════════════════
# 5. app/auth.py  [NEW]
# ══════════════════════════════════════════════════════════════════════════════
file_banner('app/auth.py', 'JWT + OPA Authorization  [NEW FILE]')
body('This is the security gateway. Every request to /predict must pass through this function. It extracts the Bearer token from the HTTP header and asks OPA whether the token is allowed to call the endpoint. If OPA says no, the request is blocked before the prediction runs.')

h2('Imports Explained')
code([
    'import os',
    'import httpx',
    'from fastapi import HTTPException, Security, Request',
    'from fastapi.security import HTTPBearer, HTTPAuthorizationCredentials',
])
var_table([
    ['httpx', 'async HTTP client', 'Used to make async HTTP calls to OPA. Faster than requests for async code.'],
    ['HTTPException', 'FastAPI class', 'Raises an HTTP error response. status_code sets the code, detail sets the message.'],
    ['Security', 'FastAPI function', 'Like Depends but marks the dependency as a security scheme (shows in /docs).'],
    ['Request', 'FastAPI class', 'Gives access to full request info: method, path, headers, etc.'],
    ['HTTPBearer', 'class', 'Extracts the Bearer token from the Authorization header automatically.'],
    ['HTTPAuthorizationCredentials', 'class', 'Holds scheme ("Bearer") and credentials (the raw token string).'],
])

h2('Configuration')
code(['OPA_URL  = os.getenv("OPA_URL", "http://opa:8181")', 'security = HTTPBearer()'])
var_table([
    ['OPA_URL', 'str', 'Address of the OPA server. Default points to the Docker service name "opa". Change for local dev: set OPA_URL=http://localhost:8181.'],
    ['security', 'HTTPBearer instance', 'Reusable extractor object. Defined once at module level — not inside the function.'],
])

fn_banner('authorize(request, credentials) -> str  [async]')
code([
    'async def authorize(',
    '    request: Request,',
    '    credentials: HTTPAuthorizationCredentials = Security(security),',
    '):',
    '    token = credentials.credentials',
    '    payload = {',
    '        "input": {',
    '            "token": f"Bearer {token}",',
    '            "method": request.method,',
    '            "path": request.url.path,',
    '        }',
    '    }',
    '    try:',
    '        async with httpx.AsyncClient(timeout=5.0) as client:',
    '            resp = await client.post(',
    '                f"{OPA_URL}/v1/data/aiopscare/authz/allow",',
    '                json=payload,',
    '            )',
    '        if not resp.json().get("result", False):',
    '            raise HTTPException(status_code=403, detail="Forbidden: insufficient role")',
    '    except httpx.RequestError:',
    '        raise HTTPException(status_code=503, detail="Authorization service unavailable")',
    '    return token',
])
var_table([
    ['credentials.credentials', 'str', 'The raw JWT token string extracted from "Authorization: Bearer <token>" header.'],
    ['payload["input"]', 'dict', 'The data OPA evaluates against the policy. Must match field names used in policy.rego.'],
    ['"token"', 'str', 'Full "Bearer <token>" string. OPA policy calls io.jwt.decode() on this.'],
    ['"method"', 'str', 'HTTP method (GET, POST). Used in policy to allow GET / without auth.'],
    ['"path"', 'str', 'URL path (/predict). Used to apply different rules per endpoint.'],
    ['httpx.AsyncClient(timeout=5.0)', 'context manager', 'Creates an HTTP client with a 5-second timeout. Closed automatically after the with block.'],
    ['/v1/data/aiopscare/authz/allow', 'OPA REST endpoint', 'This URL maps to the package aiopscare.authz, rule allow in policy.rego.'],
    ['resp.json().get("result", False)', 'bool', 'OPA returns {"result": true} or {"result": false}. Default False if key missing.'],
    ['HTTPException(403)', 'exception', 'FastAPI catches this and returns HTTP 403 Forbidden to the client.'],
    ['HTTPException(503)', 'exception', 'Returned when OPA is unreachable — service unavailable, not an auth failure.'],
    ['return token', 'str', 'Returns the token to the calling endpoint. main.py receives it as the token parameter.'],
])
tip('To disable auth for development: comment out the Depends(authorize) in main.py — do NOT delete this file.')
tip('To add a new protected endpoint: just add token: str = Depends(authorize) as a parameter to any route function.')
doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 6. ml/train.py
# ══════════════════════════════════════════════════════════════════════════════
file_banner('ml/train.py', 'Model Training Pipeline')
body('Trains the sepsis prediction ML model and saves it to disk. Run this script once before starting the app. The saved file (model.joblib) is what model_loader.py loads at runtime.')

h2('Imports')
code([
    'import os',
    'from sklearn.linear_model import LogisticRegression',
    'from sklearn.preprocessing import StandardScaler',
    'from sklearn.pipeline import Pipeline',
    'from sklearn.model_selection import cross_val_score',
    'import pandas as pd',
    'import joblib',
])
var_table([
    ['LogisticRegression', 'sklearn model', 'Binary classifier. Predicts 0 (no sepsis) or 1 (sepsis). Good for small datasets.'],
    ['StandardScaler', 'sklearn transformer', 'Scales features to mean=0, std=1. Critical when features have different ranges (HR:70-120, Temp:36-40).'],
    ['Pipeline', 'sklearn class', 'Chains scaler + classifier into one object. When you call pipeline.predict(), it scales first then classifies.'],
    ['cross_val_score', 'function', 'Evaluates model by splitting data into 5 folds, training on 4, testing on 1, rotating. More reliable than a single split.'],
    ['joblib', 'library', 'Serializes (saves) and deserializes (loads) Python objects efficiently. Better than pickle for numpy arrays.'],
])

h2('Training Data')
code([
    'data = {',
    '    "heart_rate":       [80, 95, 110, 70, 120, 85, 105, 72, 115, 90],',
    '    "temperature":      [36.5, 38.2, 39.1, 36.8, 40.0, 37.0, 38.8, 36.6, 39.5, 37.5],',
    '    "respiratory_rate": [18, 22, 30, 16, 35, 19, 28, 15, 32, 21],',
    '    "sepsis":           [0, 1, 1, 0, 1, 0, 1, 0, 1, 0],',
    '}',
])
note('This is synthetic data with 10 samples. For a real project use MIMIC-III or eICU datasets which have thousands of real ICU patient records.')
tip('To use a real CSV: replace the data dict with df = pd.read_csv("data/icu_patients.csv") and adjust column names.')

h2('Feature and Label Split')
code([
    'X = df[["heart_rate", "temperature", "respiratory_rate"]]',
    'y = df["sepsis"]',
])
var_table([
    ['X', 'DataFrame (n x 3)', 'Input features — the vitals. Shape: (10 rows, 3 columns). Column order matters for prediction.'],
    ['y', 'Series (n,)', 'Target labels — 0 or 1. Must be a 1D array of the same length as X.'],
])
tip('To add a feature (e.g., blood pressure): add it as a column in data, add it to X columns list, add it to PatientData schema, and pass it in model.predict() in handler.py and main.py.')

h2('Pipeline and Training')
code([
    'pipeline = Pipeline([',
    '    ("scaler", StandardScaler()),',
    '    ("classifier", LogisticRegression(random_state=42)),',
    '])',
    'scores = cross_val_score(pipeline, X, y, cv=5, scoring="accuracy")',
    'pipeline.fit(X, y)',
])
var_table([
    ['Pipeline([...])', 'list of tuples', 'Each tuple is (name, transformer/estimator). Name can be anything — used to access steps later.'],
    ['random_state=42', 'int', 'Makes results reproducible. Any integer works — 42 is convention. Change freely.'],
    ['cv=5', 'int', 'Number of cross-validation folds. Increase to 10 for more reliable estimates on small data.'],
    ['scoring="accuracy"', 'str', 'Metric used. Change to "roc_auc" for imbalanced classes or "f1" for recall focus.'],
    ['pipeline.fit(X, y)', 'method', 'Trains the full pipeline on ALL data after cross-validation gives a performance estimate.'],
])

h2('Saving the Model')
code([
    'output_path = os.path.join(os.path.dirname(__file__), "model.joblib")',
    'joblib.dump(pipeline, output_path)',
])
body('os.path.dirname(__file__) is the ml/ directory. So model.joblib is saved inside ml/. This matches what model_loader.py expects.')
tip('To save with a version: change to f"model_v{version}.joblib" and update model_loader.py to point to the new filename.')
doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 7. events/producer.py
# ══════════════════════════════════════════════════════════════════════════════
file_banner('events/producer.py', 'Kafka Event Producer')
body('Simulates a patient monitoring device sending vitals data to Kafka. In a real system this would be replaced by actual ICU equipment streaming data continuously.')

h2('Full Code Explained')
code([
    'from kafka import KafkaProducer',
    'import json',
    '',
    'producer = KafkaProducer(',
    '    bootstrap_servers="localhost:9092",',
    '    value_serializer=lambda v: json.dumps(v).encode("utf-8")',
    ')',
    '',
    'patient_event = {',
    '    "heart_rate": 120,',
    '    "temperature": 39.2,',
    '    "respiratory_rate": 30',
    '}',
    '',
    'producer.send("patient-events", patient_event)',
    'producer.flush()',
])
var_table([
    ['KafkaProducer', 'class', 'Connects to Kafka broker and sends messages to topics.'],
    ['bootstrap_servers', 'str', 'Address of the Kafka broker. Use "kafka:29092" when running inside Docker Compose.'],
    ['value_serializer', 'lambda', 'Converts Python dict to JSON bytes before sending. Kafka only transmits bytes.'],
    ['json.dumps(v)', 'function', 'Serializes dict to JSON string. v is the message value passed to producer.send().'],
    ['.encode("utf-8")', 'method', 'Converts string to bytes (required by Kafka).'],
    ['"patient-events"', 'str', 'The Kafka topic name. Consumer must subscribe to the same topic name.'],
    ['producer.flush()', 'method', 'Waits until all buffered messages are sent before the script exits. Important for one-shot scripts.'],
])
tip('To send continuous data: wrap producer.send() in a while True loop with time.sleep(1) to simulate a stream.')
tip('To run inside Docker: change bootstrap_servers to "kafka:29092" (internal listener).')

# ══════════════════════════════════════════════════════════════════════════════
# 8. events/consumer.py  [UPDATED]
# ══════════════════════════════════════════════════════════════════════════════
file_banner('events/consumer.py', 'Kafka Consumer + OpenFaaS Trigger  [UPDATED]')
body('Listens for patient events from Kafka and forwards each event to the OpenFaaS serverless function via HTTP. This is the event-driven trigger pattern — Kafka events drive serverless function invocations.')

h2('Full Code Explained')
code([
    'import os, json, httpx',
    'from kafka import KafkaConsumer',
    '',
    'KAFKA_BROKER      = os.getenv("KAFKA_BROKER", "kafka:29092")',
    'OPENFAAS_URL      = os.getenv("OPENFAAS_URL", "http://openfaas-gateway:8080")',
    'FUNCTION_ENDPOINT = f"{OPENFAAS_URL}/function/sepsis-predict"',
    '',
    'consumer = KafkaConsumer(',
    '    "patient-events",',
    '    bootstrap_servers=KAFKA_BROKER,',
    '    auto_offset_reset="earliest",',
    '    value_deserializer=lambda x: json.loads(x.decode("utf-8"))',
    ')',
    '',
    'for message in consumer:',
    '    patient_data = message.value',
    '    response = httpx.post(FUNCTION_ENDPOINT, json=patient_data, timeout=10.0)',
    '    result = response.json()',
    '    if result["sepsis_prediction"] == 1:',
    '        print("ACTION: High sepsis risk -- escalating to ICU team.")',
])
var_table([
    ['KAFKA_BROKER', 'str (env var)', 'Kafka address. Default kafka:29092 works inside Docker. Use localhost:9092 for local dev.'],
    ['OPENFAAS_URL', 'str (env var)', 'OpenFaaS gateway URL. Change if gateway runs on a different host or port.'],
    ['FUNCTION_ENDPOINT', 'str', 'Full URL to invoke the sepsis-predict function. /function/{function-name} is the OpenFaaS routing pattern.'],
    ['"patient-events"', 'str', 'Topic to subscribe to. Must match the topic used in producer.py.'],
    ['auto_offset_reset="earliest"', 'str', 'If no offset recorded, start reading from the beginning of the topic. Use "latest" to only read new messages.'],
    ['value_deserializer', 'lambda', 'Reverses what the producer did: bytes -> string -> dict.'],
    ['for message in consumer', 'infinite loop', 'Blocks and waits for messages. Each iteration processes one message.'],
    ['message.value', 'dict', 'The deserialized patient event dict from Kafka.'],
    ['httpx.post(..., json=...)', 'HTTP call', 'Sends the patient data to the OpenFaaS function as a JSON POST request.'],
    ['timeout=10.0', 'float', 'Seconds to wait for the function to respond. Increase for slow models.'],
    ['response.json()', 'dict', 'Parses the JSON response from the function into a Python dict.'],
])
tip('To add logging to a file: replace print() with Python logging module calls. Use logging.basicConfig(filename="consumer.log").')
tip('To handle errors: wrap the httpx.post call in try/except httpx.RequestError to avoid crashing when OpenFaaS is down.')
doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 9. functions/predict/handler.py  [NEW]
# ══════════════════════════════════════════════════════════════════════════════
file_banner('functions/predict/handler.py', 'OpenFaaS Serverless Function Logic  [NEW FILE]')
body('This is the serverless function. It is identical in logic to what was previously in main.py, but now runs as an independent, stateless container behind the OpenFaaS gateway. The gateway calls handle() for every invocation.')

h2('Full Code Explained')
code([
    'import json, joblib, os',
    '',
    'MODEL_PATH          = os.getenv("MODEL_PATH", "/app/ml/model.joblib")',
    'HEART_RATE_THRESHOLD = 100',
    'TEMP_THRESHOLD       = 38.0',
    'RR_THRESHOLD         = 24',
    '',
    'model = joblib.load(MODEL_PATH)',
    '',
    'def handle(event, context):',
    '    try:',
    '        body = json.loads(event.body)',
    '        heart_rate       = float(body["heart_rate"])',
    '        temperature      = float(body["temperature"])',
    '        respiratory_rate = float(body["respiratory_rate"])',
    '',
    '        prediction = int(model.predict([[heart_rate, temperature, respiratory_rate]])[0])',
    '',
    '        alerts = []',
    '        if heart_rate > HEART_RATE_THRESHOLD:',
    '            alerts.append("High Heart Rate Detected")',
    '        if temperature > TEMP_THRESHOLD:',
    '            alerts.append("High Fever Detected")',
    '        if respiratory_rate > RR_THRESHOLD:',
    '            alerts.append("Abnormal Respiratory Rate")',
    '',
    '        return {',
    '            "statusCode": 200,',
    '            "body": json.dumps({"sepsis_prediction": prediction, "alerts": alerts})',
    '        }',
    '    except (KeyError, ValueError) as e:',
    '        return {"statusCode": 400, "body": json.dumps({"error": f"Invalid input: {e}"})}',
    '    except Exception as e:',
    '        return {"statusCode": 500, "body": json.dumps({"error": str(e)})}',
])
var_table([
    ['MODEL_PATH', 'str (env var)', 'Path to model.joblib inside the container. Set via docker-compose volumes.'],
    ['HEART_RATE_THRESHOLD etc.', 'module constants', 'Threshold values. Defined at module level so they are easy to find and change.'],
    ['model', 'module-level', 'Loaded once when the container starts. Shared across all invocations — no per-request loading.'],
    ['event', 'object', 'Passed by OpenFaaS. Has a .body attribute containing the raw request bytes.'],
    ['context', 'object', 'Passed by OpenFaaS. Contains metadata like function name. Not used here but required by signature.'],
    ['event.body', 'bytes', 'The raw HTTP request body. json.loads() parses it into a dict.'],
    ['float(body["key"])', 'conversion', 'Explicit float conversion handles cases where JSON sends integers instead of floats.'],
    ['model.predict([[...]])[0]', 'ndarray access', '[0] gets the first (and only) prediction from the result array.'],
    ['statusCode', 'int', 'OpenFaaS reads this to set the HTTP response status. 200=OK, 400=bad input, 500=error.'],
    ['"body"', 'str', 'OpenFaaS uses this as the HTTP response body. Must be a string — hence json.dumps().'],
    ['KeyError / ValueError', 'exceptions', 'KeyError = missing field in JSON. ValueError = cannot convert to float. Both return 400.'],
])
tip('To add a new vital: add float(body["new_vital"]) extraction, add it to model.predict([[...]]), add a threshold check.')
tip('To change thresholds: edit the three constants at the top of the file — no need to touch the function logic.')

# ══════════════════════════════════════════════════════════════════════════════
# 10. functions/predict/entrypoint.py  [NEW]
# ══════════════════════════════════════════════════════════════════════════════
file_banner('functions/predict/entrypoint.py', 'OpenFaaS HTTP Wrapper  [NEW FILE]')
body('A thin Flask web server that wraps handler.py. OpenFaaS gateway needs to call the function over HTTP — this file provides that HTTP interface. It follows the OpenFaaS Classic Watchdog pattern where each function is an HTTP microservice.')

h2('Full Code Explained')
code([
    'from flask import Flask, request, Response',
    'from handler import handle',
    '',
    'app = Flask(__name__)',
    '',
    'class Event:',
    '    def __init__(self, body: bytes):',
    '        self.body = body',
    '',
    '@app.post("/")',
    'def invoke():',
    '    event = Event(body=request.data)',
    '    result = handle(event, context=None)',
    '    return Response(',
    '        result["body"],',
    '        status=result["statusCode"],',
    '        mimetype="application/json",',
    '    )',
    '',
    'if __name__ == "__main__":',
    '    app.run(host="0.0.0.0", port=5000)',
])
var_table([
    ['Flask(__name__)', 'Flask app', 'Creates the HTTP server. __name__ tells Flask the name of the current module.'],
    ['Event class', 'wrapper', 'Mimics the event object that real OpenFaaS passes to handle(). Only needs a .body attribute.'],
    ['request.data', 'bytes', 'Flask attribute — the raw bytes of the incoming HTTP request body.'],
    ['Event(body=request.data)', 'object', 'Wraps the raw bytes in an Event so handle() can access event.body consistently.'],
    ['context=None', 'None', 'Passes None as context since this wrapper does not provide OpenFaaS metadata.'],
    ['Response(...)', 'Flask response', 'Constructs the HTTP response using the statusCode and body from handle().'],
    ['mimetype="application/json"', 'str', 'Sets Content-Type header so the client knows to parse the response as JSON.'],
    ['host="0.0.0.0"', 'str', 'Listens on all network interfaces inside the container — required for Docker networking.'],
    ['port=5000', 'int', 'The port this container listens on. Must match what docker-compose routes to.'],
])
tip('To change the port: change 5000 here AND update the functions_provider_url in docker-compose.yml.')
doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 11. security/opa/policy.rego  [NEW]
# ══════════════════════════════════════════════════════════════════════════════
file_banner('security/opa/policy.rego', 'Open Policy Agent Authorization Rules  [NEW FILE]')
body('Written in Rego — OPA\'s policy language. This file defines WHO can call WHICH endpoint. OPA evaluates this policy for every request when auth.py asks "is this token allowed?"')

h2('Full Policy Explained')
code([
    'package aiopscare.authz',
    '',
    'default allow = false',
    '',
    'allow {',
    '    input.method == "GET"',
    '    input.path == "/"',
    '}',
    '',
    'allow {',
    '    input.method == "GET"',
    '    input.path == "/metrics"',
    '}',
    '',
    'allow {',
    '    input.method == "POST"',
    '    input.path == "/predict"',
    '    token.payload.realm_access.roles[_] == "icu-staff"',
    '}',
    '',
    'token := {"payload": payload} {',
    '    [_, encoded] := split(input.token, " ")',
    '    [_, payload, _] := io.jwt.decode(encoded)',
    '}',
])
var_table([
    ['package aiopscare.authz', 'declaration', 'Namespace for this policy. Maps to the URL /v1/data/aiopscare/authz/allow in OPA REST API.'],
    ['default allow = false', 'default rule', 'If no allow rule matches, result is false. Deny-by-default is secure.'],
    ['allow { ... }', 'rule block', 'Multiple allow blocks with different conditions. OPA returns true if ANY block fully matches.'],
    ['input.method', 'str', 'HTTP method sent by auth.py in the input object. Must match exactly.'],
    ['input.path', 'str', 'URL path sent by auth.py. Add new blocks to protect additional endpoints.'],
    ['input.token', 'str', 'Full "Bearer <jwt>" string sent by auth.py.'],
    ['split(input.token, " ")', 'Rego builtin', 'Splits "Bearer eyJ..." into ["Bearer", "eyJ..."]. [_, encoded] destructures the result.'],
    ['io.jwt.decode(encoded)', 'Rego builtin', 'Decodes the JWT without verifying signature. Returns [header, payload, signature].'],
    ['token.payload.realm_access.roles', 'JWT claim', 'Keycloak puts assigned roles here. [_] means "any element in this array".'],
    ['"icu-staff"', 'str', 'Required role name. Must match exactly what is defined in Keycloak realm-export.json.'],
])
tip('To protect a new endpoint: add a new allow { input.path == "/new-path" ... } block.')
tip('To add a new role: create it in Keycloak realm-export.json and add a new allow block checking for that role name.')
tip('To skip auth for an endpoint: add allow { input.path == "/health" } with no token check.')

# ══════════════════════════════════════════════════════════════════════════════
# 12. security/keycloak/realm-export.json  [NEW]
# ══════════════════════════════════════════════════════════════════════════════
file_banner('security/keycloak/realm-export.json', 'Keycloak Identity Configuration  [NEW FILE]')
body('JSON configuration file auto-imported by Keycloak at startup. Defines the authentication realm, API client, user roles, and test users. Think of it as the "database seed" for the identity system.')

h2('Structure Explained')
code([
    '{',
    '  "realm": "aiopscare",',
    '  "enabled": true,',
    '  "sslRequired": "none",',
    '',
    '  "clients": [{',
    '    "clientId": "aiopscare-api",',
    '    "secret": "aiopscare-secret",',
    '    "directAccessGrantsEnabled": true',
    '  }],',
    '',
    '  "roles": {',
    '    "realm": [',
    '      { "name": "icu-staff" },',
    '      { "name": "admin" }',
    '    ]',
    '  },',
    '',
    '  "users": [{',
    '    "username": "icu-user",',
    '    "credentials": [{ "type": "password", "value": "icu-password" }],',
    '    "realmRoles": ["icu-staff"]',
    '  }]',
    '}',
])
var_table([
    ['"realm"', 'str', 'Unique name for this Keycloak tenant. Appears in all token URLs: /realms/aiopscare/...'],
    ['"sslRequired": "none"', 'str', 'Disables HTTPS requirement for development. Change to "external" in production.'],
    ['"clientId"', 'str', 'The application identifier used when requesting tokens. Must match curl -d client_id=...'],
    ['"secret"', 'str', 'Client secret for confidential clients. Change this to a strong random string in production.'],
    ['"directAccessGrantsEnabled"', 'bool', 'Allows username/password token requests (Resource Owner Password flow). Convenient for testing.'],
    ['"realm" roles', 'list', 'Roles available in this realm. Add new roles here as objects with "name" field.'],
    ['"icu-staff"', 'str', 'Must match the role name checked in policy.rego. If you rename it, update both files.'],
    ['"username" / "credentials"', 'str', 'Test user for development. Remove or change credentials before production deployment.'],
    ['"realmRoles"', 'list', 'Roles assigned to this user. These appear in the JWT under realm_access.roles.'],
])
tip('To add a new user: copy the users array entry and change username, password, and realmRoles.')
tip('To add a new role: add an object to realm.roles array AND assign it to users via realmRoles.')
doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 13. infrastructure/docker-compose.yml  [UPDATED]
# ══════════════════════════════════════════════════════════════════════════════
file_banner('infrastructure/docker-compose.yml', 'All Services Orchestration  [UPDATED]')
body('Defines all 11 services that make up the system. Docker Compose reads this file and starts every container with the right configuration, networking, and dependencies.')

h2('Key Concepts')
var_table([
    ['image:', 'str', 'Pre-built Docker image to pull from registry. No build step needed.'],
    ['build: context:', 'path', 'Build a custom image from a Dockerfile in the given directory.'],
    ['depends_on:', 'list', 'Start this service only after listed services have started (not health-checked — just started).'],
    ['environment:', 'map', 'Sets environment variables inside the container. These are read by os.getenv() in Python code.'],
    ['volumes:', 'list', 'Mounts host files/directories into the container. format: host_path:container_path'],
    ['ports:', 'list', 'Maps host port to container port. format: "host:container". Remove to hide service from host.'],
    ['restart: unless-stopped', 'policy', 'Automatically restarts the container if it crashes, until you explicitly stop it.'],
])

h2('Kafka — Dual Listener Configuration')
code([
    'KAFKA_LISTENERS: PLAINTEXT://0.0.0.0:9092,PLAINTEXT_INTERNAL://0.0.0.0:29092',
    'KAFKA_ADVERTISED_LISTENERS: PLAINTEXT://localhost:9092,PLAINTEXT_INTERNAL://kafka:29092',
    'KAFKA_LISTENER_SECURITY_PROTOCOL_MAP: PLAINTEXT:PLAINTEXT,PLAINTEXT_INTERNAL:PLAINTEXT',
    'KAFKA_INTER_BROKER_LISTENER_NAME: PLAINTEXT_INTERNAL',
])
var_table([
    ['PLAINTEXT://localhost:9092', 'external listener', 'For access from your host machine (e.g., running producer.py locally).'],
    ['PLAINTEXT_INTERNAL://kafka:29092', 'internal listener', 'For access from other Docker containers. "kafka" resolves to the Kafka container IP inside Docker.'],
    ['KAFKA_INTER_BROKER_LISTENER_NAME', 'str', 'Which listener Kafka brokers use to talk to each other. Must be one of the defined listener names.'],
])
note('This was the critical bug in the original config. localhost:9092 inside a container refers to that container itself, not Kafka.')

h2('Keycloak Service')
code([
    'keycloak:',
    '  image: quay.io/keycloak/keycloak:24.0.3',
    '  command: start-dev --import-realm',
    '  environment:',
    '    KC_DB: postgres',
    '    KC_DB_URL: jdbc:postgresql://postgres:5432/keycloak',
    '    KEYCLOAK_ADMIN: admin',
    '    KEYCLOAK_ADMIN_PASSWORD: admin',
    '  volumes:',
    '    - ../security/keycloak:/opt/keycloak/data/import',
])
var_table([
    ['start-dev', 'command', 'Starts Keycloak in development mode (no HTTPS required). Change to start for production.'],
    ['--import-realm', 'flag', 'Tells Keycloak to import JSON files from the /opt/keycloak/data/import directory on startup.'],
    ['KC_DB_URL', 'str', 'JDBC connection string. postgres is the Docker service name — resolves automatically.'],
    ['/opt/keycloak/data/import', 'mount target', 'The directory Keycloak reads realm JSON files from. Mapped from ../security/keycloak on host.'],
])

h2('OPA Service')
code([
    'opa:',
    '  image: openpolicyagent/opa:latest',
    '  command: run --server --log-level=info /policies',
    '  volumes:',
    '    - ../security/opa:/policies',
])
var_table([
    ['run --server', 'command', 'Starts OPA in server mode — exposes REST API on port 8181.'],
    ['--log-level=info', 'flag', 'Logs every policy decision. Change to debug for more detail or error for less noise.'],
    ['/policies', 'path', 'Directory OPA loads .rego files from. Mapped from ../security/opa on host.'],
])

h2('OpenFaaS Gateway Service')
code([
    'openfaas-gateway:',
    '  image: ghcr.io/openfaas/gateway:0.27.3',
    '  environment:',
    '    basic_auth: "false"',
    '    functions_provider_url: "http://predict-fn:5000/"',
    '    direct_functions: "true"',
])
var_table([
    ['basic_auth: "false"', 'str', 'Disables gateway authentication for development. Enable and set credentials for production.'],
    ['functions_provider_url', 'str', 'Where the gateway routes function invocations to. Points to our predict-fn Flask server.'],
    ['direct_functions: "true"', 'str', 'Gateway calls functions directly via HTTP without an additional provider layer.'],
])

h2('Named Volumes')
code([
    'volumes:',
    '  kafka-data:',
    '  postgres-data:',
    '  prometheus-data:',
    '  grafana-data:',
])
body('Named volumes persist data across docker compose down/up. Without these, Kafka messages, Keycloak users, Prometheus metrics, and Grafana dashboards would be lost every restart.')
tip('To wipe all data and start fresh: docker compose down -v (the -v flag removes volumes).')
doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 14. monitoring/prometheus/prometheus.yml  [UPDATED]
# ══════════════════════════════════════════════════════════════════════════════
file_banner('monitoring/prometheus/prometheus.yml', 'Prometheus Metrics Config  [UPDATED]')

h2('Full Config Explained')
code([
    'global:',
    '  scrape_interval: 15s',
    '',
    'scrape_configs:',
    '  - job_name: "fastapi"',
    '    static_configs:',
    '      - targets: ["fastapi:8000"]',
    '',
    '  - job_name: "openfaas-gateway"',
    '    static_configs:',
    '      - targets: ["openfaas-gateway:8080"]',
])
var_table([
    ['scrape_interval: 15s', 'duration', 'How often Prometheus polls each target for metrics. 15s is standard. Use 5s for real-time dashboards.'],
    ['job_name', 'str', 'Label applied to all metrics from this target. Appears in Grafana as a filter dimension.'],
    ['targets', 'list[str]', 'host:port of services to scrape. Must use Docker service names inside Compose networks.'],
    ['"fastapi:8000"', 'str', 'Was hardcoded IP 172.21.100.237:8000. Now uses Docker service name — works on any machine.'],
    ['"openfaas-gateway:8080"', 'str', 'Added to also scrape function invocation metrics from the OpenFaaS gateway.'],
])
tip('To monitor a new service: add a new - job_name block with the service name and port.')
tip('To scrape a service outside Docker: use its actual IP or hostname in targets.')

# ══════════════════════════════════════════════════════════════════════════════
# 15. scripts/zap_scan.sh  [NEW]
# ══════════════════════════════════════════════════════════════════════════════
file_banner('scripts/zap_scan.sh', 'OWASP ZAP Security Scan  [NEW FILE]')
body('Shell script that runs an automated security scan against the FastAPI API using OWASP ZAP — the industry standard web application security scanner. Saves an HTML report documenting any vulnerabilities found.')

h2('Full Script Explained')
code([
    '#!/usr/bin/env bash',
    'TARGET=${1:-"http://localhost:8000"}',
    'REPORT_DIR="docs/security"',
    'REPORT_FILE="$REPORT_DIR/zap-report.html"',
    '',
    'mkdir -p "$REPORT_DIR"',
    '',
    'docker run --rm \\',
    '  --network host \\',
    '  -v "$(pwd)/$REPORT_DIR:/zap/wrk:rw" \\',
    '  ghcr.io/zaproxy/zaproxy:stable \\',
    '  zap-baseline.py \\',
    '    -t "$TARGET" \\',
    '    -r "zap-report.html" \\',
    '    -I',
])
var_table([
    ['#!/usr/bin/env bash', 'shebang', 'Tells the OS to run this script with Bash.'],
    ['${1:-"http://localhost:8000"}', 'variable', 'First argument to script, defaulting to localhost:8000. Run as: bash zap_scan.sh http://other-host:8000'],
    ['REPORT_DIR', 'str', 'Where to save the HTML report. Creates docs/security/ directory if it does not exist.'],
    ['docker run --rm', 'flag', 'Runs ZAP in a Docker container and removes the container when done.'],
    ['--network host', 'flag', 'Gives ZAP container access to the host network so it can reach localhost:8000.'],
    ['-v $(pwd)/...:/zap/wrk:rw', 'volume mount', 'Mounts local report directory into ZAP container so the report file is saved to host.'],
    ['zap-baseline.py', 'ZAP script', 'Runs a passive scan (does not attack, only observes responses). Safe to run against any server.'],
    ['-t "$TARGET"', 'flag', 'Target URL to scan.'],
    ['-r "zap-report.html"', 'flag', 'Output report filename inside /zap/wrk (which maps to our local report directory).'],
    ['-I', 'flag', 'Ignore warnings — exits with code 0 even if issues found. Remove to make CI fail on findings.'],
])
tip('To run a deeper active scan: replace zap-baseline.py with zap-full-scan.py — but this will send actual attack payloads.')
tip('To integrate with CI: remove -I flag so the script exits non-zero when vulnerabilities are found.')

# ══════════════════════════════════════════════════════════════════════════════
# 16. Dockerfile  [FILLED]
# ══════════════════════════════════════════════════════════════════════════════
file_banner('Dockerfile', 'FastAPI Container Image  [FILLED FROM EMPTY]')
body('Defines how to build the Docker image for the FastAPI application. Every instruction creates a layer in the image. Docker caches layers — if requirements.txt does not change, pip install is skipped on rebuild.')

h2('Full Dockerfile Explained')
code([
    'FROM python:3.11-slim',
    '',
    'WORKDIR /app',
    '',
    'COPY requirements.txt .',
    'RUN pip install --no-cache-dir -r requirements.txt',
    '',
    'COPY . .',
    '',
    'RUN python ml/train.py',
    '',
    'EXPOSE 8000',
    '',
    'CMD ["uvicorn", "app.main:app", "--host", "0.0.0.0", "--port", "8000"]',
])
var_table([
    ['FROM python:3.11-slim', 'base image', 'Starts from an official Python image. slim reduces image size by omitting docs and test files.'],
    ['WORKDIR /app', 'instruction', 'Sets working directory inside container. All subsequent paths are relative to /app.'],
    ['COPY requirements.txt .', 'instruction', 'Copies requirements.txt first — before the rest of the code. This way Docker caches the pip install layer separately.'],
    ['--no-cache-dir', 'pip flag', 'Skips pip download cache — reduces image size since the cache is not needed after build.'],
    ['COPY . .', 'instruction', 'Copies everything else (app/, ml/, events/, etc.) into /app. Runs after pip install so code changes do not bust the dependency cache layer.'],
    ['RUN python ml/train.py', 'instruction', 'Trains and saves model.joblib inside the image at build time. The model is baked into the image.'],
    ['EXPOSE 8000', 'instruction', 'Documents that the container listens on port 8000. Does not actually open the port — that is done by ports: in docker-compose.yml.'],
    ['CMD [...]', 'instruction', 'Default command when container starts. uvicorn starts the ASGI server. --host 0.0.0.0 required for Docker networking.'],
])
tip('To rebuild after code changes: docker compose up --build. The pip install layer is cached and skipped if requirements.txt did not change.')
tip('To use a different Python version: change 3.11-slim to 3.12-slim.')

# ══════════════════════════════════════════════════════════════════════════════
# FOOTER
# ══════════════════════════════════════════════════════════════════════════════
sec = doc.sections[0]
footer = sec.footer
fp = footer.paragraphs[0]
fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = fp.add_run('AIOpsCare Master Code Notes  |  Final Year Project — Topic 7')
r.font.name='Calibri'; r.font.size=Pt(8)
r.font.color.rgb = RGBColor(0x80,0x80,0x80)

output = r'C:\Users\Acer\AIOpsCare_Master_Code_Notes.docx'
doc.save(output)
print(f'Saved: {output}')
